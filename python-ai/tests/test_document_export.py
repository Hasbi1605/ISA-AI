import os
import sys
from io import BytesIO

import pytest
from docx import Document as DocxDocument
from fastapi.testclient import TestClient
from openpyxl import load_workbook

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
os.environ["AI_SERVICE_TOKEN"] = "test_internal_api_secret"

from app.documents_api import app
from app.services.document_export import export_content


client = TestClient(app)
AUTH_HEADERS = {"Authorization": "Bearer test_internal_api_secret"}


def _sample_html() -> str:
    return """
    <article>
      <h2>Ringkasan</h2>
      <p>Ini adalah jawaban contoh.</p>
      <table>
        <tr><th>Nama</th><th>Nilai</th></tr>
        <tr><td>A</td><td>10</td></tr>
        <tr><td>B</td><td>20</td></tr>
      </table>
    </article>
    """


def test_export_content_creates_pdf_docx_xlsx_and_csv():
    html = _sample_html()

    pdf = export_content(html, "pdf", "jawaban-ai")
    assert pdf.filename == "jawaban-ai.pdf"
    assert pdf.mime_type == "application/pdf"
    assert pdf.content.startswith(b"%PDF")

    docx = export_content(html, "docx", "jawaban-ai")
    assert docx.filename == "jawaban-ai.docx"
    assert docx.mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    docx_document = DocxDocument(BytesIO(docx.content))
    paragraphs = [paragraph.text for paragraph in docx_document.paragraphs if paragraph.text.strip()]
    assert any("Jawaban Ai" in paragraph or "jawaban ai" in paragraph.lower() for paragraph in paragraphs)
    assert any("Ini adalah jawaban contoh." in paragraph for paragraph in paragraphs)

    xlsx = export_content(html, "xlsx", "jawaban-ai")
    assert xlsx.filename == "jawaban-ai.xlsx"
    workbook = load_workbook(BytesIO(xlsx.content))
    sheet = workbook.active
    assert sheet["A1"].value == "Nama"
    assert sheet["B2"].value == "10"

    csv_export = export_content(html, "csv", "jawaban-ai")
    assert csv_export.filename == "jawaban-ai.csv"
    assert b"Nama" in csv_export.content
    assert b"10" in csv_export.content


def test_documents_export_route_returns_binary_download():
    response = client.post(
        "/api/documents/export",
        headers=AUTH_HEADERS,
        json={
            "content_html": _sample_html(),
            "target_format": "pdf",
            "file_name": "jawaban-ai",
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("application/pdf")
    assert response.headers["content-disposition"].startswith('attachment; filename="jawaban-ai.pdf"')
    assert response.content.startswith(b"%PDF")


def test_export_content_rejects_external_resource_urls(monkeypatch):
    calls = {"count": 0}

    class FakeWeasyHTML:
        def __init__(self, *, string: str, url_fetcher=None):
            self.string = string
            self.url_fetcher = url_fetcher

        def write_pdf(self) -> bytes:
            calls["count"] += 1
            return b"%PDF-1.4 fake"

    monkeypatch.setattr("app.services.document_export.WeasyHTML", FakeWeasyHTML)

    dangerous_html = """
    <article>
      <p>Konten tidak aman.</p>
      <img src="file:///etc/passwd" />
      <a href="file:///etc/passwd">Lampiran lokal</a>
    </article>
    """

    with pytest.raises(ValueError):
        export_content(dangerous_html, "pdf", "jawaban-ai")


def test_documents_export_route_rejects_external_resource_urls(monkeypatch):
    class FakeWeasyHTML:
        def __init__(self, *, string: str, url_fetcher=None):
            self.string = string
            self.url_fetcher = url_fetcher

        def write_pdf(self) -> bytes:
            return b"%PDF-1.4 fake"

    monkeypatch.setattr("app.services.document_export.WeasyHTML", FakeWeasyHTML)

    response = client.post(
        "/api/documents/export",
        headers=AUTH_HEADERS,
        json={
            "content_html": "<img src=\"file:///etc/passwd\"><a href=\"file:///etc/passwd\">x</a>",
            "target_format": "pdf",
            "file_name": "jawaban-ai",
        },
    )

    assert response.status_code == 400


def test_documents_export_route_rejects_oversized_html():
    response = client.post(
        "/api/documents/export",
        headers=AUTH_HEADERS,
        json={
            "content_html": "A" * 512001,
            "target_format": "pdf",
            "file_name": "jawaban-ai",
        },
    )

    assert response.status_code in {400, 422}


def test_documents_extract_tables_route_reads_upload():
    html = _sample_html()
    pdf_artifact = export_content(html, "pdf", "sample-table")

    response = client.post(
        "/api/documents/extract-tables",
        headers=AUTH_HEADERS,
        files={"file": ("sample-table.pdf", pdf_artifact.content, "application/pdf")},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "success"
    assert payload["tables"]
    first_table = payload["tables"][0]
    assert first_table["header"] == ["Nama", "Nilai"]
    assert first_table["rows"][0] == ["A", "10"]


@pytest.mark.parametrize("filename", ["../berbahaya.pdf", "/etc/passwd"])
def test_documents_extract_tables_route_rejects_unsafe_filename(monkeypatch, filename):
    class DummyFile(BytesIO):
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

    monkeypatch.setattr("builtins.open", lambda *_args, **_kwargs: DummyFile())
    monkeypatch.setattr("app.routers.documents.extract_tables_from_file", lambda *_args, **_kwargs: [])

    response = client.post(
        "/api/documents/extract-tables",
        headers=AUTH_HEADERS,
        files={"file": (filename, b"%PDF-1.4 fake", "application/pdf")},
    )

    assert response.status_code == 400


def test_documents_extract_tables_route_supports_xlsx_upload():
    html = _sample_html()
    xlsx_artifact = export_content(html, "xlsx", "sample-table")

    response = client.post(
        "/api/documents/extract-tables",
        headers=AUTH_HEADERS,
        files={
            "file": (
                "sample-table.xlsx",
                xlsx_artifact.content,
                "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ),
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "success"
    assert payload["tables"]
    first_table = payload["tables"][0]
    assert first_table["header"] == ["Nama", "Nilai"]
    assert first_table["rows"][0] == ["A", "10"]


def test_documents_extract_content_route_returns_full_docx_html():
    document = DocxDocument()
    document.add_paragraph("Paragraf awal dokumen lengkap.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Nama"
    table.cell(0, 1).text = "Nilai"
    table.cell(1, 0).text = "A"
    table.cell(1, 1).text = "10"

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = client.post(
        "/api/documents/extract-content",
        headers=AUTH_HEADERS,
        files={
            "file": (
                "dokumen-lengkap.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "success"
    assert "Paragraf awal dokumen lengkap." in payload["content_html"]
    assert "<table>" in payload["content_html"]
    assert "Nama" in payload["content_html"]


@pytest.mark.parametrize("filename", ["../berbahaya.docx", "/etc/passwd"])
def test_documents_extract_content_route_rejects_unsafe_filename(monkeypatch, filename):
    class DummyFile(BytesIO):
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

    monkeypatch.setattr("builtins.open", lambda *_args, **_kwargs: DummyFile())
    calls = {"count": 0}

    def _fake_extract_document_content_html(*_args, **_kwargs):
        calls["count"] += 1
        return "<article></article>"

    monkeypatch.setattr("app.routers.documents.extract_document_content_html", _fake_extract_document_content_html)

    response = client.post(
        "/api/documents/extract-content",
        headers=AUTH_HEADERS,
        files={
            "file": (
                filename,
                b"PK\x03\x04 fake docx",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ),
        },
    )

    assert response.status_code == 400
    assert calls["count"] == 0


def test_export_content_uses_safe_url_fetcher(monkeypatch):
    captured = {}

    class FakeWeasyHTML:
        def __init__(self, *, string: str, url_fetcher=None):
            captured["url_fetcher"] = url_fetcher
            self.string = string

        def write_pdf(self) -> bytes:
            return b"%PDF-1.4 safe"

    monkeypatch.setattr("app.services.document_export.WeasyHTML", FakeWeasyHTML)

    export = export_content("<p>Halo</p>", "pdf", "jawaban-ai")

    assert export.content.startswith(b"%PDF")
    assert captured.get("url_fetcher") is not None


def test_documents_extract_content_route_supports_csv_for_file_conversion():
    response = client.post(
        "/api/documents/extract-content",
        headers=AUTH_HEADERS,
        files={
            "file": (
                "biaya.csv",
                b"Nama,Nilai\nA,10\n",
                "text/csv",
            ),
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["status"] == "success"
    assert "<table>" in payload["content_html"]
    assert "Nama" in payload["content_html"]
    assert "10" in payload["content_html"]

    pdf = export_content(payload["content_html"], "pdf", "biaya")
    assert pdf.filename == "biaya.pdf"
    assert pdf.content.startswith(b"%PDF")

    docx = export_content(payload["content_html"], "docx", "biaya")
    docx_document = DocxDocument(BytesIO(docx.content))
    table_text = "\n".join(
        cell.text
        for table in docx_document.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "Nama" in table_text
    assert "10" in table_text


def test_documents_delete_route_requires_user_scope(monkeypatch):
    captured = {}

    def fake_delete_document_vectors(filename, user_id=None, document_id=None, cleanup_legacy=False):
        captured["filename"] = filename
        captured["user_id"] = user_id
        return True, "deleted"

    monkeypatch.setattr("app.routers.documents.delete_document_vectors", fake_delete_document_vectors)

    response = client.delete(
        "/api/documents/agenda.pdf?user_id=user-42",
        headers=AUTH_HEADERS,
    )

    assert response.status_code == 200
    assert captured == {"filename": "agenda.pdf", "user_id": "user-42"}


def test_documents_delete_route_rejects_missing_user_scope(monkeypatch):
    monkeypatch.setattr(
        "app.routers.documents.delete_document_vectors",
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("should not delete")),
    )

    response = client.delete(
        "/api/documents/agenda.pdf",
        headers=AUTH_HEADERS,
    )

    assert response.status_code == 422


def test_documents_process_route_rejects_unsafe_filename(monkeypatch):
    def fail_if_called(*_args, **_kwargs):
        raise AssertionError("run_document_process should not be called")

    monkeypatch.setattr("app.routers.documents.run_document_process", fail_if_called)

    response = client.post(
        "/api/documents/process",
        headers=AUTH_HEADERS,
        data={"user_id": "user-1"},
        files={"file": ("../evil.pdf", b"%PDF-1.4 fake", "application/pdf")},
    )

    assert response.status_code == 400


def test_documents_process_route_cleans_temp_file_on_failure(monkeypatch):
    removed_paths = []
    exists_calls = {"count": 0}

    def fake_run(*_args, **_kwargs):
        raise RuntimeError("boom")

    def fake_exists(_path):
        exists_calls["count"] += 1
        return True

    def fake_remove(path):
        removed_paths.append(path)

    class DummyFile(BytesIO):
        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

    monkeypatch.setattr("app.routers.documents.run_document_process", fake_run)
    monkeypatch.setattr("app.routers.documents.os.path.exists", fake_exists)
    monkeypatch.setattr("app.routers.documents.os.remove", fake_remove)
    monkeypatch.setattr("builtins.open", lambda *_args, **_kwargs: DummyFile())

    response = client.post(
        "/api/documents/process",
        headers=AUTH_HEADERS,
        data={"user_id": "user-1"},
        files={"file": ("aman.pdf", b"%PDF-1.4 fake", "application/pdf")},
    )

    assert response.status_code == 500
    assert exists_calls["count"] >= 1
    assert removed_paths


def test_documents_extract_tables_rejects_upload_exceeding_max_size(monkeypatch):
    """extract-tables must return 413 when the uploaded file exceeds MAX_UPLOAD_BYTES."""
    import app.routers.documents as docs_module

    original_limit = docs_module.MAX_UPLOAD_BYTES
    monkeypatch.setattr(docs_module, "MAX_UPLOAD_BYTES", 10)  # 10-byte limit for the test

    oversized_content = b"A" * 11  # 1 byte over the limit

    response = client.post(
        "/api/documents/extract-tables",
        headers=AUTH_HEADERS,
        files={"file": ("big-file.pdf", oversized_content, "application/pdf")},
    )

    assert response.status_code == 413


def test_documents_extract_content_rejects_upload_exceeding_max_size(monkeypatch):
    """extract-content must return 413 when the uploaded file exceeds MAX_UPLOAD_BYTES."""
    import app.routers.documents as docs_module

    monkeypatch.setattr(docs_module, "MAX_UPLOAD_BYTES", 10)  # 10-byte limit for the test

    oversized_content = b"B" * 11  # 1 byte over the limit

    response = client.post(
        "/api/documents/extract-content",
        headers=AUTH_HEADERS,
        files={"file": ("big-doc.docx", oversized_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )

    assert response.status_code == 413


@pytest.mark.parametrize("dangerous_value", [
    "=CMD|'/C calc'!A0",
    "+1+1",
    "-1",
    "@SUM(A1:A2)",
    "\t=hidden",
    "\r=CR",
])
def test_export_csv_sanitizes_formula_injection_in_cell_values(dangerous_value):
    """CSV export must prefix formula-triggering cell values to prevent injection."""
    from app.services.document_export import _sanitize_cell_value, _FORMULA_PREFIXES

    # Verify the raw value is indeed dangerous.
    assert dangerous_value[0] in _FORMULA_PREFIXES

    sanitized = _sanitize_cell_value(dangerous_value)

    # After sanitization the value must start with a tab (safe prefix).
    assert sanitized.startswith("\t"), f"Expected tab prefix for {dangerous_value!r}, got {sanitized!r}"
    # The original text must still be present.
    assert dangerous_value in sanitized


def test_export_xlsx_formula_injection_cells_are_prefixed():
    """XLSX export must sanitize formula-triggering values in table rows and headers."""
    from io import BytesIO
    from openpyxl import load_workbook

    dangerous_html = """
    <table>
      <tr><th>=SUM(A1:A2)</th><th>Nilai</th></tr>
      <tr><td>+kalkulasi</td><td>100</td></tr>
      <tr><td>normal text</td><td>@formula</td></tr>
    </table>
    """

    xlsx = export_content(dangerous_html, "xlsx", "test-inject")
    workbook = load_workbook(BytesIO(xlsx.content))
    sheet = workbook.active

    # All formula-triggering cells must be prefixed with a tab.
    header_a1 = sheet["A1"].value or ""
    assert header_a1.startswith("\t"), f"Header A1 should be tab-prefixed: {header_a1!r}"
    assert "=SUM(A1:A2)" in header_a1

    row2_a = sheet["A2"].value or ""
    assert row2_a.startswith("\t"), f"Row2 A should be tab-prefixed: {row2_a!r}"

    row3_b = sheet["B3"].value or ""
    assert row3_b.startswith("\t"), f"Row3 B should be tab-prefixed: {row3_b!r}"

    # Safe values must remain unchanged.
    assert sheet["B1"].value == "Nilai"
    assert sheet["B2"].value == "100"
    assert sheet["A3"].value == "normal text"


def test_export_csv_formula_injection_cells_are_prefixed():
    """CSV export must sanitize formula-triggering values in table rows and headers."""
    dangerous_html = """
    <table>
      <tr><th>=HYPERLINK(\"http://evil.com\")</th><th>Safe</th></tr>
      <tr><td>-deleteAll</td><td>ok</td></tr>
    </table>
    """

    csv_export = export_content(dangerous_html, "csv", "test-inject")
    content = csv_export.content.decode("utf-8")

    # Formula-starting values must be prefixed with a tab in the CSV output.
    assert "\t=HYPERLINK" in content, f"Formula header not sanitized: {content!r}"
    assert "\t-deleteAll" in content, f"Formula row not sanitized: {content!r}"

    # Safe values must remain unchanged.
    assert "Safe" in content
    assert "ok" in content
