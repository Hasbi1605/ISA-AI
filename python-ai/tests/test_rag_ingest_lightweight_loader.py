import csv

from docx import Document as DocxDocument
from openpyxl import Workbook

from app.services.rag_ingest import _load_documents_lightweight


def test_lightweight_loader_reads_docx_without_unstructured(tmp_path):
    path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Ringkasan proyek ISTA AI")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Nama"
    table.cell(0, 1).text = "Nilai"
    table.cell(1, 0).text = "Akurasi"
    table.cell(1, 1).text = "97%"
    document.save(path)

    docs = _load_documents_lightweight(str(path), path.name)

    assert len(docs) == 1
    assert "Ringkasan proyek ISTA AI" in docs[0].page_content
    assert "Akurasi | 97%" in docs[0].page_content


def test_lightweight_loader_reads_xlsx_and_csv(tmp_path):
    xlsx_path = tmp_path / "biaya.xlsx"
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Personil"
    sheet.append(["Nama", "Total"])
    sheet.append(["Hasbi", 12500000])
    workbook.save(xlsx_path)

    csv_path = tmp_path / "data.csv"
    with csv_path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["Kolom", "Isi"])
        writer.writerow(["Status", "Selesai"])

    xlsx_docs = _load_documents_lightweight(str(xlsx_path), xlsx_path.name)
    csv_docs = _load_documents_lightweight(str(csv_path), csv_path.name)

    assert len(xlsx_docs) == 1
    assert "Personil" in xlsx_docs[0].page_content
    assert "Hasbi | 12500000" in xlsx_docs[0].page_content
    assert len(csv_docs) == 1
    assert "Status | Selesai" in csv_docs[0].page_content


def test_lightweight_loader_rejects_unsupported_file_type(tmp_path):
    path = tmp_path / "archive.zip"
    path.write_bytes(b"not a supported document")

    try:
        _load_documents_lightweight(str(path), path.name)
    except ValueError as exc:
        assert "Unsupported document type" in str(exc)
    else:
        raise AssertionError("Unsupported file type should raise ValueError")
