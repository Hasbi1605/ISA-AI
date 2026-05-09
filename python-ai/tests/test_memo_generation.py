import os
import sys
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from fastapi.testclient import TestClient

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
os.environ["AI_SERVICE_TOKEN"] = "test_internal_api_secret"

from app.documents_api import app
from app.services.memo_generation import MemoDraft
from app.services.memo_generation import build_memo_prompt
from app.services.memo_generation import generate_memo_docx


def _all_document_text(document):
    texts = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    texts.append(cell.text)
    for section in document.sections:
        texts.extend(paragraph.text for paragraph in section.footer.paragraphs if paragraph.text)
    return "\n".join(texts)


def _find_table_containing(document, text):
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text == text:
                    return table
    raise AssertionError(f"Table containing {text!r} not found")


def _has_table_containing(document, text):
    try:
        _find_table_containing(document, text)
    except AssertionError:
        return False
    return True


def _table_indent_twips(table):
    table_indent = table._tbl.tblPr.find(qn("w:tblInd"))
    assert table_indent is not None
    return int(table_indent.get(qn("w:w")))


def _cell_margin_twips(cell, margin_name):
    tc_mar = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
    assert tc_mar is not None
    margin = tc_mar.find(qn(f"w:{margin_name}"))
    assert margin is not None
    return int(margin.get(qn("w:w")))


def _cell_width_twips(cell):
    width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
    assert width is not None
    return int(width.get(qn("w:w")))


def _spacing_after_table_twips(table):
    paragraph = table._tbl.getnext()
    assert paragraph is not None
    spacing = paragraph.find(qn("w:pPr")).find(qn("w:spacing"))
    assert spacing is not None
    return int(spacing.get(qn("w:after"), "0"))


def _signature_qr_center_ratio(document, table):
    indent_in = _table_indent_twips(table) / 1440
    page_width_in = document.sections[0].page_width.inches
    left_margin_in = document.sections[0].left_margin.inches
    qr_center_in = left_margin_in + indent_in + 0.35 + (0.86 / 2)
    return qr_center_in / page_width_in


def _signature_spacer_before_twips(table):
    spacer = table._tbl.getprevious()
    assert spacer is not None
    spacing = spacer.find(qn("w:pPr")).find(qn("w:spacing"))
    assert spacing is not None
    return int(spacing.get(qn("w:before")))


def _separator_after_twips(document):
    for paragraph in document.paragraphs:
        p_pr = paragraph._p.get_or_add_pPr()
        if p_pr.find(qn("w:pBdr")) is None:
            continue
        spacing = p_pr.find(qn("w:spacing"))
        assert spacing is not None
        return int(spacing.get(qn("w:after")))
    raise AssertionError("Separator paragraph not found")


def _paragraph_space_before_twips(paragraph):
    spacing = paragraph._p.get_or_add_pPr().find(qn("w:spacing"))
    assert spacing is not None
    return int(spacing.get(qn("w:before"), "0"))


def _find_paragraph(document, text):
    for paragraph in document.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise AssertionError(f"Paragraph {text!r} not found")


def test_generate_memo_docx_builds_official_memorandum_document():
    configuration = {
        "number": "M-02/I-Yog/IT.02/04/2026",
        "recipient": "Kepala Pusat Pengembangan dan Layanan Sistem Informasi",
        "sender": "Kepala Istana Kepresidenan Yogyakarta",
        "subject": "Penyampaian Nama PIC Aplikasi Virtual Meeting",
        "date": "3 April 2026",
        "basis": "Menindaklanjuti memorandum Bapak nomor M-01/PPLSI/IT.02.00/04/2026.",
        "content": "Sampaikan nama, NIP, pangkat/gol., dan jabatan PIC.",
        "closing": "Atas perhatian dan kerja sama Bapak, kami ucapkan terima kasih.",
        "signatory": "Deni Mulyana",
        "carbon_copy": "Kepala Subbagian Tata Usaha, Istana Kepresidenan Yogyakarta, Sekretariat Presiden",
        "page_size": "folio",
    }

    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penyampaian Nama PIC Aplikasi Virtual Meeting",
        context="Bahas PIC aplikasi virtual meeting.",
        text_generator=lambda prompt: "Mohon setiap unit menyiapkan laporan progres.\n- Bawa data pendukung.",
        configuration=configuration,
    )

    document = Document(BytesIO(draft.content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]

    assert draft.filename == "penyampaian-nama-pic-aplikasi-virtual-meeting.docx"
    assert paragraphs[:3] == [
        "KEMENTERIAN SEKRETARIAT NEGARA RI",
        "SEKRETARIAT PRESIDEN",
        "ISTANA KEPRESIDENAN YOGYAKARTA",
    ]
    assert "MEMORANDUM" in paragraphs
    assert "Nomor M-02/I-Yog/IT.02/04/2026" in paragraphs
    assert "Mohon setiap unit menyiapkan laporan progres." in paragraphs
    assert any("Bawa data pendukung." in paragraph for paragraph in paragraphs)
    assert "Atas perhatian dan kerja sama Bapak, kami ucapkan terima kasih." in paragraphs
    assert "Deni Mulyana" in _all_document_text(document)
    assert "Dokumen ini telah ditandatangani secara elektronik" in document.sections[0].footer.paragraphs[0].text
    assert document.sections[0].page_height == Inches(14)
    assert draft.page_size == "folio"
    assert document.styles["Normal"].font.name == "Arial"
    assert document.paragraphs[0].runs[0].font.name == "Arial"
    assert document.tables[0].cell(0, 2).text == configuration["recipient"]
    assert document.tables[0].cell(2, 2).text == configuration["subject"]
    qr_table = _find_table_containing(document, "QR\nTTE")
    assert qr_table.cell(0, 1).text == "QR\nTTE"
    assert qr_table.cell(1, 0).text == "Deni Mulyana"
    assert 5550 <= _table_indent_twips(qr_table) <= 5700
    assert 0.67 <= _signature_qr_center_ratio(document, qr_table) <= 0.72
    assert "Penyampaian Nama PIC Aplikasi Virtual Meeting" in draft.searchable_text


def test_generate_memo_docx_adds_formal_fallback_closing_when_blank():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Permohonan Data Rapat",
        context="Mohon dibuatkan memo permohonan data rapat.",
        text_generator=lambda prompt: "Mohon unit terkait menyiapkan data rapat paling lambat 10 Mei 2026.",
        configuration={
            "number": "M-04/I-Yog/UM.01/05/2026",
            "recipient": "Kepala Subbagian Tata Usaha",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Permohonan Data Rapat",
            "date": "6 Mei 2026",
            "content": "Minta data rapat paling lambat 10 Mei 2026.",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text]
    fallback_closing = "Demikian disampaikan untuk menjadi perhatian dan tindak lanjut sebagaimana mestinya."

    assert "Demikian, mohon arahan lebih lanjut." not in paragraphs
    assert "Demikian, mohon arahan lebih lanjut." not in draft.searchable_text
    assert fallback_closing in paragraphs
    assert fallback_closing in draft.searchable_text
    assert "Deni Mulyana" in _all_document_text(document)


def test_generate_memo_docx_moves_generated_closing_to_official_closing_block():
    closing = "Demikian disampaikan, atas perhatian dan kerja samanya diucapkan terima kasih."

    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penyampaian Data Pegawai Pendamping Kegiatan",
        context="Data pegawai pendamping kegiatan.",
        text_generator=lambda prompt: (
            "Sehubungan dengan kebutuhan pendampingan kegiatan integrasi aplikasi, "
            "dapat kami sampaikan data pegawai pendamping sebagai berikut.\n"
            "nama: Muhammad Hasbi Ash Shiddiqi\n"
            "NIP: 231210013\n"
            f"{closing}"
        ),
        configuration={
            "number": "EVAL-11/IST/YK/05/2026",
            "recipient": "Kepala Bagian SDM",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penyampaian Data Pegawai Pendamping Kegiatan",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    closing_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text == closing]

    assert len(closing_paragraphs) == 1
    assert _paragraph_space_before_twips(closing_paragraphs[0]) >= 280
    assert draft.searchable_text.count(closing) == 1


def test_generate_memo_docx_rejects_ai_unavailable_fallback_message():
    try:
        generate_memo_docx(
            memo_type="memo_internal",
            title="Revisi Ubah Penerima Memo",
            context="Revisi penerima memo.",
            text_generator=lambda prompt: "❌ Maaf, semua layanan AI sedang tidak tersedia. Silakan coba lagi nanti.",
            configuration={
                "number": "EVAL-32/IST/YK/05/2026",
                "recipient": "Kepala Bagian Administrasi",
                "sender": "Kepala Istana Kepresidenan Yogyakarta",
                "subject": "Revisi Ubah Penerima Memo",
                "date": "7 Mei 2026",
                "signatory": "Deni Mulyana",
                "page_size": "letter",
            },
        )
    except ValueError as exc:
        assert "Layanan AI" in str(exc)
    else:
        raise AssertionError("Expected AI unavailable fallback to be rejected")


def test_generate_memo_docx_auto_page_size_uses_generated_body_length():
    long_body = "\n".join(
        [
            "1. " + "Koordinasi lintas unit perlu dilakukan secara tertib dan terdokumentasi. " * 4,
            "2. " + "Setiap unit diminta menyiapkan data pendukung dan batas waktu pelaksanaan. " * 4,
            "3. " + "Hasil pembahasan disampaikan kembali sebagai bahan tindak lanjut pimpinan. " * 4,
        ]
    )

    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Koordinasi Kegiatan",
        context="Buat memo koordinasi kegiatan.",
        text_generator=lambda prompt: long_body,
        configuration={
            "number": "M-05/I-Yog/UM.01/05/2026",
            "recipient": "Kepala Unit Terkait",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Koordinasi Kegiatan",
            "date": "7 Mei 2026",
            "content": "Buat memo koordinasi kegiatan.",
            "signatory": "Deni Mulyana",
            "page_size": "auto",
            "page_size_mode": "auto",
        },
    )

    document = Document(BytesIO(draft.content))

    assert draft.page_size == "folio"
    assert document.sections[0].page_height == Inches(14)


def test_generate_memo_docx_auto_page_size_keeps_short_many_copies_on_letter():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penyampaian Data Singkat",
        context="Memo singkat dengan banyak tembusan.",
        text_generator=lambda prompt: "Sehubungan dengan kebutuhan koordinasi, data singkat disampaikan untuk menjadi perhatian.",
        configuration={
            "number": "EVAL-24/IST/YK/05/2026",
            "recipient": "Kepala Bagian Administrasi",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penyampaian Data Singkat",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "carbon_copy": "\n".join([
                "Kepala Bagian Tata Usaha",
                "Kepala Subbagian Kepegawaian",
                "Kepala Subbagian Rumah Tangga",
                "Koordinator Layanan",
                "Arsip",
            ]),
            "page_size": "auto",
            "page_size_mode": "auto",
        },
    )

    document = Document(BytesIO(draft.content))

    assert draft.page_size == "letter"
    assert document.sections[0].page_height == Inches(11)


def test_build_memo_prompt_keeps_ai_inside_body_scope():
    prompt = build_memo_prompt(
        memo_type="memo_internal",
        title="Permohonan Penempatan Jabatan Fungsional Pranata Komputer",
        context="Mohon penempatan pegawai tetap di Istana Kepresidenan Yogyakarta.",
        configuration={
            "number": "M-09/I-Yog/KP/03/2025",
            "recipient": "Deputi Bidang Administrasi dan Pengelolaan Istana",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "date": "17 Maret 2025",
            "content": "Tulis 3 poin bernomor.",
        },
    )

    assert "Nomor: M-09/I-Yog/KP/03/2025" in prompt
    assert "Yth.: Deputi Bidang Administrasi dan Pengelolaan Istana" in prompt
    assert "Tulis hanya isi utama memo" in prompt
    assert "tanpa kop, nomor, Yth., Dari, Hal, Tanggal" in prompt
    assert "Arahan tambahan:" in prompt
    assert "kecuali user mengisinya di field Penutup" in prompt
    assert "kontrol kerja, bukan bagian naskah" in prompt
    assert "Sehubungan hal tersebut" in prompt


def test_build_memo_prompt_prioritizes_revision_instruction_and_current_context():
    prompt = build_memo_prompt(
        memo_type="memo_internal",
        title="Penyampaian Keberatan Untuk Keperluan tersebut",
        context="Isi memo saat ini:\nTembusan:\n1. Kepala A\n2. Kepala B\n\nInstruksi revisi wajib diterapkan:\ntambahkan tembusan nomor 3, Kepala C",
        configuration={
            "number": "M/2312/22D/409L/YK",
            "recipient": "Kepala Komdigi",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "date": "6 Mei 2026",
            "content": "Isi lama dari konfigurasi.",
            "revision_instruction": "tambahkan tembusan nomor 3, Kepala C",
        },
    )

    assert "Isi memo saat ini:" in prompt
    assert "Isi lama dari konfigurasi." not in prompt
    assert "Instruksi revisi wajib diterapkan:" in prompt
    assert "tambahkan tembusan nomor 3, Kepala C" in prompt
    assert "Jangan meregenerasi seluruh memo" in prompt
    assert "ubah hanya bagian yang disebut" in prompt
    assert "baseline, uji, skenario evaluasi, dan auto format" in prompt
    assert "Jangan menulis blok Tembusan" in prompt


def test_generate_memo_docx_sanitizes_evaluation_artifacts_and_forbidden_sections():
    closing = "Demikian disampaikan, atas perhatian dan kerja samanya diucapkan terima kasih."
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Revisi Ubah Penutup",
        context="Revisi memo evaluasi.",
        text_generator=lambda prompt: (
            "Sehubungan dengan persiapan kegiatan, dimohon melakukan koordinasi.\n"
            "Memo ini juga sebagai baseline untuk uji revisi penutup.\n"
            f"{closing}\n"
            f"{closing}\n"
            "Tembusan:\n"
            "Kepala Bagian Keamanan"
        ),
        configuration={
            "number": "EVAL-33/IST/YK/05/2026",
            "recipient": "Kepala Bagian Protokol",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Revisi Ubah Penutup",
            "date": "7 Mei 2026",
            "closing": closing,
            "signatory": "Deni Mulyana",
            "carbon_copy": "Kepala Bagian Keamanan",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert "baseline" not in all_text.lower()
    assert "uji revisi" not in all_text.lower()
    assert all_text.count(closing) == 1
    assert all_text.count("Tembusan:") == 1
    assert all_text.count("Kepala Bagian Keamanan") == 1


def test_generate_memo_docx_strips_dangling_closing_fragment():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Kebutuhan Teknis Ruang Rapat Hybrid",
        context="Kebutuhan teknis ruang rapat hybrid.",
        text_generator=lambda prompt: (
            "Sehubungan dengan kebutuhan pelaksanaan rapat hybrid, diminta agar unit terkait "
            "menyiapkan perangkat audio, kamera, dan jaringan pendukung. Dengan"
        ),
        configuration={
            "number": "EVAL-15/IST/YK/05/2026",
            "recipient": "Kepala Subbagian Tata Usaha",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Kebutuhan Teknis Ruang Rapat Hybrid",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    all_text = _all_document_text(Document(BytesIO(draft.content)))

    assert "Dengan\n" not in all_text
    assert "\nDengan\n" not in draft.searchable_text
    assert "jaringan pendukung" in draft.searchable_text


def test_generate_memo_docx_strips_revision_and_additional_instruction_artifacts():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Revisi Typo Nama Orang",
        context="Revisi typo nama orang.",
        text_generator=lambda prompt: (
            "Menindaklanjuti proses pendataan pegawai pendamping kegiatan integrasi aplikasi, "
            "terdapat kekeliruan penulisan identitas yang perlu diperbaiki. "
            "Adapun data pegawai yang benar adalah sebagai berikut:\n"
            "nama: Muhammad Hasbi Ash Shiddiqi\n"
            "NIP: 231210013\n"
            "Perbaikan data tersebut dilakukan untuk keperluan pendampingan kegiatan integrasi aplikasi. "
            "Perbaiki typo nama menjadi Muhammad Hasbi Ash Shiddiqi, bagian lain jangan diubah.\n"
            "Hal ini disampaikan dengan mempertahankan seluruh data orang tersebut di atas tanpa perubahan."
        ),
        configuration={
            "number": "EVAL-38/IST/YK/05/2026",
            "recipient": "Kepala Bagian SDM",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Revisi Typo Nama Orang",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "carbon_copy": "Kepala Subbagian Kepegawaian",
            "page_size": "letter",
            "revision_instruction": "Perbaiki typo nama menjadi Muhammad Hasbi Ash Shiddiqi, bagian lain jangan diubah.",
            "additional_instruction": "Pertahankan seluruh data orang tanpa perubahan.",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert "Muhammad Hasbi Ash Shiddiqi" in all_text
    assert "231210013" in all_text
    assert "perbaiki typo" not in all_text.lower()
    assert "bagian lain jangan" not in all_text.lower()
    assert "pertahankan seluruh data" not in all_text.lower()
    assert "tanpa perubahan" not in all_text.lower()
    assert "perbaiki typo" not in draft.searchable_text.lower()


def test_generate_memo_docx_removes_carbon_copy_from_body_but_keeps_configuration_copy():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Revisi Kapitalisasi Tembusan",
        context="Kapitalisasi tembusan.",
        text_generator=lambda prompt: (
            "Menindaklanjuti kebutuhan distribusi informasi, diminta agar:\n"
            "1. Informasi disampaikan kepada unit terkait.\n"
            "2. Setiap unit memastikan tindak lanjut.\n"
            "Tembusan:\n"
            "Kementrian Polisi\n"
            "Kemenkue\n"
            "Purbaya Yudhi Sadewa"
        ),
        configuration={
            "number": "EVAL-40/IST/YK/05/2026",
            "recipient": "Kepala Bagian Administrasi",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Revisi Kapitalisasi Tembusan",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "carbon_copy": "Kementrian Polisi\nKemenkue\nPurbaya Yudhi Sadewa",
            "page_size": "folio",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert all_text.count("Tembusan:") == 1
    assert all_text.count("Kementrian Polisi") == 1
    assert all_text.count("Kemenkue") == 1
    assert all_text.count("Purbaya Yudhi Sadewa") == 1


def test_generate_memo_docx_removes_bare_configured_carbon_copy_line_from_body():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Revisi Kapitalisasi Tembusan",
        context="Kapitalisasi tembusan.",
        text_generator=lambda prompt: (
            "Menindaklanjuti kebutuhan distribusi informasi, diminta agar setiap unit memastikan tindak lanjut.\n"
            "Kepala Bagian Keamanan"
        ),
        configuration={
            "number": "EVAL-40/IST/YK/05/2026",
            "recipient": "Kepala Bagian Administrasi",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Revisi Kapitalisasi Tembusan",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "carbon_copy": "Kepala Bagian Keamanan",
            "page_size": "letter",
        },
    )

    all_text = _all_document_text(Document(BytesIO(draft.content)))

    assert all_text.count("Kepala Bagian Keamanan") == 1
    assert draft.searchable_text.count("Kepala Bagian Keamanan") == 1


def test_generate_memo_docx_uses_body_override_without_calling_text_generator():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Revisi Ubah Penerima Memo",
        context="Isi fallback tidak boleh dipakai.",
        text_generator=lambda prompt: (_ for _ in ()).throw(AssertionError("text generator should not be called")),
        configuration={
            "number": "EVAL-32/IST/YK/05/2026",
            "recipient": "Kepala Pusat Pengembangan Layanan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Revisi Ubah Penerima Memo",
            "date": "7 Mei 2026",
            "body_override": "Isi memo saat ini tetap dipertahankan.",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    assert "Isi memo saat ini tetap dipertahankan." in draft.searchable_text


def test_generate_memo_docx_treats_mohon_tindak_lanjut_as_generated_closing():
    closing = "Mohon tindak lanjut sesuai poin-poin tersebut agar proses pembaruan aplikasi dapat terlaksana dengan lancar dan terkoordinasi dengan baik."

    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Koordinasi Teknis Pembaruan Aplikasi Internal",
        context="Koordinasi pembaruan aplikasi internal.",
        text_generator=lambda prompt: (
            "Sehubungan dengan rencana pembaruan aplikasi internal persuratan, dapat kami sampaikan hal-hal sebagai berikut.\n"
            "1. Lakukan koordinasi teknis dengan tim terkait.\n"
            "2. Pastikan seluruh data penting telah dicadangkan.\n"
            f"{closing}"
        ),
        configuration={
            "number": "EVAL-04/IST/YK/05/2026",
            "recipient": "Kepala Pusat Pengembangan dan Layanan Sistem Informasi",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Koordinasi Teknis Pembaruan Aplikasi Internal",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    closing_paragraph = [paragraph for paragraph in document.paragraphs if paragraph.text == closing][0]

    assert _paragraph_space_before_twips(closing_paragraph) >= 280
    assert draft.searchable_text.count(closing) == 1


def test_generate_memo_docx_formats_pic_data_as_key_value_table():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penyampaian Data Pegawai",
        context="Data PIC pegawai.",
        text_generator=lambda prompt: (
            "Menindaklanjuti proses pendataan pegawai, dengan ini disampaikan data sebagai berikut:\n"
            "1. Nama pegawai yang benar: Muhammad Hasbi Ash Shiddiqi\n"
            "2. NIP: 231210013\n"
            "3. Jabatan: Pengadministrasi Perkantoran"
        ),
        configuration={
            "number": "EVAL-38/IST/YK/05/2026",
            "recipient": "Kepala Bagian SDM",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penyampaian Data Pegawai",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    data_table = _find_table_containing(document, "nama")

    assert data_table.cell(0, 0).text == "nama"
    assert data_table.cell(0, 1).text == ":"
    assert data_table.cell(0, 2).text == "Muhammad Hasbi Ash Shiddiqi"
    assert data_table.cell(1, 0).text == "NIP"
    assert data_table.cell(1, 2).text == "231210013"
    assert data_table.cell(2, 0).text == "jabatan"
    assert data_table.cell(2, 2).text == "Pengadministrasi Perkantoran"
    assert _cell_margin_twips(data_table.cell(0, 2), "start") >= 80


def test_generate_memo_docx_keeps_schedule_inside_key_value_table():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penugasan Staf Pendamping Kegiatan",
        context="Penugasan staf pendamping kegiatan.",
        text_generator=lambda prompt: (
            "Dalam rangka pelaksanaan kegiatan integrasi aplikasi layanan internal, perlu dilakukan "
            "penugasan staf pendamping kegiatan untuk memastikan kelancaran proses integrasi.\n"
            "nama: Andi Susanto\n"
            "NIP: 197605172006041001\n"
            "jabatan: Kasubag Teknologi Informasi\n"
            "unit kerja: Bagian SDM\n"
            "5. Jadwal pendampingan: Senin hingga Jumat, pukul 08.00-16.00 WIB."
        ),
        configuration={
            "number": "EVAL-06/IST/YK/05/2026",
            "recipient": "Kepala Bagian SDM",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penugasan Staf Pendamping Kegiatan",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    data_table = _find_table_containing(document, "jadwal pendampingan")

    assert data_table.cell(0, 0).text == "nama"
    assert data_table.cell(4, 0).text == "jadwal pendampingan"
    assert data_table.cell(4, 2).text == "Senin hingga Jumat, pukul 08.00-16.00 WIB"
    assert _cell_width_twips(data_table.cell(4, 0)) >= 2000


def test_generate_memo_docx_formats_period_inline_key_value_and_spaces_following_paragraph():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Pemindahan Lokasi Penugasan Sementara",
        context="Pemindahan lokasi penugasan sementara.",
        text_generator=lambda prompt: (
            "Sehubungan dengan penyesuaian kebutuhan layanan sementara, dapat kami sampaikan bahwa "
            "lokasi asal: Ruang Layanan Administrasi lokasi tujuan: Ruang Koordinasi Terpadu. "
            "Periode: 13 sampai 17 Mei 2026.\n"
            "Selama periode tersebut, setiap pegawai wajib melaporkan progres tugas harian."
        ),
        configuration={
            "number": "EVAL-13/IST/YK/05/2026",
            "recipient": "Kepala Subbagian Kepegawaian",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Pemindahan Lokasi Penugasan Sementara",
            "date": "7 Mei 2026",
            "content": (
                "Lokasi asal: Ruang Layanan Administrasi\n"
                "Lokasi tujuan: Ruang Koordinasi Terpadu\n"
                "Periode: 13 sampai 17 Mei 2026"
            ),
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    data_table = _find_table_containing(document, "periode")

    assert data_table.cell(0, 0).text == "lokasi asal"
    assert data_table.cell(1, 0).text == "lokasi tujuan"
    assert data_table.cell(1, 2).text == "Ruang Koordinasi Terpadu"
    assert data_table.cell(2, 0).text == "periode"
    assert data_table.cell(2, 2).text == "13 sampai 17 Mei 2026"
    assert _spacing_after_table_twips(data_table) >= 160


def test_generate_memo_docx_strips_source_json_urls_and_markdown_artifacts():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Permohonan Persetujuan Jadwal Pemeliharaan Sistem",
        context="Minta persetujuan jadwal pemeliharaan sistem.",
        text_generator=lambda prompt: (
            "Sehubungan dengan kebutuhan pemeliharaan sistem aplikasi internal, kami memohon persetujuan.\n"
            "1. Persetujuan pelaksanaan pemeliharaan sistem pada tanggal yang telah ditentukan.\n"
            "2. Penyampaian pemberitahuan kepada seluruh unit pengguna.\n"
            '[SOURCES:[{"type":"web","title":"Contoh Memo","url":"https://example.com/memo",'
            '"snippet":"contoh memo internal"}]]\n'
            "3. **Koordinasi teknis** dilakukan oleh unit terkait."
        ),
        configuration={
            "number": "EVAL-09/IST/YK/05/2026",
            "recipient": "Kepala Istana Kepresidenan Yogyakarta",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Permohonan Persetujuan Jadwal Pemeliharaan Sistem",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert "[SOURCES" not in all_text
    assert '"type":"web"' not in all_text
    assert "https://example.com" not in all_text
    assert "**" not in all_text
    assert "Koordinasi teknis" in all_text
    assert "[SOURCES" not in draft.searchable_text


def test_generate_memo_docx_formats_inline_pic_data_as_key_value_table():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penyampaian Kontak PIC Layanan",
        context="Data PIC layanan.",
        text_generator=lambda prompt: (
            "Dalam rangka mempercepat koordinasi layanan internal, berikut disampaikan kontak PIC:\n"
            "1. Nama: Eko Prasetyo 2. NIP: 199411172025211057 "
            "3. Pangkat/golongan: V 4. Jabatan: Pengadministrasi Perkantoran "
            "5. Nomor kontak: 0812-0000-2026"
        ),
        configuration={
            "number": "EVAL-19/IST/YK/05/2026",
            "recipient": "Kepala Unit Layanan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penyampaian Kontak PIC Layanan",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    data_table = _find_table_containing(document, "nama")

    assert data_table.cell(0, 0).text == "nama"
    assert data_table.cell(0, 2).text == "Eko Prasetyo"
    assert data_table.cell(1, 0).text == "NIP"
    assert data_table.cell(2, 0).text == "pangkat/gol."
    assert data_table.cell(3, 0).text == "jabatan"
    assert data_table.cell(4, 0).text == "nomor kontak"


def test_generate_memo_docx_formats_comma_separated_inline_pic_data_as_key_value_table():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penyampaian Kontak PIC Layanan",
        context="Data PIC layanan.",
        text_generator=lambda prompt: (
            "Untuk mempercepat koordinasi layanan internal, perlu dilakukan peningkatan koordinasi antar unit. "
            "Sehubungan hal tersebut, dapat kami sampaikan sebagai berikut. "
            "Nama: Eko Prasetyo, NIP: 199411172025211057, pangkat/gol.: V, "
            "jabatan: Pengadministrasi Perkantoran pada Subbagian Tata Usaha, "
            "nomor kontak: 0812-0000-2026."
        ),
        configuration={
            "number": "EVAL-19/IST/YK/05/2026",
            "recipient": "Kepala Unit Layanan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penyampaian Kontak PIC Layanan",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    data_table = _find_table_containing(document, "nama")

    assert data_table.cell(0, 2).text == "Eko Prasetyo"
    assert data_table.cell(1, 0).text == "NIP"
    assert data_table.cell(1, 2).text == "199411172025211057"
    assert data_table.cell(2, 0).text == "pangkat/gol."
    assert data_table.cell(3, 0).text == "jabatan"
    assert data_table.cell(4, 0).text == "nomor kontak"
    assert data_table.cell(4, 2).text == "0812-0000-2026"


def test_generate_memo_docx_prefers_complete_configured_pic_rows_over_model_narrative():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penyampaian Data Pegawai Pendamping Kegiatan",
        context="Data pegawai pendamping kegiatan.",
        text_generator=lambda prompt: (
            "Sehubungan dengan kebutuhan pendampingan kegiatan integrasi aplikasi, dapat kami sampaikan sebagai berikut.\n"
            "nama: Muhammad Hasbi Ash Shiddiqi\n"
            "NIP: 231210013\n"
            "jabatan: Analis Sistem Informasi\n"
            "unit kerja: Subbagian Tata Usaha. Pegawai tersebut ditugaskan untuk pendampingan kegiatan integrasi aplikasi."
        ),
        configuration={
            "number": "EVAL-11/IST/YK/05/2026",
            "recipient": "Kepala Bagian SDM",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penyampaian Data Pegawai Pendamping Kegiatan",
            "date": "7 Mei 2026",
            "content": (
                "Nama: Muhammad Hasbi Ash Shiddiqi\n"
                "NIP: 231210013\n"
                "Jabatan: Analis Sistem Informasi\n"
                "Unit kerja: Subbagian Tata Usaha\n"
                "Keperluan: Pendampingan kegiatan integrasi aplikasi."
            ),
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    data_table = _find_table_containing(document, "keperluan")

    assert data_table.cell(0, 0).text == "nama"
    assert data_table.cell(3, 0).text == "unit kerja"
    assert data_table.cell(3, 2).text == "Subbagian Tata Usaha"
    assert data_table.cell(4, 0).text == "keperluan"
    assert data_table.cell(4, 2).text == "Pendampingan kegiatan integrasi aplikasi"
    assert "Pegawai tersebut ditugaskan" not in data_table.cell(3, 2).text


def test_generate_memo_docx_keeps_blank_signatory_blank():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Konfirmasi Kehadiran Rapat Singkat",
        context="Konfirmasi kehadiran rapat.",
        text_generator=lambda prompt: "Menindaklanjuti rencana rapat singkat, dimohon mengonfirmasi kehadiran.",
        configuration={
            "number": "EVAL-21/IST/YK/05/2026",
            "recipient": "Kepala Subbagian Persuratan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Konfirmasi Kehadiran Rapat Singkat",
            "date": "7 Mei 2026",
            "signatory": "",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    assert not _has_table_containing(document, "QR\nTTE")
    assert "Deni Mulyana" not in _all_document_text(document)
    assert "Deni Mulyana" not in draft.searchable_text
    assert "Dokumen ini telah ditandatangani secara elektronik" not in document.sections[0].footer.paragraphs[0].text
    assert "Dokumen ini telah ditandatangani secara elektronik" not in draft.searchable_text


def test_generate_memo_docx_uses_compact_layout_without_manual_page_break_for_long_folio_with_carbon_copy():
    long_body = "\n".join(
        [
            f"{index}. "
            + "Koordinasi dan penataan layanan administrasi internal perlu dilakukan secara bertahap. " * 7
            for index in range(1, 7)
        ]
    )

    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Memo Panjang Format Folio",
        context="Buat memo panjang format folio.",
        text_generator=lambda prompt: long_body,
        configuration={
            "number": "EVAL-25/IST/YK/05/2026",
            "recipient": "Kepala Pusat Pengembangan Layanan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Memo Panjang Format Folio",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "carbon_copy": "Kepala Biro Perencanaan\nKepala Bagian Administrasi\nKoordinator Layanan",
            "page_size": "folio",
        },
    )

    document = Document(BytesIO(draft.content))
    page_breaks = document._element.xpath(".//w:br[@w:type='page']")
    qr_table = _find_table_containing(document, "QR\nTTE")

    assert not page_breaks
    assert qr_table.cell(1, 0).text == "Deni Mulyana"
    assert "Tembusan:" in _all_document_text(document)
    assert _separator_after_twips(document) >= 780
    assert _signature_spacer_before_twips(qr_table) >= 1000


def test_generate_memo_docx_keeps_compact_folio_signature_lower_for_medium_numbered_body():
    body = "\n".join(
        [
            "Menindaklanjuti kebutuhan konsolidasi rencana kerja lintas unit, disampaikan langkah sebagai berikut:",
            "1. Menyusun rencana kerja secara ringkas.",
            "2. Menunjuk PIC serta menetapkan jadwal pelaksanaan.",
            "3. Mengidentifikasi kendala dan mengajukan usulan perbaikan.",
            "4. Menyampaikan laporan lengkap kepada pimpinan sesuai jadwal.",
            "5. Melakukan evaluasi berkala atas tindak lanjut setiap unit.",
        ]
    )

    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Revisi Memo Menjadi Lebih Singkat",
        context="Revisi ringkas dengan lima poin.",
        text_generator=lambda prompt: body,
        configuration={
            "number": "EVAL-36/IST/YK/05/2026",
            "recipient": "Kepala Biro Perencanaan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Revisi Memo Menjadi Lebih Singkat",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "folio",
        },
    )

    document = Document(BytesIO(draft.content))
    qr_table = _find_table_containing(document, "QR\nTTE")

    assert _signature_spacer_before_twips(qr_table) >= 2200
    assert 0.67 <= _signature_qr_center_ratio(document, qr_table) <= 0.72


def test_generate_memo_docx_cleans_broken_official_language_fragments():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Permohonan Persetujuan Jadwal",
        context="Persetujuan jadwal kegiatan.",
        text_generator=lambda prompt: (
            "Sehubungan dengan rencana kegiatan, kami ingin menyampaikan berikut. "
            "Mohon untuk mem dan menindaklanjuti hal tersebut. Terimakasih."
        ),
        configuration={
            "number": "EVAL-09/IST/YK/05/2026",
            "recipient": "Kepala Istana Kepresidenan Yogyakarta",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Permohonan Persetujuan Jadwal",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    text = draft.searchable_text

    assert "mohon untuk mem dan" not in text.lower()
    assert "kami ingin menyampaikan berikut" not in text.lower()
    assert "dapat kami sampaikan sebagai berikut" in text.lower()
    assert "Terimakasih" not in text


def test_generate_memo_docx_anchors_short_signature_lower_without_changing_horizontal_qr():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penyampaian Kontak PIC Layanan",
        context="Data PIC layanan.",
        text_generator=lambda prompt: (
            "Dalam rangka mempercepat koordinasi layanan internal, dengan ini disampaikan "
            "kontak PIC layanan sebagai berikut:\n"
            "nama: Eko Prasetyo\n"
            "NIP: 199411172025211057\n"
            "jabatan: Pengadministrasi Perkantoran"
        ),
        configuration={
            "number": "EVAL-19/IST/YK/05/2026",
            "recipient": "Kepala Unit Layanan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penyampaian Kontak PIC Layanan",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    qr_table = _find_table_containing(document, "QR\nTTE")

    assert _signature_spacer_before_twips(qr_table) >= 2000
    assert 0.67 <= _signature_qr_center_ratio(document, qr_table) <= 0.72


def test_generate_memo_docx_enforces_short_revision_constraints():
    generated = "\n".join(
        [
            "Menindaklanjuti kebutuhan konsolidasi rencana kerja lintas unit, disampaikan beberapa langkah untuk meningkatkan efisiensi dan efektifitas pelaksanaan rencana kerja.",
            "Berikut adalah beberapa keputusan yang perlu diambil:",
            "1. Menyusun rencana kerja secara ringkas, disertai data pendukung yang relevan, sehingga memudahkan pelaksanaan dan pemantauan.",
            "2. Menunjuk PIC serta menetapkan jadwal pelaksanaan, dengan memastikan bahwa PIC memiliki kemampuan dan sumber daya yang memadai.",
            "3. Mengidentifikasi kendala dan mengajukan usulan perbaikan, sehingga dapat dilakukan antisipasi dan penyelesaian masalah secara efektif.",
            "4. Menyampaikan laporan lengkap kepada pimpinan sesuai jadwal, sehingga memungkinkan pemantauan dan evaluasi berkala.",
            "5. Melakukan evaluasi berkala atas tindak lanjut setiap unit, sehingga dapat dilakukan penyesuaian dan perbaikan.",
            "Dalam pelaksanaan rencana kerja, perlu diingat bahwa efisiensi dan efektifitas adalah kunci untuk mencapai tujuan.",
        ]
    )

    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Revisi Memo Menjadi Lebih Singkat",
        context="Revisi memo menjadi lebih singkat.",
        text_generator=lambda prompt: generated,
        configuration={
            "number": "EVAL-36/IST/YK/05/2026",
            "recipient": "Kepala Biro Perencanaan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Revisi Memo Menjadi Lebih Singkat",
            "date": "7 Mei 2026",
            "basis": "Menindaklanjuti kebutuhan konsolidasi rencana kerja lintas unit.",
            "content": (
                "1. Menyusun rencana kerja secara ringkas, disertai data pendukung yang relevan.\n"
                "2. Menunjuk PIC serta menetapkan jadwal pelaksanaan.\n"
                "3. Mengidentifikasi kendala dan mengajukan usulan perbaikan.\n"
                "4. Menyampaikan laporan lengkap kepada pimpinan sesuai jadwal.\n"
                "5. Melakukan evaluasi berkala atas tindak lanjut setiap unit."
            ),
            "signatory": "Deni Mulyana",
            "carbon_copy": "Para Kepala Bagian",
            "page_size": "auto",
            "page_size_mode": "auto",
            "revision_instruction": "Buat isi memo menjadi lebih singkat maksimal dua paragraf, metadata jangan berubah.",
            "additional_instruction": "Buat isi memo menjadi lebih singkat maksimal dua paragraf.",
        },
    )

    assert "efisiensi dan efektifitas adalah kunci" not in draft.searchable_text
    assert "penyusunan rencana kerja secara ringkas" in draft.searchable_text
    assert "evaluasi berkala atas tindak lanjut setiap unit" in draft.searchable_text
    assert "Menyusun rencana kerja secara ringkas" not in draft.searchable_text
    assert "Melakukan evaluasi berkala" not in draft.searchable_text
    assert "hal-hal yang perlu ditindaklanjuti meliputi" in draft.searchable_text
    assert "\n1. Menyusun rencana kerja" not in draft.searchable_text
    assert "\n5. Melakukan evaluasi berkala" not in draft.searchable_text
    assert len(draft.searchable_text.split()) < len(generated.split())


def test_generate_memo_docx_formats_activity_details_as_official_key_value_block():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Undangan Rapat Persiapan Kegiatan Kunjungan",
        context="Undangan rapat persiapan kegiatan kunjungan.",
        text_generator=lambda prompt: (
            "Sehubungan dengan persiapan kegiatan kunjungan tamu pada tanggal 16 Mei 2026, "
            "dapat kami sampaikan sebagai berikut.\n"
            "1. Mohon mengundang unit terkait untuk menghadiri rapat persiapan di ruang rapat utama.\n"
            "2. Rapat akan dilaksanakan pada:\n"
            "3. Agenda rapat meliputi pembahasan alur tamu, kesiapan ruang, serta pendampingan acara."
        ),
        configuration={
            "number": "EVAL-05/IST/YK/05/2026",
            "recipient": "Kepala Bagian Protokol",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Undangan Rapat Persiapan Kegiatan Kunjungan",
            "date": "7 Mei 2026",
            "content": (
                "Undang unit terkait untuk rapat persiapan di ruang rapat utama pada 10 Mei 2026 "
                "pukul 09.00 WIB dengan agenda alur tamu, kesiapan ruang, dan pendampingan acara."
            ),
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)
    detail_table = _find_table_containing(document, "hari/tanggal")

    assert detail_table.cell(0, 2).text == "10 Mei 2026"
    assert detail_table.cell(1, 0).text == "pukul"
    assert detail_table.cell(1, 2).text == "09.00 WIB"
    assert detail_table.cell(2, 0).text == "tempat"
    assert detail_table.cell(2, 2).text == "ruang rapat utama"
    assert detail_table.cell(3, 0).text == "agenda"
    assert "alur tamu" in detail_table.cell(3, 2).text
    assert _spacing_after_table_twips(detail_table) >= 160
    assert "Rapat akan dilaksanakan pada" not in all_text
    assert "Agenda rapat meliputi" not in all_text
    assert "1.\tMohon mengundang unit terkait" in all_text


def test_generate_memo_docx_removes_redundant_sub_bullets_after_activity_key_values():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Undangan Rapat Persiapan Kegiatan Kunjungan",
        context="Undangan rapat persiapan kegiatan kunjungan.",
        text_generator=lambda prompt: (
            "Sehubungan dengan persiapan kegiatan kunjungan tamu, dapat kami sampaikan sebagai berikut.\n"
            "1. Mohon mengundang unit terkait untuk menghadiri rapat persiapan.\n"
            "- Pembahasan alur tamu.\n"
            "- Kesiapan ruang.\n"
            "- Pendampingan acara."
        ),
        configuration={
            "number": "EVAL-05/IST/YK/05/2026",
            "recipient": "Kepala Bagian Protokol",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Undangan Rapat Persiapan Kegiatan Kunjungan",
            "date": "7 Mei 2026",
            "content": (
                "Undang unit terkait untuk rapat persiapan di ruang rapat utama pada 10 Mei 2026 "
                "pukul 09.00 WIB dengan agenda alur tamu, kesiapan ruang, dan pendampingan acara."
            ),
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert _has_table_containing(document, "hari/tanggal")
    assert "1.\tMohon mengundang unit terkait" in all_text
    assert "Pembahasan alur tamu" not in all_text
    assert "Kesiapan ruang" not in all_text
    assert "Pendampingan acara" not in all_text


def test_generate_memo_docx_removes_repeated_existing_schedule_table_details():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Revisi Tanggal Kegiatan",
        context="Revisi tanggal kegiatan.",
        text_generator=lambda prompt: (
            "Menindaklanjuti persiapan rapat koordinasi kegiatan, sehubungan hal tersebut dapat kami sampaikan sebagai berikut.\n"
            "1. Tanggal pelaksanaan rapat: 13 Mei 2026\n"
            "hari/tanggal: 13 Mei 2026\n"
            "pukul: 09.00 WIB\n"
            "tempat: ruang rapat utama\n"
            "agenda: rapat meliputi pembagian peran serta pengecekan kesiapan setiap bagian terkait\n"
            "Agenda rapat meliputi pembagian peran serta pengecekan kesiapan setiap bagian terkait."
        ),
        configuration={
            "number": "EVAL-39/IST/YK/05/2026",
            "recipient": "Kepala Bagian Protokol",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Revisi Tanggal Kegiatan",
            "date": "7 Mei 2026",
            "content": (
                "1. Tanggal pelaksanaan rapat direvisi menjadi 13 Mei 2026 pukul 09.00 WIB, "
                "bertempat di ruang rapat utama.\n"
                "2. Agenda rapat meliputi pembagian peran serta pengecekan kesiapan setiap bagian terkait."
            ),
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)
    detail_table = _find_table_containing(document, "hari/tanggal")

    assert detail_table.cell(0, 2).text == "13 Mei 2026"
    assert detail_table.cell(1, 2).text == "09.00 WIB"
    assert detail_table.cell(2, 2).text == "ruang rapat utama"
    assert "pembagian peran" in detail_table.cell(3, 2).text
    assert "Tanggal pelaksanaan rapat:" not in all_text
    assert "Agenda rapat meliputi" not in all_text


def test_generate_memo_docx_keeps_numbering_continuous_after_redundant_key_value_item_removed():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Rapat Evaluasi Durasi Layanan",
        context="Rapat evaluasi durasi layanan.",
        text_generator=lambda prompt: (
            "Sehubungan dengan rencana evaluasi layanan, dapat kami sampaikan sebagai berikut.\n"
            "1. Estimasi durasi rapat: 2 jam.\n"
            "2. Peserta menyiapkan data layanan masing-masing.\n"
            "3. Hasil rapat disampaikan kepada pimpinan."
        ),
        configuration={
            "number": "EVAL-09/IST/YK/05/2026",
            "recipient": "Para Kepala Subbagian",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Rapat Evaluasi Durasi Layanan",
            "date": "7 Mei 2026",
            "content": (
                "Rapat dilaksanakan 12 Mei 2026 pukul 10.00 WIB, aula utama. "
                "Agenda evaluasi layanan. Estimasi durasi: 2 jam. Peserta menyiapkan data layanan."
            ),
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)
    numbered_paragraphs = [paragraph.text for paragraph in document.paragraphs if "\t" in paragraph.text]

    assert _has_table_containing(document, "estimasi durasi")
    assert "Estimasi durasi rapat" not in all_text
    assert numbered_paragraphs == [
        "1.\tPeserta menyiapkan data layanan masing-masing.",
        "2.\tHasil rapat disampaikan kepada pimpinan.",
    ]


def test_generate_memo_docx_preserves_decimal_time_in_activity_key_value_block():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Konfirmasi Jadwal Kegiatan Koordinasi",
        context="Konfirmasi jadwal kegiatan koordinasi.",
        text_generator=lambda prompt: (
            "Menindaklanjuti jadwal kegiatan koordinasi lintas unit, dengan ini kami sampaikan "
            "informasi sebagai berikut:\n"
            "1. Hari/Tanggal: Rabu, 13 Mei 2026\n"
            "2. Waktu: Pukul 09.00 WIB\n"
            "3. Tempat: Ruang Rapat Utama\n"
            "4. Agenda: Pembagian peran dan pengecekan kesiapan setiap bagian"
        ),
        configuration={
            "number": "EVAL-12/IST/YK/05/2026",
            "recipient": "Kepala Bagian Protokol",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Konfirmasi Jadwal Kegiatan Koordinasi",
            "date": "10 Mei 2026",
            "content": (
                "Tanggal kegiatan 13 Mei 2026, pukul 09.00 WIB, ruang rapat utama, "
                "agenda pembagian peran dan pengecekan kesiapan setiap bagian."
            ),
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    body_paragraphs = [paragraph.text for paragraph in document.paragraphs]
    detail_table = _find_table_containing(document, "pukul")

    assert detail_table.cell(0, 2).text == "13 Mei 2026"
    assert detail_table.cell(1, 2).text == "09.00 WIB"
    assert detail_table.cell(2, 2).text == "ruang rapat utama"
    assert "1.\tHari/Tanggal: Rabu, 13 Mei 2026" not in body_paragraphs
    assert "2.\tWaktu: Pukul 09.00 WIB" not in body_paragraphs
    assert "3.\tTempat: Ruang Rapat Utama" not in body_paragraphs
    assert "4.\tAgenda: Pembagian peran dan pengecekan kesiapan setiap bagian" not in body_paragraphs
    assert "pukul : 00 WIB" not in _all_document_text(document)
    assert "pukul: 00 WIB" not in draft.searchable_text


def test_generate_memo_docx_justifies_short_body_paragraphs():
    body = "Sehubungan dengan kebutuhan koordinasi, kami sampaikan informasi kegiatan."
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Koordinasi Singkat",
        context="Koordinasi singkat.",
        text_generator=lambda prompt: body,
        configuration={
            "number": "EVAL-ALIGN/IST/YK/05/2026",
            "recipient": "Kepala Bagian Protokol",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Koordinasi Singkat",
            "date": "10 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    paragraph = _find_paragraph(document, body)

    assert paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY


def test_generate_memo_docx_preserves_time_range_in_activity_key_value_block():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Konfirmasi Kehadiran Rapat Singkat",
        context="Konfirmasi kehadiran rapat singkat.",
        text_generator=lambda prompt: (
            "Sehubungan dengan rencana rapat singkat, dimohon agar peserta mengonfirmasi kehadiran."
        ),
        configuration={
            "number": "EVAL-21/IST/YK/05/2026",
            "recipient": "Kepala Subbagian Persuratan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Konfirmasi Kehadiran Rapat Singkat",
            "date": "7 Mei 2026",
            "content": "Rapat dilaksanakan 12 Mei 2026 pukul 09.00 s.d. 11.30 WIB.",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    detail_table = _find_table_containing(document, "pukul")

    assert detail_table.cell(0, 2).text == "12 Mei 2026"
    assert detail_table.cell(1, 2).text == "09.00 s.d. 11.30 WIB"


def test_generate_memo_docx_trims_activity_period_and_removes_redundant_detail_list():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Pemindahan Lokasi Penugasan Sementara",
        context="Pemindahan lokasi penugasan sementara.",
        text_generator=lambda prompt: (
            "Dalam rangka penyesuaian kebutuhan layanan sementara, sehubungan hal tersebut, dapat kami sampaikan sebagai berikut.\n"
            "lokasi asal: Ruang Layanan Administrasi\n"
            "lokasi tujuan: Ruang Koordinasi Terpadu\n"
            "periode: 13 sampai 17 Mei 2026\n"
            "1. Penugasan sementara pegawai dari lokasi asal di Ruang Layanan Administrasi ke lokasi tujuan di Ruang Koordinasi Terpadu.\n"
            "2. Pegawai wajib melaporkan progres harian kepada atasan langsung."
        ),
        configuration={
            "number": "EVAL-13/IST/YK/05/2026",
            "recipient": "Kepala Subbagian Kepegawaian",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Pemindahan Lokasi Penugasan Sementara",
            "date": "7 Mei 2026",
            "content": (
                "Lokasi asal: Ruang Layanan Administrasi. Lokasi tujuan: Ruang Koordinasi Terpadu. "
                "Periode: 13 sampai 17 Mei 2026. Pegawai tetap melaporkan progres harian kepada atasan langsung."
            ),
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)
    detail_table = _find_table_containing(document, "periode")

    assert detail_table.cell(0, 2).text == "Ruang Layanan Administrasi"
    assert detail_table.cell(1, 2).text == "Ruang Koordinasi Terpadu"
    assert detail_table.cell(2, 2).text == "13 sampai 17 Mei 2026"
    assert _spacing_after_table_twips(detail_table) >= 240
    assert "Penugasan sementara pegawai dari lokasi asal" not in all_text
    assert "1.\tPegawai wajib melaporkan progres harian" in all_text


def test_generate_memo_docx_strips_bare_manual_closing_instruction_from_body():
    closing = "Demikian untuk menjadi bahan gerak cepat bersama."
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penutup Manual Tidak Lazim",
        context="Penutup manual tidak lazim.",
        text_generator=lambda prompt: (
            "Sehubungan dengan kebutuhan tindak lanjut cepat, dapat kami sampaikan sebagai berikut.\n"
            "1. Unit terkait agar segera menyiapkan data yang dibutuhkan.\n"
            "Penutup manual apa adanya.\n"
            f"{closing}"
        ),
        configuration={
            "number": "EVAL-30/IST/YK/05/2026",
            "recipient": "Kepala Bagian Administrasi",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penutup Manual Tidak Lazim",
            "date": "7 Mei 2026",
            "content": "Sampaikan arahan agar unit terkait menyiapkan data, menetapkan PIC, dan melaporkan progres.",
            "closing": closing,
            "signatory": "Deni Mulyana",
            "page_size": "letter",
            "additional_instruction": "Pertahankan penutup manual apa adanya.",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert "Penutup manual apa adanya" not in all_text
    assert all_text.count(closing) == 1


def test_generate_memo_docx_strips_manual_closing_instruction_variant_from_body():
    closing = "Demikian untuk menjadi bahan gerak cepat bersama."
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penutup Manual Tidak Lazim",
        context="Penutup manual tidak lazim.",
        text_generator=lambda prompt: (
            "Menindaklanjuti kebutuhan tindak lanjut cepat, dapat kami sampaikan sebagai berikut.\n"
            "1. Unit terkait agar segera menyiapkan data yang diperlukan.\n"
            "2. Harap menetapkan PIC untuk pelaksanaan kegiatan dimaksud.\n"
            "3. Progres pelaksanaan agar dilaporkan secara berkala sesuai ketentuan.\n"
            "Penutup manual tidak lazim apa adanya.\n"
            f"{closing}"
        ),
        configuration={
            "number": "EVAL-30/IST/YK/05/2026",
            "recipient": "Kepala Bagian Administrasi",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penutup Manual Tidak Lazim",
            "date": "7 Mei 2026",
            "content": "Sampaikan arahan agar unit terkait menyiapkan data, menetapkan PIC, dan melaporkan progres.",
            "closing": closing,
            "signatory": "Deni Mulyana",
            "page_size": "letter",
            "additional_instruction": "Pertahankan penutup manual apa adanya.",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert "Penutup manual tidak lazim apa adanya" not in all_text
    assert all_text.count(closing) == 1


def test_generate_memo_docx_replaces_empty_person_placeholders_with_generic_text():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penugasan Staf Pendamping Kegiatan",
        context="Penugasan staf pendamping kegiatan.",
        text_generator=lambda prompt: (
            "Dalam rangka pelaksanaan kegiatan integrasi aplikasi layanan internal, sehubungan hal tersebut, dapat kami sampaikan sebagai berikut.\n"
            "1. Nama:\n"
            "2. NIP:\n"
            "3. Jabatan:\n"
            "4. Unit Kerja:\n"
            "5. Jadwal Pendampingan:"
        ),
        configuration={
            "number": "EVAL-06/IST/YK/05/2026",
            "recipient": "Kepala Bagian SDM",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penugasan Staf Pendamping Kegiatan",
            "date": "7 Mei 2026",
            "content": "Tugaskan satu staf pendamping kegiatan, cantumkan nama, NIP, jabatan, unit kerja, dan jadwal pendampingan.",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert "Nama:" not in all_text
    assert "NIP:" not in all_text
    assert "Jabatan:" not in all_text
    assert "Staf pendamping dan kelengkapan data pegawai ditetapkan oleh unit terkait." in all_text


def test_generate_memo_docx_removes_invented_pic_names_when_not_configured():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penyampaian Hasil Rapat Koordinasi Layanan",
        context="Penyampaian hasil rapat koordinasi layanan.",
        text_generator=lambda prompt: (
            "Berdasarkan hasil rapat koordinasi layanan tanggal 6 Mei 2026, dapat kami sampaikan sebagai berikut:\n"
            "1. Penunjukan Person In Charge (PIC) untuk tiap layanan sebagai berikut:\n"
            "- Subbagian Administrasi: Bapak Andi Prasetyo\n"
            "- Subbagian Keuangan: Ibu Maria Dewi\n"
            "2. Jadwal pelaporan mingguan ditetapkan setiap hari Senin."
        ),
        configuration={
            "number": "EVAL-02/IST/YK/05/2026",
            "recipient": "Para Kepala Subbagian",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penyampaian Hasil Rapat Koordinasi Layanan",
            "date": "7 Mei 2026",
            "content": "Sampaikan hasil rapat berupa pembagian PIC, jadwal pelaporan mingguan, dan penyampaian kendala melalui kanal prioritas.",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert "Bapak Andi Prasetyo" not in all_text
    assert "Ibu Maria Dewi" not in all_text
    assert "PIC) pada tiap layanan agar ditetapkan oleh masing-masing unit" in all_text
    assert "-\tSubbagian" not in all_text


def test_generate_memo_docx_replaces_unconfigured_incident_time_and_impacted_users():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penyampaian Kendala Akses Sistem Persuratan",
        context="Penyampaian kendala akses sistem persuratan.",
        text_generator=lambda prompt: (
            "Sehubungan dengan kendala akses sistem persuratan pada 6 Mei 2026, dapat kami sampaikan sebagai berikut.\n"
            "1. Kendala akses sistem persuratan terjadi pada tanggal 6 Mei 2026, pukul 09.30 WIB.\n"
            "2. Pengguna yang terdampak adalah seluruh staf persuratan.\n"
            "3. Tindak lanjut dilakukan melalui koordinasi dengan unit teknis."
        ),
        configuration={
            "number": "EVAL-08/IST/YK/05/2026",
            "recipient": "Kepala Subbagian Persuratan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penyampaian Kendala Akses Sistem Persuratan",
            "date": "7 Mei 2026",
            "content": (
                "Sampaikan kendala akses sistem persuratan pada 6 Mei 2026, "
                "waktu kejadian, pengguna terdampak, dan tindak lanjut."
            ),
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert "09.30 WIB" not in all_text
    assert "seluruh staf persuratan" not in all_text
    assert "Waktu kejadian ditetapkan berdasarkan laporan unit terkait." in all_text
    assert "Pengguna terdampak diidentifikasi oleh unit terkait." in all_text


def test_generate_memo_docx_preserves_configured_numbering_over_ordinal_rewrite():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Memo Panjang Format Folio Eksplisit",
        context="Memo panjang format folio eksplisit.",
        text_generator=lambda prompt: (
            "Sehubungan dengan kebutuhan peningkatan kualitas layanan administrasi internal, dapat kami sampaikan hal-hal sebagai berikut.\n"
            "Pertama, perlu dilakukan penyederhanaan alur administrasi.\n"
            "Kedua, percepatan respons layanan administrasi menjadi prioritas."
        ),
        configuration={
            "number": "EVAL-25/IST/YK/05/2026",
            "recipient": "Kepala Pusat Pengembangan Layanan",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Memo Panjang Format Folio Eksplisit",
            "date": "7 Mei 2026",
            "basis": "Menindaklanjuti kebutuhan peningkatan kualitas layanan administrasi internal.",
            "content": (
                "1. Penyederhanaan Alur Administrasi: Layanan administrasi internal perlu diselaraskan.\n"
                "2. Percepatan Respons Layanan: Layanan administrasi internal perlu dipercepat."
            ),
            "signatory": "Deni Mulyana",
            "page_size": "folio",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert "Pertama" not in all_text
    assert "Kedua" not in all_text
    assert "1.\tPenyederhanaan Alur Administrasi" in all_text
    assert "2.\tPercepatan Respons Layanan" in all_text


def test_generate_memo_docx_strips_manual_closing_instruction_notes():
    closing = "Demikian untuk menjadi bahan gerak cepat bersama."
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Penutup Manual Tidak Lazim",
        context="Penutup manual tidak lazim.",
        text_generator=lambda prompt: (
            "Sehubungan dengan kebutuhan tindak lanjut yang harus dilakukan secara cepat, dapat kami sampaikan arahan sebagai berikut.\n"
            "1. Unit terkait diminta untuk segera menyiapkan data pendukung.\n"
            "Catatan: Penutup manual tetap dipertahankan sesuai dengan format yang berlaku.\n"
            f"{closing}"
        ),
        configuration={
            "number": "EVAL-30/IST/YK/05/2026",
            "recipient": "Kepala Bagian Administrasi",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Penutup Manual Tidak Lazim",
            "date": "7 Mei 2026",
            "content": "Sampaikan arahan agar unit terkait menyiapkan data, menetapkan PIC, dan melaporkan progres.",
            "closing": closing,
            "signatory": "Deni Mulyana",
            "page_size": "letter",
            "additional_instruction": "Pertahankan penutup manual apa adanya.",
        },
    )

    document = Document(BytesIO(draft.content))
    all_text = _all_document_text(document)

    assert "Catatan:" not in all_text
    assert all_text.count(closing) == 1


def test_generate_memo_docx_moves_dimohon_sentence_to_closing_block():
    closing = "Dimohon kehadiran tepat waktu dan kesiapan peserta agar rapat dapat berjalan efektif."
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Daftar Peserta Rapat Evaluasi Layanan",
        context="Daftar peserta rapat evaluasi layanan.",
        text_generator=lambda prompt: (
            "Sehubungan dengan akan dilaksanakannya rapat evaluasi layanan internal, dapat kami sampaikan sebagai berikut:\n"
            "1. Peserta rapat wajib dihadiri oleh Kepala Bagian Administrasi.\n"
            f"2. Rapat akan diselenggarakan pada tanggal 12 Mei 2026 pukul 10.00 WIB. {closing}"
        ),
        configuration={
            "number": "EVAL-14/IST/YK/05/2026",
            "recipient": "Para Kepala Subbagian",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Daftar Peserta Rapat Evaluasi Layanan",
            "date": "7 Mei 2026",
            "content": "Peserta wajib: Kepala Bagian Administrasi. Rapat dilaksanakan 12 Mei 2026 pukul 10.00 WIB.",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    closing_paragraphs = [paragraph for paragraph in document.paragraphs if paragraph.text == closing]

    assert len(closing_paragraphs) == 1
    assert _paragraph_space_before_twips(closing_paragraphs[0]) >= 280
    assert draft.searchable_text.count(closing) == 1


def test_generate_memo_docx_keeps_carbon_copy_close_to_signature_on_short_document():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Revisi Tanggal Kegiatan",
        context="Revisi tanggal kegiatan.",
        text_generator=lambda prompt: "Menindaklanjuti persiapan rapat koordinasi kegiatan, dapat kami sampaikan agar unit terkait menyesuaikan jadwal pelaksanaan.",
        configuration={
            "number": "EVAL-39/IST/YK/05/2026",
            "recipient": "Kepala Bagian Protokol",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Revisi Tanggal Kegiatan",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "carbon_copy": "Kepala Bagian Keamanan",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    tembusan = _find_paragraph(document, "Tembusan:")

    assert 360 <= _paragraph_space_before_twips(tembusan) <= 560


def test_generate_memo_docx_italicizes_foreign_terms_without_italicizing_system_names():
    draft = generate_memo_docx(
        memo_type="memo_internal",
        title="Undangan Partisipasi Kegiatan Setneg Book Club",
        context="Undangan kegiatan literasi.",
        text_generator=lambda prompt: (
            "Kegiatan membahas e-book dan proses sign up untuk meningkatkan critical thinking "
            "melalui SATUSEHAT Mobile."
        ),
        configuration={
            "number": "EVAL-50/IST/YK/05/2026",
            "recipient": "Pejabat terlampir",
            "sender": "Kepala Istana Kepresidenan Yogyakarta",
            "subject": "Undangan Partisipasi Kegiatan Setneg Book Club",
            "date": "7 Mei 2026",
            "signatory": "Deni Mulyana",
            "page_size": "letter",
        },
    )

    document = Document(BytesIO(draft.content))
    runs = [run for paragraph in document.paragraphs for run in paragraph.runs]

    assert any(run.text == "e-book" and run.font.italic for run in runs)
    assert any(run.text == "sign up" and run.font.italic for run in runs)
    assert any(run.text == "critical thinking" and run.font.italic for run in runs)
    assert any("SATUSEHAT Mobile" in run.text and not run.font.italic for run in runs)


def test_generate_memo_docx_rejects_unknown_type():
    try:
        generate_memo_docx(
            memo_type="surat_bebas",
            title="Judul",
            context="Konteks",
            text_generator=lambda prompt: "Isi",
        )
    except ValueError as exc:
        assert "Jenis memo" in str(exc)
    else:
        raise AssertionError("Expected ValueError")


def test_generate_memo_endpoint_requires_token():
    client = TestClient(app)

    response = client.post("/api/memos/generate-body", json={
        "memo_type": "memo_internal",
        "title": "Judul",
        "context": "Konteks",
    })

    assert response.status_code == 401


def test_generate_memo_endpoint_handles_unicode_searchable_text(monkeypatch):
    def fake_generate_memo_docx(memo_type, title, context, configuration=None):
        assert configuration == {"number": "M-01/I-Yog/IT.02/05/2026"}

        return MemoDraft(
            filename="memo-unicode.docx",
            content=b"docx-bytes",
            searchable_text="Paragraf awal\n- Poin penting • arahan — selesai",
            page_size="letter",
        )

    monkeypatch.setattr("app.routers.memos.generate_memo_docx", fake_generate_memo_docx)
    client = TestClient(app)

    response = client.post(
        "/api/memos/generate-body",
        headers={"Authorization": "Bearer test_internal_api_secret"},
        json={
            "memo_type": "memo_internal",
            "title": "Judul Unicode",
            "context": "Konteks",
            "configuration": {"number": "M-01/I-Yog/IT.02/05/2026"},
        },
    )

    assert response.status_code == 200
    assert response.content == b"docx-bytes"
    assert "X-Memo-Searchable-Text" not in response.headers
    encoded = response.headers["X-Memo-Searchable-Text-B64"]
    assert isinstance(encoded.encode("latin-1"), bytes)
    assert response.headers["X-Memo-Page-Size"] == "letter"
