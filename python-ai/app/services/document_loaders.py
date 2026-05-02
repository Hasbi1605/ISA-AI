import csv
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from langchain_core.documents import Document
from openpyxl import load_workbook


def _clean_text(value) -> str:
    return "" if value is None else str(value).strip()


def _table_to_text(rows: list[list[object]]) -> str:
    lines: list[str] = []
    for row in rows:
        cleaned = [_clean_text(cell) for cell in row]
        if any(cleaned):
            lines.append(" | ".join(cleaned))
    return "\n".join(lines)


def _load_pdf_documents(file_path: str, filename: str) -> list[Document]:
    docs: list[Document] = []

    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            parts: list[str] = []
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(text)

            for table in page.extract_tables() or []:
                table_text = _table_to_text(table)
                if table_text:
                    parts.append(table_text)

            content = "\n\n".join(parts).strip()
            if content:
                docs.append(
                    Document(
                        page_content=content,
                        metadata={"source": filename, "page": page_number},
                    )
                )

    return docs


def _load_docx_documents(file_path: str, filename: str) -> list[Document]:
    document = DocxDocument(file_path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        table_text = _table_to_text(
            [[cell.text for cell in row.cells] for row in table.rows]
        )
        if table_text:
            parts.append(table_text)

    content = "\n\n".join(parts).strip()
    return [Document(page_content=content, metadata={"source": filename})] if content else []


def _load_xlsx_documents(file_path: str, filename: str) -> list[Document]:
    workbook = load_workbook(file_path, data_only=True, read_only=True)
    docs: list[Document] = []

    try:
        for sheet in workbook.worksheets:
            rows = [
                [cell for cell in row]
                for row in sheet.iter_rows(values_only=True)
            ]
            table_text = _table_to_text(rows)
            if table_text:
                docs.append(
                    Document(
                        page_content=f"{sheet.title}\n\n{table_text}",
                        metadata={"source": filename, "sheet": sheet.title},
                    )
                )
    finally:
        workbook.close()

    return docs


def _load_csv_documents(file_path: str, filename: str) -> list[Document]:
    rows: list[list[str]] = []

    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            with open(file_path, newline="", encoding=encoding) as handle:
                sample = handle.read(4096)
                handle.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample)
                except csv.Error:
                    dialect = csv.excel
                rows = [row for row in csv.reader(handle, dialect)]
            break
        except UnicodeDecodeError:
            continue

    content = _table_to_text(rows)
    return [Document(page_content=content, metadata={"source": filename})] if content else []


def _load_text_documents(file_path: str, filename: str) -> list[Document]:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            content = Path(file_path).read_text(encoding=encoding).strip()
            if content:
                return [Document(page_content=content, metadata={"source": filename})]
            return []
        except UnicodeDecodeError:
            continue

    return []


def load_documents_lightweight(file_path: str, filename: str) -> list[Document]:
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _load_pdf_documents(file_path, filename)
    if suffix == ".docx":
        return _load_docx_documents(file_path, filename)
    if suffix == ".xlsx":
        return _load_xlsx_documents(file_path, filename)
    if suffix == ".csv":
        return _load_csv_documents(file_path, filename)
    if suffix in {".txt", ".md", ".markdown"}:
        return _load_text_documents(file_path, filename)

    raise ValueError(f"Unsupported document type for lightweight ingest: {suffix or 'unknown'}")
