from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from typing import Any, Callable, Mapping

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SUPPORTED_MEMO_TYPES = {
    "memo_internal": "Memo Internal",
    "nota_dinas": "Nota Dinas",
    "arahan": "Arahan",
}

HEADER_LINES = [
    "KEMENTERIAN SEKRETARIAT NEGARA RI",
    "SEKRETARIAT PRESIDEN",
    "ISTANA KEPRESIDENAN YOGYAKARTA",
]

FOOTER_NOTICE = (
    "Dokumen ini telah ditandatangani secara elektronik menggunakan sertifikat elektronik\n"
    "yang diterbitkan oleh Balai Sertifikasi Elektronik (BSrE)."
)

DEFAULT_GENERATED_CLOSING = "Demikian disampaikan untuk menjadi perhatian dan tindak lanjut sebagaimana mestinya."

PERSON_DATA_LABEL_PATTERN = (
    r"nama(?:\s+(?:pic|pegawai)(?:\s+yang\s+benar)?)?"
    r"|nip(?:\s+yang\s+bersangkutan)?"
    r"|pangkat/gol(?:ongan|\.)?"
    r"|jabatan"
    r"|unit\s+kerja"
    r"|jadwal(?:\s+(?:pendampingan|kegiatan|pelaksanaan))?"
    r"|hari/tanggal"
    r"|pukul"
    r"|tempat"
    r"|agenda"
    r"|(?:estimasi\s+)?durasi(?:\s+(?:rapat|kegiatan|pelaksanaan))?"
    r"|batas\s+waktu"
    r"|lokasi(?:\s+(?:asal|tujuan))?"
    r"|periode(?:\s+(?:kegiatan|pelaksanaan))?"
    r"|waktu\s+pelaksanaan"
    r"|waktu\s+kejadian"
    r"|pengguna\s+terdampak"
    r"|peserta"
    r"|nomor\s+kontak"
    r"|keperluan"
)

PERSON_DATA_PATTERN = re.compile(
    r"^(?:\d+[.)]\s*)?"
    rf"(?P<label>{PERSON_DATA_LABEL_PATTERN})"
    r"\s*(?::|adalah)\s*(?P<value>.+)$",
    re.IGNORECASE,
)

SIGNATURE_TABLE_INDENT_IN = 3.92
SIGNATURE_TABLE_WIDTH_IN = 1.56
SIGNATURE_SPACE_BEFORE_SHORT_PT = 110
SIGNATURE_SPACE_BEFORE_MEDIUM_PT = 76
SIGNATURE_SPACE_BEFORE_DEFAULT_PT = 56
SIGNATURE_SPACE_BEFORE_COMPACT_PT = 84
SIGNATURE_SPACE_BEFORE_COMPACT_DENSE_PT = 56
SEPARATOR_SPACE_AFTER_PT = 39

AI_UNAVAILABLE_MARKERS = (
    "semua layanan ai sedang tidak tersedia",
    "silakan coba lagi nanti",
)

INDONESIAN_SMALL_NUMBERS = {
    "satu": 1,
    "dua": 2,
    "tiga": 3,
    "empat": 4,
    "lima": 5,
    "enam": 6,
}

INSTRUCTION_ARTIFACT_PATTERNS = (
    r"\bperbaiki\s+typo\b",
    r"\bbagian\s+lain\s+jangan\s+diubah\b",
    r"\bmetadata\s+jangan\s+berubah\b",
    r"\bjangan\s+(?:diubah|berubah|mengubah|ubah)\b",
    r"\b(?:pertahankan|mempertahankan)\s+(?:seluruh\s+)?data\b.*\btanpa\s+perubahan\b",
    r"\btanpa\s+perubahan\b.*\bdata\b",
    r"\binstruksi\s+revisi\b",
    r"\barahan\s+tambahan\b",
    r"\bkontrol\s+kerja\b",
    r"\bpenutup\s+manual\s+apa\s+adanya\b",
    r"\bpenutup\s+manual\b.*\bapa\s+adanya\b",
    r"\b(?:pertahankan|mempertahankan)\s+penutup\s+manual\b",
    r"^catatan\s*:.*\b(?:penutup|manual|dipertahankan|format|instruksi|revisi)\b",
)

FOREIGN_ITALIC_PATTERNS = (
    r"\be-book\b",
    r"\bsign\s+up\b",
    r"\bcritical\s+thinking\b",
    r"\bKnow\s+Your\s+Customer\b",
    r"\bonline\b",
    r"\boffline\b",
    r"\bhybrid\b",
    r"\bvirtual\s+meeting\b",
)

ORDINAL_ITEM_PATTERN = re.compile(
    r"^(?:pertama|kedua|ketiga|keempat|kelima|keenam|ketujuh|kedelapan|kesembilan|kesepuluh),\s+(.+)$",
    re.IGNORECASE,
)

ACTIVITY_KEY_VALUE_LABEL_KEYS = {
    "hari/tanggal",
    "pukul",
    "tempat",
    "agenda",
    "durasi",
    "estimasi durasi",
    "jadwal",
    "jadwal pendampingan",
    "batas waktu",
    "lokasi",
    "lokasi asal",
    "lokasi tujuan",
    "periode",
    "waktu pelaksanaan",
    "waktu kejadian",
}


@dataclass(slots=True)
class MemoDraft:
    filename: str
    content: bytes
    searchable_text: str
    page_size: str


def normalize_memo_type(memo_type: str) -> str:
    normalized = (memo_type or "").strip().lower().replace("-", "_")
    if normalized not in SUPPORTED_MEMO_TYPES:
        raise ValueError("Jenis memo tidak didukung.")
    return normalized


def build_memo_prompt(
    memo_type: str,
    title: str,
    context: str,
    configuration: Mapping[str, Any] | None = None,
) -> str:
    label = SUPPORTED_MEMO_TYPES[normalize_memo_type(memo_type)]
    config = _normalize_configuration(configuration, title, context)
    content_source = context.strip() if config["revision_instruction"] else (config["content"] or context.strip())
    revision_section = (
        "Instruksi revisi wajib diterapkan:\n"
        f"{config['revision_instruction']}\n\n"
        if config["revision_instruction"]
        else ""
    )
    closing_rule = (
        "- Jangan menulis kalimat penutup akhir karena bagian penutup sudah disediakan konfigurasi.\n"
        if config["closing"]
        else "- Jangan menulis kalimat penutup akhir kecuali user mengisinya di field Penutup.\n"
    )
    revision_rules = (
        "- Karena ini revisi, pertahankan kalimat, urutan, nomor butir, dan informasi yang tidak diminta berubah secara eksplisit.\n"
        "- Jangan meregenerasi seluruh memo; ubah hanya bagian yang disebut dalam instruksi revisi.\n"
        if config["revision_instruction"]
        else ""
    )

    return (
        "Tulis isi memorandum resmi dalam Bahasa Indonesia dengan gaya naskah dinas.\n"
        f"Jenis: {label}\n"
        f"Nomor: {config['number']}\n"
        f"Yth.: {config['recipient']}\n"
        f"Dari: {config['sender']}\n"
        f"Hal: {config['subject']}\n"
        f"Tanggal: {config['date']}\n\n"
        "Konteks/dasar:\n"
        f"{config['basis'] or '-'}\n\n"
        "Isi atau poin wajib:\n"
        f"{content_source}\n\n"
        f"{revision_section}"
        "Arahan tambahan:\n"
        f"{config['additional_instruction'] or '-'}\n\n"
        "Aturan keluaran:\n"
        "- Tulis hanya isi utama memo, tanpa kop, nomor, Yth., Dari, Hal, Tanggal, tanda tangan, tembusan, atau footer.\n"
        "- Gunakan paragraf formal yang singkat, jelas, dan mengikuti contoh memorandum manual.\n"
        "- Gunakan rumusan naskah dinas yang hemat, misalnya 'Sehubungan hal tersebut, dapat kami sampaikan sebagai berikut.' bila sesuai konteks.\n"
        "- Hindari frasa generik atau terlalu operasional seperti 'beberapa hal yang perlu diperhatikan' bila data dapat langsung disampaikan.\n"
        "- Jika ada beberapa butir keputusan/permohonan, gunakan daftar bernomor 1., 2., 3.\n"
        "- Jika input sudah memakai penomoran 1., 2., 3., pertahankan nomor dan urutan tersebut; jangan ubah menjadi Pertama/Kedua/Ketiga.\n"
        "- Awali dengan dasar atau tindak lanjut bila konteks menyediakannya.\n"
        "- Jangan mengarang nama orang, NIP, jabatan, nomor kontak, unit kerja, atau PIC bila tidak tertulis eksplisit di konfigurasi.\n"
        "- Instruksi revisi dan arahan tambahan adalah kontrol kerja, bukan bagian naskah; jangan salin frasa seperti 'jangan diubah', 'metadata jangan berubah', atau 'perbaiki typo'.\n"
        "- Perlakukan kata seperti baseline, uji, skenario evaluasi, dan auto format sebagai instruksi internal; jangan salin ke naskah memo.\n"
        "- Jangan menulis blok Tembusan karena tembusan diambil dari konfigurasi.\n"
        "- Jangan mencantumkan sumber, URL, JSON, kutipan tool, atau blok [SOURCES: ...] dalam naskah memo.\n"
        "- Untuk data PIC/pegawai, tulis setiap label dari konfigurasi sebagai baris terpisah; jangan menggabungkan nama, NIP, jabatan, unit kerja, keperluan, jadwal, atau nomor kontak ke dalam paragraf naratif.\n"
        "- Untuk detail kegiatan seperti hari/tanggal, pukul, dan tempat, tulis setiap label sebagai baris terpisah seperti naskah dinas resmi.\n"
        "- Jika field Penutup berisi teks, jangan ubah atau hilangkan kalimat penutup tersebut.\n"
        f"{revision_rules}"
        "- Jangan gunakan markdown, tabel, salam pembuka, atau salam penutup.\n"
        f"{closing_rule}"
    )


def generate_memo_docx(
    memo_type: str,
    title: str,
    context: str,
    text_generator: Callable[[str], str] | None = None,
    configuration: Mapping[str, Any] | None = None,
) -> MemoDraft:
    normalized_type = normalize_memo_type(memo_type)
    clean_title = _clean_title(title)
    clean_context = (context or "").strip()

    if not clean_context:
        raise ValueError("Konteks memo wajib diisi.")

    config = _normalize_configuration(configuration, clean_title, clean_context)
    generator = text_generator or _default_text_generator
    prompt = build_memo_prompt(normalized_type, clean_title, clean_context, config)
    raw_body = config["body_override"] or generator(prompt)
    body = _sanitize_memo_body(_normalize_generated_text(raw_body), config)
    body = _enforce_revision_constraints(body, config)
    body = _preserve_configured_numbered_items(body, config)
    body, generated_closing = _separate_generated_closing(body, config)
    if generated_closing and not config["closing"]:
        config["closing"] = generated_closing
    elif not config["closing"]:
        config["closing"] = DEFAULT_GENERATED_CLOSING
    config["page_size"] = _resolve_page_size(config, body)

    document = _build_official_memo_document(normalized_type, config, body)

    buffer = BytesIO()
    document.save(buffer)

    searchable_text = _build_searchable_text(normalized_type, config, body)

    return MemoDraft(
        filename=f"{_slugify(clean_title)}.docx",
        content=buffer.getvalue(),
        searchable_text=searchable_text,
        page_size=config["page_size"],
    )


def _default_text_generator(prompt: str) -> str:
    from app.llm_manager import get_llm_stream

    chunks: list[str] = []
    for chunk in get_llm_stream([{"role": "user", "content": prompt}]):
        chunks.append(chunk)

    return "".join(chunks)


def _build_official_memo_document(memo_type: str, config: dict[str, str], body: str) -> Document:
    document = Document()
    document.core_properties.title = config["subject"]
    document.core_properties.subject = SUPPORTED_MEMO_TYPES[memo_type]

    _apply_section_layout(document.sections[0], config["page_size"])
    _apply_base_styles(document)
    _add_header(document)
    _add_document_title(document, memo_type, config["number"])
    _add_metadata(document, config)
    compact_layout = _should_use_compact_layout(config, body)
    _add_separator(document, compact=compact_layout)
    _add_body(document, body, compact=compact_layout, config=config)
    _add_closing(document, config["closing"], compact=compact_layout)
    signature_space_before = _signature_space_before(config, body, compact=compact_layout)
    _add_signature_placeholder(
        document,
        config["signatory"],
        compact=compact_layout,
        space_before_pt=signature_space_before,
    )
    _add_carbon_copy(document, config["carbon_copy"], compact=compact_layout)
    if config["signatory"].strip():
        _add_footer(document.sections[0])

    return document


def _apply_section_layout(section, page_size: str) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11 if page_size == "letter" else 14)
    section.top_margin = Inches(0.5)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(0.98)
    section.bottom_margin = Inches(0.38)
    section.footer_distance = Inches(0.22)


def _apply_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    _set_style_font(normal, "Arial")

    for style_name in ["Header", "Footer"]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        _set_style_font(style, "Arial")


def _add_header(document: Document) -> None:
    for index, line in enumerate(HEADER_LINES):
        paragraph = document.add_paragraph()
        _format_paragraph(
            paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing_pt=14.95,
            space_after_pt=33 if index == len(HEADER_LINES) - 1 else 0,
        )
        run = paragraph.add_run(line)
        _format_run(run, size_pt=12, bold=True, color="4D4D4D")


def _add_document_title(document: Document, memo_type: str, number: str) -> None:
    title = "MEMORANDUM" if memo_type == "memo_internal" else SUPPORTED_MEMO_TYPES[memo_type].upper()

    paragraph = document.add_paragraph()
    _format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=13.8)
    run = paragraph.add_run(title)
    _format_run(run, size_pt=11, bold=True)

    paragraph = document.add_paragraph()
    _format_paragraph(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing_pt=13.8,
        space_after_pt=10,
    )
    run = paragraph.add_run(f"Nomor {number}")
    _format_run(run, size_pt=11)


def _add_metadata(document: Document, config: dict[str, str]) -> None:
    rows = [
        ("Yth.", config["recipient"]),
        ("Dari", config["sender"]),
        ("Hal", config["subject"]),
        ("Tanggal", config["date"]),
    ]
    table = document.add_table(rows=len(rows), cols=3)
    table.autofit = False

    widths = [Inches(0.62), Inches(0.12), Inches(5.52)]
    for row_index, (label, value) in enumerate(rows):
        cells = table.rows[row_index].cells
        for cell_index, width in enumerate(widths):
            cells[cell_index].width = width
            cells[cell_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(cells[cell_index], top=0, start=0, bottom=0, end=0)

        _set_cell_text(cells[0], label)
        _set_cell_text(cells[1], ":")
        _set_cell_text(cells[2], value)


def _add_separator(document: Document, *, compact: bool = False) -> None:
    paragraph = document.add_paragraph()
    _format_paragraph(
        paragraph,
        line_spacing_pt=1,
        space_before_pt=16,
        space_after_pt=SEPARATOR_SPACE_AFTER_PT,
    )
    _set_paragraph_bottom_border(paragraph)


def _add_body(
    document: Document,
    body: str,
    *,
    compact: bool = False,
    config: dict[str, str] | None = None,
) -> None:
    blocks = _split_blocks(body)
    configured_key_values = _extract_configured_key_value_items(config or {})
    configured_key_values_used = False

    if len(configured_key_values) >= 2 and not _blocks_contain_key_value_table(blocks):
        insert_at = 1 if blocks and not _is_structured_body_block(blocks[0]) else 0
        configured_blocks = [f"{label}: {value}" for label, value in configured_key_values]
        blocks = [*blocks[:insert_at], *configured_blocks, *blocks[insert_at:]]
        blocks = _remove_redundant_blocks_after_configured_key_values(
            blocks,
            configured_key_values,
            start_index=insert_at + len(configured_blocks),
        )
    elif len(configured_key_values) >= 2:
        blocks = _clean_redundant_blocks_around_existing_key_values(blocks, configured_key_values)

    index = 0
    line_spacing_pt = 13.1 if compact else 13.8

    while index < len(blocks):
        key_value_items = _collect_person_data_blocks(blocks, index)
        if len(key_value_items) >= 2:
            table_items = key_value_items
            if len(configured_key_values) >= 2 and not configured_key_values_used:
                table_items = _merge_key_value_items(key_value_items, configured_key_values)
                configured_key_values_used = True

            _add_key_value_body_table(
                document,
                table_items,
                compact=compact,
                add_space_before=index > 0,
                add_space_after=_should_add_space_after_key_value_table(blocks, index + len(key_value_items)),
            )
            index += len(key_value_items)
            continue

        block = blocks[index]
        numbered = re.match(r"^(\d+)[.)]\s+(.+)$", block)
        bulleted = re.match(r"^([-*•])\s+(.+)$", block)

        if numbered:
            paragraph = document.add_paragraph()
            _format_paragraph(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                line_spacing_pt=line_spacing_pt,
                left_indent_in=0.28,
                first_line_indent_in=-0.28,
                keep_together=True,
            )
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.28))
            _append_text_run(paragraph, f"{numbered.group(1)}.\t{numbered.group(2).strip()}")
            index += 1
            continue

        if bulleted:
            paragraph = document.add_paragraph()
            _format_paragraph(
                paragraph,
                alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                line_spacing_pt=line_spacing_pt,
                left_indent_in=0.28,
                first_line_indent_in=-0.28,
                keep_together=True,
            )
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.28))
            _append_text_run(paragraph, f"•\t{bulleted.group(2).strip()}")
            index += 1
            continue

        paragraph = document.add_paragraph()
        _format_paragraph(
            paragraph,
            alignment=_body_alignment(block),
            line_spacing_pt=line_spacing_pt,
            first_line_indent_in=0.5,
            keep_together=True,
        )
        _append_text_run(paragraph, block)
        index += 1


def _blocks_contain_key_value_table(blocks: list[str]) -> bool:
    for index in range(len(blocks)):
        if len(_collect_person_data_blocks(blocks, index)) >= 2:
            return True
    return False


def _key_value_table_span(blocks: list[str]) -> tuple[int, int] | None:
    for index in range(len(blocks)):
        key_value_items = _collect_person_data_blocks(blocks, index)
        if len(key_value_items) >= 2:
            return index, index + len(key_value_items)
    return None


def _clean_redundant_blocks_around_existing_key_values(
    blocks: list[str],
    configured_items: list[tuple[str, str]],
) -> list[str]:
    span = _key_value_table_span(blocks)
    if not span:
        return blocks

    start_index, end_index = span
    configured_labels = {_key_value_label_key(label) for label, _value in configured_items}
    configured_values = _normalized_configured_values(configured_items)

    if (
        start_index > 0
        and configured_labels & ACTIVITY_KEY_VALUE_LABEL_KEYS
        and _looks_like_redundant_activity_detail_block(
            blocks[start_index - 1],
            configured_labels,
            configured_values,
        )
    ):
        blocks = [*blocks[: start_index - 1], *blocks[start_index:]]
        start_index -= 1
        end_index -= 1

    return _remove_redundant_blocks_after_configured_key_values(
        blocks,
        configured_items,
        start_index=end_index,
    )


def _should_add_space_after_key_value_table(blocks: list[str], next_index: int) -> bool:
    if next_index >= len(blocks):
        return False

    next_block = blocks[next_index].strip()
    return bool(next_block)


def _remove_redundant_blocks_after_configured_key_values(
    blocks: list[str],
    configured_items: list[tuple[str, str]],
    *,
    start_index: int,
) -> list[str]:
    if not configured_items or start_index >= len(blocks):
        return blocks

    configured_labels = {_key_value_label_key(label) for label, _value in configured_items}
    configured_values = _normalized_configured_values(configured_items)
    has_activity_details = bool(configured_labels & ACTIVITY_KEY_VALUE_LABEL_KEYS)
    output = blocks[:start_index]
    removed = False

    for block in blocks[start_index:]:
        parsed = _parse_person_data_block(block)
        if parsed and _key_value_label_key(parsed[0]) in configured_labels:
            removed = True
            continue

        if has_activity_details and _looks_like_redundant_activity_detail_block(
            block,
            configured_labels,
            configured_values,
        ):
            removed = True
            continue

        output.append(block)

    return _renumber_numbered_block_sequences(output, from_index=start_index) if removed else blocks


def _looks_like_redundant_activity_detail_block(
    block: str,
    configured_labels: set[str],
    configured_values: set[str],
) -> bool:
    stripped = block.strip()
    is_list_block = re.match(r"^(?:\d+[.)]|[-*•])\s+", stripped) is not None
    normalized = _normalize_comparison_text(re.sub(r"^\s*(?:\d+[.)]|[-*•])\s*", "", stripped))
    if not normalized:
        return False

    matched_value_count = _configured_value_match_count(normalized, configured_values)
    configured_token_coverage = _configured_value_token_coverage(normalized, configured_values)
    has_activity_label = any(label in normalized for label in configured_labels & ACTIVITY_KEY_VALUE_LABEL_KEYS)
    starts_with_detail = normalized.startswith(
        (
            "hari/tanggal",
            "tanggal",
            "pukul",
            "jam",
            "tempat",
            "agenda",
            "durasi",
            "estimasi durasi",
            "lokasi",
            "periode",
            "waktu pelaksanaan",
            "waktu kejadian",
            "rapat akan dilaksanakan",
            "rapat dilaksanakan",
            "kegiatan akan dilaksanakan",
            "kegiatan dilaksanakan",
        )
    )

    if not is_list_block:
        return (
            (starts_with_detail and (has_activity_label or matched_value_count >= 1))
            or matched_value_count >= 2
            or configured_token_coverage >= 0.66
            or (normalized.startswith("agenda ") and matched_value_count >= 1)
        )

    detail_starters = (
        "hari/tanggal",
        "tanggal",
        "pukul",
        "jam",
        "tempat",
        "agenda",
        "durasi",
        "estimasi durasi",
        "lokasi",
        "periode",
        "waktu pelaksanaan",
        "waktu kejadian",
        "rapat akan dilaksanakan",
        "rapat dilaksanakan",
        "kegiatan akan dilaksanakan",
        "kegiatan dilaksanakan",
    )
    if not normalized.startswith(detail_starters):
        return matched_value_count >= 2 or configured_token_coverage >= 0.66

    if normalized.endswith("pada:") or normalized.endswith("pada"):
        return bool(configured_labels & {"hari/tanggal", "pukul", "tempat"})

    if has_activity_label:
        return True

    return matched_value_count >= 1


def _normalized_configured_values(configured_items: list[tuple[str, str]]) -> set[str]:
    return {
        _normalize_comparison_text(value)
        for _label, value in configured_items
        if len(_normalize_comparison_text(value)) >= 5
    }


def _configured_value_match_count(normalized_block: str, configured_values: set[str]) -> int:
    return sum(
        1
        for value in configured_values
        if _configured_value_matches_block(normalized_block, value)
    )


def _configured_value_token_coverage(normalized_block: str, configured_values: set[str]) -> float:
    block_tokens = set(_significant_comparison_tokens(_flatten_comparison_text(normalized_block)))
    if not block_tokens:
        return 0

    configured_tokens: set[str] = set()
    for value in configured_values:
        configured_tokens.update(_significant_comparison_tokens(_flatten_comparison_text(value)))

    if not configured_tokens:
        return 0

    return len(block_tokens & configured_tokens) / len(block_tokens)


def _configured_value_matches_block(normalized_block: str, normalized_value: str) -> bool:
    if not normalized_block or not normalized_value:
        return False
    if normalized_value in normalized_block:
        return True

    flat_block = _flatten_comparison_text(normalized_block)
    flat_value = _flatten_comparison_text(normalized_value)
    if flat_value and flat_value in flat_block:
        return True

    value_tokens = _significant_comparison_tokens(flat_value)
    if len(value_tokens) < 3:
        return False

    block_tokens = set(_significant_comparison_tokens(flat_block))
    matched = sum(1 for token in set(value_tokens) if token in block_tokens)
    return matched >= 3 and matched / len(set(value_tokens)) >= 0.72


def _flatten_comparison_text(text: str) -> str:
    return re.sub(r"[^0-9a-zA-ZÀ-ÿ/]+", " ", text or "").strip()


def _significant_comparison_tokens(text: str) -> list[str]:
    stopwords = {
        "dan",
        "serta",
        "yang",
        "dengan",
        "pada",
        "untuk",
        "dari",
        "ke",
        "di",
        "dalam",
        "akan",
        "telah",
        "rapat",
        "kegiatan",
    }
    return [
        token
        for token in re.findall(r"[0-9a-zA-ZÀ-ÿ/]+", text or "")
        if len(token) > 2 and token.lower() not in stopwords
    ]


def _renumber_numbered_block_sequences(blocks: list[str], *, from_index: int = 0) -> list[str]:
    output = blocks[:from_index]
    next_number = 1

    for block in blocks[from_index:]:
        match = re.match(r"^\s*\d+[.)]\s+(.+)$", block)
        if match:
            output.append(f"{next_number}. {match.group(1).strip()}")
            next_number += 1
            continue

        output.append(block)

    return output


def _extract_configured_key_value_items(config: dict[str, str]) -> list[tuple[str, str]]:
    content = config.get("content", "")
    if not content:
        return []

    normalized = _normalize_person_data_sequences(content)
    items: list[tuple[str, str]] = []
    for block in _split_blocks(normalized):
        parsed = _parse_person_data_block(block)
        if parsed:
            items.append(parsed)

    items.extend(_extract_activity_key_value_items(content))
    return _dedupe_key_value_items(items)


def _merge_key_value_items(
    generated_items: list[tuple[str, str]],
    configured_items: list[tuple[str, str]],
) -> list[tuple[str, str]]:
    merged = _dedupe_key_value_items(configured_items)
    configured_labels = {_key_value_label_key(label) for label, _value in merged}

    for label, value in generated_items:
        if _key_value_label_key(label) in configured_labels:
            continue
        merged.append((label, value))

    return merged


def _dedupe_key_value_items(items: list[tuple[str, str]]) -> list[tuple[str, str]]:
    output: list[tuple[str, str]] = []
    seen: set[str] = set()

    for label, value in items:
        key = _key_value_label_key(label)
        clean_value = _clean_key_value_value(label, value)
        if not key or not clean_value or key in seen:
            continue
        output.append((label, clean_value))
        seen.add(key)

    return output


def _key_value_label_key(label: str) -> str:
    return _normalize_comparison_text(label)


def _extract_activity_key_value_items(text: str) -> list[tuple[str, str]]:
    clean = re.sub(r"\s+", " ", text or "").strip()
    if not clean:
        return []

    items: list[tuple[str, str]] = []
    date_pattern = r"\d{1,2}\s+[A-Za-zÀ-ÿ]+\s+\d{4}"
    time_pattern = (
        r"\d{1,2}[.:]\d{2}"
        r"(?:\s*(?:s\.?d\.?|-|hingga|sampai)\s*\d{1,2}[.:]\d{2})?"
        r"\s*WIB"
    )
    date_match = re.search(
        rf"\b(?:pada|tanggal(?:\s+kegiatan|\s+pelaksanaan|\s+rapat)?|"
        rf"dilaksanakan(?:\s+pada)?|rapat\s+dilaksanakan|kegiatan\s+dilaksanakan)"
        rf"\s+(?:hari\s+)?(?P<date>{date_pattern})\b",
        clean,
        flags=re.IGNORECASE,
    )
    if not date_match:
        date_match = re.search(
            rf"\b(?P<date>{date_pattern})\b(?=[, ]+(?:pukul|jam|bertempat|di|ruang|agenda))",
            clean,
            flags=re.IGNORECASE,
        )
    time_match = re.search(
        rf"\b(?:pukul|jam)\s+(?P<time>{time_pattern})\b",
        clean,
        flags=re.IGNORECASE,
    )
    place_match = re.search(
        rf"\bdi\s+(?P<place>[^,.]+?)\s+pada\s+(?:hari\s+)?{date_pattern}\b",
        clean,
        flags=re.IGNORECASE,
    )
    if not place_match and time_match:
        place_match = re.search(
            rf"\b(?:pukul|jam)\s+{time_pattern}\s*,\s*(?P<place>(?:ruang|gedung|aula|tempat|lokasi)[^,.;]+)",
            clean,
            flags=re.IGNORECASE,
        )
    if not place_match:
        place_match = re.search(
            r"\bbertempat\s+di\s+(?P<place>[^,.;]+)",
            clean,
            flags=re.IGNORECASE,
        )
    agenda_match = re.search(
        r"\b(?:dengan\s+)?agenda\s*:?\s+(?P<agenda>[^.]+)",
        clean,
        flags=re.IGNORECASE,
    )
    duration_match = re.search(
        r"\b(?P<label>estimasi\s+durasi|durasi(?:\s+(?:rapat|kegiatan|pelaksanaan))?)\s*:?\s+"
        r"(?P<duration>[^,.;]+)",
        clean,
        flags=re.IGNORECASE,
    )
    if not duration_match:
        duration_match = re.search(
            r"\b(?:selama|berdurasi)\s+(?P<duration>\d+\s*(?:jam|menit))\b",
            clean,
            flags=re.IGNORECASE,
        )

    if date_match:
        items.append(("hari/tanggal", _clean_activity_value(date_match.group("date"))))
    if time_match:
        items.append(("pukul", _clean_activity_value(time_match.group("time"))))
    if place_match:
        items.append(("tempat", _clean_activity_value(place_match.group("place"))))
    if agenda_match:
        items.append(("agenda", _clean_activity_value(agenda_match.group("agenda"))))
    if duration_match:
        label = duration_match.groupdict().get("label") or "estimasi durasi"
        items.append((_normalize_person_data_label(label), _clean_activity_value(duration_match.group("duration"))))

    return items if len(items) >= 2 else []


def _clean_activity_value(value: str) -> str:
    clean = re.sub(r"\s+", " ", value or "").strip(" ,;")
    clean = re.sub(r"^\d+[.)]\s+", "", clean).strip()
    return clean.rstrip(" ,;.")


def _clean_key_value_value(label: str, value: str) -> str:
    key = _key_value_label_key(label)
    if key in ACTIVITY_KEY_VALUE_LABEL_KEYS:
        clean = _clean_activity_value(value)
        return _trim_activity_value_tail(key, clean)
    return _clean_inline_person_data_value(value)


def _trim_activity_value_tail(key: str, value: str) -> str:
    clean = value.strip()
    if not clean:
        return ""

    if key in {"hari/tanggal", "periode", "waktu pelaksanaan", "batas waktu"}:
        date_match = re.match(
            r"(?P<value>.*?\d{1,2}\s+[A-Za-zÀ-ÿ]+\s+\d{4}"
            r"(?:\s*(?:pukul|jam)\s+\d{1,2}[.:]\d{2}(?:\s*WIB)?)?)"
            r"(?:\.\s+[A-ZÀ-Ý].*)?$",
            clean,
        )
        if date_match:
            clean = date_match.group("value")

    if key in {"lokasi", "lokasi asal", "lokasi tujuan"}:
        clean = re.split(r"\.\s+(?=[A-ZÀ-Ý])", clean, maxsplit=1)[0]

    return clean.rstrip(" ,;.")


def _add_key_value_body_table(
    document: Document,
    items: list[tuple[str, str]],
    *,
    compact: bool = False,
    add_space_before: bool = False,
    add_space_after: bool = False,
) -> None:
    if add_space_before:
        spacer = document.add_paragraph()
        _format_paragraph(
            spacer,
            line_spacing_pt=1,
            space_before_pt=3 if compact else 5,
            space_after_pt=5 if compact else 7,
        )

    table = document.add_table(rows=len(items), cols=3)
    table.autofit = False
    _set_table_indent(table, 0.5)
    _set_table_width(table, 5.22)

    widths = [Inches(1.45), Inches(0.18), Inches(3.59)]
    for row_index, (label, value) in enumerate(items):
        row = table.rows[row_index]
        _set_row_cant_split(row)
        cells = row.cells
        for cell_index, width in enumerate(widths):
            cells[cell_index].width = width
            cells[cell_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(cells[cell_index], top=10, start=0, bottom=10, end=0)
        _set_cell_margins(cells[2], top=10, start=80, bottom=10, end=0)

        _set_cell_text(cells[0], label)
        _set_cell_text(cells[1], ":")
        _set_cell_text(cells[2], value)

    if add_space_after:
        spacer = document.add_paragraph()
        _format_paragraph(
            spacer,
            line_spacing_pt=1,
            space_before_pt=0,
            space_after_pt=12 if compact else 16,
        )


def _add_closing(document: Document, closing: str, *, compact: bool = False) -> None:
    if not closing:
        return

    paragraph = document.add_paragraph()
    _format_paragraph(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        line_spacing_pt=13.1 if compact else 13.8,
        first_line_indent_in=0.5,
        space_before_pt=14 if compact else 20,
        keep_together=True,
        keep_with_next=True,
    )
    _append_text_run(paragraph, closing)


def _add_signature_placeholder(
    document: Document,
    signatory: str,
    *,
    compact: bool = False,
    space_before_pt: float | None = None,
) -> None:
    if not signatory.strip():
        return

    if space_before_pt is None:
        space_before_pt = SIGNATURE_SPACE_BEFORE_COMPACT_PT if compact else SIGNATURE_SPACE_BEFORE_DEFAULT_PT

    spacer = document.add_paragraph()
    _format_paragraph(
        spacer,
        line_spacing_pt=1,
        space_before_pt=space_before_pt,
        space_after_pt=0,
        keep_together=True,
        keep_with_next=True,
    )

    table = document.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_indent(table, SIGNATURE_TABLE_INDENT_IN)
    _set_table_width(table, SIGNATURE_TABLE_WIDTH_IN)

    widths = [Inches(0.35), Inches(0.86), Inches(0.35)]
    for row in table.rows:
        _set_row_cant_split(row)
        for cell_index, width in enumerate(widths):
            cell = row.cells[cell_index]
            cell.width = width
            _set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

    table.rows[0].height = Inches(0.86)
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    qr_cell = table.cell(0, 1)
    qr_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_border(qr_cell, color="777777")
    paragraph = qr_cell.paragraphs[0]
    _format_paragraph(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing_pt=8,
        keep_together=True,
        keep_with_next=True,
    )
    run = paragraph.add_run("QR\nTTE")
    _format_run(run, size_pt=6, color="777777")

    name_cell = table.cell(1, 0).merge(table.cell(1, 2))
    name = name_cell.paragraphs[0]
    _format_paragraph(
        name,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        line_spacing_pt=13.8,
        space_before_pt=10,
        keep_together=True,
        keep_with_next=True,
    )
    _append_text_run(name, signatory)


def _add_carbon_copy(document: Document, carbon_copy: str, *, compact: bool = False) -> None:
    lines = _split_lines(carbon_copy)
    if not lines:
        return

    heading = document.add_paragraph()
    _format_paragraph(
        heading,
        line_spacing_pt=13.1 if compact else 13.8,
        space_before_pt=24 if compact else 28,
        keep_with_next=True,
        keep_together=True,
    )
    _append_text_run(heading, "Tembusan:")

    should_auto_number = len(lines) > 1 and not any(re.match(r"^\d+[.)]\s+", line) for line in lines)
    for index, line in enumerate(lines, start=1):
        text = f"{index}. {line}" if should_auto_number else line
        numbered = re.match(r"^(\d+)[.)]\s+(.+)$", text)

        paragraph = document.add_paragraph()
        if numbered:
            _format_paragraph(
                paragraph,
                line_spacing_pt=13.1 if compact else 13.8,
                left_indent_in=0.28,
                first_line_indent_in=-0.28,
                keep_together=True,
                keep_with_next=index < len(lines),
            )
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(0.28))
            _append_text_run(paragraph, f"{numbered.group(1)}.\t{numbered.group(2).strip()}")
        else:
            _format_paragraph(
                paragraph,
                line_spacing_pt=13.1 if compact else 13.8,
                keep_together=True,
                keep_with_next=index < len(lines),
            )
            _append_text_run(paragraph, text)


def _add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    _format_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing_pt=8)
    run = paragraph.add_run(FOOTER_NOTICE)
    _format_run(run, size_pt=7)


def _set_cell_text(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    _format_paragraph(paragraph, line_spacing_pt=13.8)
    _append_text_run(paragraph, text)


def _format_paragraph(
    paragraph,
    *,
    alignment=None,
    line_spacing_pt: float = 13.8,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
    first_line_indent_in: float | None = None,
    left_indent_in: float | None = None,
    keep_together: bool = False,
    keep_with_next: bool = False,
) -> None:
    fmt = paragraph.paragraph_format
    if alignment is not None:
        paragraph.alignment = alignment
    fmt.line_spacing = Pt(line_spacing_pt)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    if first_line_indent_in is not None:
        fmt.first_line_indent = Inches(first_line_indent_in)
    if left_indent_in is not None:
        fmt.left_indent = Inches(left_indent_in)
    fmt.keep_together = keep_together
    fmt.keep_with_next = keep_with_next


def _append_text_run(paragraph, text: str):
    return _append_official_text_runs(paragraph, text, size_pt=11)


def _append_official_text_runs(paragraph, text: str, *, size_pt: float = 11):
    parts = _split_foreign_italic_parts(text)
    last_run = None
    for part, italic in parts:
        if not part:
            continue
        last_run = paragraph.add_run(part)
        _format_run(last_run, size_pt=size_pt, italic=italic)

    if last_run is None:
        last_run = paragraph.add_run("")
        _format_run(last_run, size_pt=size_pt)

    return last_run


def _split_foreign_italic_parts(text: str) -> list[tuple[str, bool]]:
    if not text:
        return [("", False)]

    combined = re.compile("|".join(f"({pattern})" for pattern in FOREIGN_ITALIC_PATTERNS), re.IGNORECASE)
    parts: list[tuple[str, bool]] = []
    cursor = 0
    for match in combined.finditer(text):
        if match.start() > cursor:
            parts.append((text[cursor : match.start()], False))
        parts.append((match.group(0), True))
        cursor = match.end()

    if cursor < len(text):
        parts.append((text[cursor:], False))

    return parts


def _format_run(run, *, size_pt: float, bold: bool = False, italic: bool = False, color: str | None = None) -> None:
    run.font.name = "Arial"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), "Arial")


def _set_style_font(style, font_name: str) -> None:
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ["ascii", "hAnsi", "eastAsia", "cs"]:
        r_fonts.set(qn(f"w:{key}"), font_name)


def _set_paragraph_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)


def _set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_border(cell, *, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ["top", "left", "bottom", "right"]:
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _set_table_indent(table, indent_in: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)

    tbl_ind.set(qn("w:w"), str(int(indent_in * 1440)))
    tbl_ind.set(qn("w:type"), "dxa")


def _set_table_width(table, width_in: float) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)

    tbl_w.set(qn("w:w"), str(int(width_in * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "1")


def _normalize_configuration(
    configuration: Mapping[str, Any] | None,
    title: str,
    context: str,
) -> dict[str, str]:
    raw = dict(configuration or {})

    page_size_mode = _clean_config_value(raw.get("page_size_mode")).lower()
    page_size = _clean_config_value(raw.get("page_size"), default="folio").lower()

    if page_size_mode == "auto" or page_size == "auto":
        page_size = "auto"
    elif page_size not in {"folio", "letter"}:
        page_size = "folio"

    return {
        "number": _clean_config_value(raw.get("number"), default="-"),
        "recipient": _clean_config_value(raw.get("recipient"), default="-"),
        "sender": _clean_config_value(raw.get("sender"), default="Kepala Istana Kepresidenan Yogyakarta"),
        "subject": _clean_config_value(raw.get("subject"), default=title),
        "date": _clean_config_value(raw.get("date"), default="-"),
        "basis": _clean_config_value(raw.get("basis")),
        "content": _clean_config_value(raw.get("content"), default=context),
        "closing": _clean_config_value(raw.get("closing")),
        "signatory": _clean_config_value(raw.get("signatory")),
        "carbon_copy": _clean_config_value(raw.get("carbon_copy")),
        "body_override": _clean_config_value(raw.get("body_override")),
        "page_size": page_size,
        "page_size_mode": page_size_mode,
        "additional_instruction": _clean_config_value(raw.get("additional_instruction")),
        "revision_instruction": _clean_config_value(raw.get("revision_instruction")),
    }


def _clean_config_value(value: Any, default: str = "") -> str:
    clean = re.sub(r"[ \t]+", " ", str(value or "").strip())
    return clean if clean else default


def _build_searchable_text(memo_type: str, config: dict[str, str], body: str) -> str:
    parts = [
        *HEADER_LINES,
        "",
        "MEMORANDUM" if memo_type == "memo_internal" else SUPPORTED_MEMO_TYPES[memo_type].upper(),
        f"Nomor {config['number']}",
        "",
        f"Yth.    : {config['recipient']}",
        f"Dari    : {config['sender']}",
        f"Hal     : {config['subject']}",
        f"Tanggal : {config['date']}",
        "",
        body,
        "",
        config["closing"],
        "",
        config["signatory"],
    ]

    if config["carbon_copy"]:
        parts.extend(["", "Tembusan:", config["carbon_copy"]])

    if config["signatory"].strip():
        parts.extend(["", FOOTER_NOTICE])
    return "\n".join(part for part in parts if part is not None).strip()


def _sanitize_memo_body(body: str, config: dict[str, str]) -> str:
    clean = _strip_external_source_artifacts(body)
    clean = _strip_forbidden_body_sections(clean)
    clean = _strip_configured_carbon_copy_lines(clean, config.get("carbon_copy", ""))
    clean = _strip_markdown_artifacts(clean)
    clean = _strip_evaluation_artifacts(clean)
    clean = _strip_body_closing_sentences(clean)
    clean = _remove_configured_closing(clean, config.get("closing", ""))
    clean = _strip_instruction_artifacts(clean, config)
    clean = _strip_unconfigured_honorific_data(clean, config)
    clean = _strip_unconfigured_operational_data(clean, config)
    clean = _replace_empty_person_data_placeholders(clean)
    clean = _strip_trailing_fragments(clean)
    clean = _normalize_person_data_sequences(clean)
    clean = _clean_official_language_fragments(clean)
    clean = _dedupe_consecutive_blocks(clean)
    return _normalize_generated_text(clean)


def _enforce_revision_constraints(body: str, config: dict[str, str]) -> str:
    instruction = " ".join(
        [
            config.get("revision_instruction", ""),
            config.get("additional_instruction", ""),
        ]
    ).lower()
    if not config.get("revision_instruction") or not instruction.strip():
        return body

    max_paragraphs = _extract_max_paragraphs(instruction)
    wants_shorter = any(marker in instruction for marker in ("lebih singkat", "ringkas", "dipersingkat"))
    blocks = _split_blocks(body)
    source_word_count = _word_count(config.get("content", ""))
    violates_short_request = wants_shorter and source_word_count > 0 and _word_count(body) >= source_word_count
    violates_paragraph_limit = max_paragraphs is not None and _paragraph_block_count(blocks) > max_paragraphs

    if not violates_short_request and not violates_paragraph_limit:
        return body

    concise = _build_concise_revision_body(
        config,
        fallback_body=body,
        paragraph_limit=max_paragraphs if wants_shorter else None,
    )
    if concise:
        concise_blocks = _split_blocks(concise)
        if max_paragraphs is None or _paragraph_block_count(concise_blocks) <= max_paragraphs:
            if not wants_shorter or _word_count(concise) < _word_count(body):
                return concise

    if max_paragraphs is not None:
        limited_blocks: list[str] = []
        paragraph_count = 0
        for block in blocks:
            if not _is_structured_body_block(block):
                paragraph_count += 1
            if paragraph_count > max_paragraphs:
                break
            limited_blocks.append(block)
        return "\n".join(limited_blocks).strip()

    return body


def _extract_max_paragraphs(instruction: str) -> int | None:
    match = re.search(
        r"(?:maksimal|paling banyak|tidak lebih dari|jangan lebih dari)\s+(\d+|[a-z]+)\s+paragraf",
        instruction,
        flags=re.IGNORECASE,
    )
    if not match:
        return None

    raw_limit = match.group(1).lower()
    if raw_limit.isdigit():
        return max(1, int(raw_limit))

    return INDONESIAN_SMALL_NUMBERS.get(raw_limit)


def _build_concise_revision_body(
    config: dict[str, str],
    *,
    fallback_body: str,
    paragraph_limit: int | None = None,
) -> str:
    source = config.get("content", "").strip()
    basis = config.get("basis", "").strip().rstrip(".")
    lead = basis if basis else "Sehubungan hal tersebut"

    items = _extract_numbered_items(source) or _extract_numbered_items(fallback_body)
    if items:
        if paragraph_limit is not None:
            return _format_numbered_items_as_concise_paragraph(lead, items)
        return f"{lead}, disampaikan hal-hal sebagai berikut:\n{_format_numbered_items(items)}"

    source_blocks = _split_blocks(source)
    if source_blocks:
        clean_source = " ".join(source_blocks[:2]).strip().rstrip(".")
        return f"{lead}, {clean_source}."

    fallback_blocks = _split_blocks(fallback_body)
    return " ".join(fallback_blocks[:2]).strip()


def _format_numbered_items_as_concise_paragraph(lead: str, items: list[str]) -> str:
    clean_items = [_normalize_concise_revision_item(item) for item in items if item.strip()]
    clean_items = [item for item in clean_items if item]
    if not clean_items:
        return lead
    if len(clean_items) == 1:
        summary = clean_items[0]
    else:
        summary = "; ".join(clean_items[:-1]) + f"; dan {clean_items[-1]}"
    return f"{lead}, hal-hal yang perlu ditindaklanjuti meliputi {summary}."


def _normalize_concise_revision_item(item: str) -> str:
    clean = item.rstrip(" .;")
    clean = re.sub(r"\s+", " ", clean).strip()
    clean = re.sub(r",?\s+sehingga\b.*$", "", clean, flags=re.IGNORECASE).strip(" ,;")

    replacements = (
        (r"^menyusun\s+rencana\s+kerja\b", "penyusunan rencana kerja"),
        (r"^menunjuk\s+pic\s+serta\s+menetapkan\s+jadwal\s+pelaksanaan\b", "penetapan PIC dan jadwal pelaksanaan"),
        (r"^menunjuk\s+pic\b", "penetapan PIC"),
        (r"^menetapkan\s+pic\b", "penetapan PIC"),
        (r"^mengidentifikasi\s+kendala\s+dan\s+mengajukan\s+usulan\s+perbaikan\b", "identifikasi kendala dan usulan perbaikan"),
        (r"^mengidentifikasi\s+kendala\b", "identifikasi kendala"),
        (r"^menyampaikan\s+laporan\s+lengkap\b", "penyampaian laporan"),
        (r"^menyampaikan\s+laporan\b", "penyampaian laporan"),
        (r"^melakukan\s+evaluasi\s+berkala\b", "evaluasi berkala"),
    )
    for pattern, replacement in replacements:
        if re.match(pattern, clean, flags=re.IGNORECASE):
            clean = re.sub(pattern, replacement, clean, count=1, flags=re.IGNORECASE)
            break
    else:
        clean = clean[:1].lower() + clean[1:] if clean else clean

    clean = re.sub(r"\bPIC\b", "PIC", clean, flags=re.IGNORECASE)
    return clean.rstrip(" .;")


def _preserve_configured_numbered_items(body: str, config: dict[str, str]) -> str:
    if config.get("revision_instruction"):
        return body

    configured_items = _extract_numbered_items(config.get("content", ""))
    if len(configured_items) < 2:
        return body

    body_blocks = _split_blocks(body)
    has_ordinal_rewrite = any(ORDINAL_ITEM_PATTERN.match(block) for block in body_blocks)
    body_numbered_count = len(_extract_numbered_items(body))
    if not has_ordinal_rewrite and body_numbered_count >= len(configured_items):
        return body

    intro_blocks: list[str] = []
    for block in body_blocks:
        if _is_structured_body_block(block) or ORDINAL_ITEM_PATTERN.match(block):
            break
        if _looks_like_closing_block(block):
            break
        intro_blocks.append(block)

    if not intro_blocks:
        basis = config.get("basis", "").strip().rstrip(".")
        lead = basis if basis else "Sehubungan hal tersebut"
        intro_blocks = [f"{lead}, dapat kami sampaikan sebagai berikut."]

    return "\n".join([*intro_blocks[:1], _format_numbered_items(configured_items)]).strip()


def _extract_numbered_items(text: str) -> list[str]:
    items: list[str] = []
    for block in _split_blocks(text):
        match = re.match(r"^\d+[.)]\s+(.+)$", block)
        if match:
            items.append(match.group(1).strip())
    return items


def _join_numbered_items(items: list[str]) -> str:
    fragments = [f"{index}. {item.rstrip(' .;')}" for index, item in enumerate(items, start=1)]
    if len(fragments) == 1:
        return f"{fragments[0]}."

    return f"{'; '.join(fragments[:-1])}; dan {fragments[-1]}."


def _format_numbered_items(items: list[str]) -> str:
    return "\n".join(
        f"{index}. {item.rstrip(' .;')}."
        for index, item in enumerate(items, start=1)
        if item.strip()
    )


def _paragraph_block_count(blocks: list[str]) -> int:
    return sum(1 for block in blocks if not _is_structured_body_block(block))


def _is_structured_body_block(block: str) -> bool:
    stripped = block.strip()
    if re.match(r"^\d+[.)]\s+.+$", stripped):
        return True
    if re.match(r"^[-*•]\s+.+$", stripped):
        return True
    return _parse_person_data_block(stripped) is not None


def _word_count(text: str) -> int:
    return len(re.findall(r"\b[\w/-]+\b", text or ""))


def _strip_forbidden_body_sections(text: str) -> str:
    text = re.sub(r"(?ims)(?:^|\n)\s*Tembusan\s*:.*$", "\n", text)
    lines = text.splitlines()
    output: list[str] = []
    skipping_tembusan = False

    for line in lines:
        stripped = line.strip()
        if skipping_tembusan:
            if not stripped:
                skipping_tembusan = False
            continue

        if re.match(r"^tembusan\s*:", stripped, re.IGNORECASE):
            skipping_tembusan = True
            continue

        if re.match(r"^(yth\.?|dari|hal|tanggal|nomor)\s*:", stripped, re.IGNORECASE):
            continue

        if re.match(r"^(memorandum|qr|tte)$", stripped, re.IGNORECASE):
            continue

        output.append(line)

    return "\n".join(output).strip()


def _strip_configured_carbon_copy_lines(text: str, carbon_copy: str) -> str:
    carbon_copy_keys = {
        _normalize_comparison_text(re.sub(r"^\s*\d+[.)]\s*", "", line))
        for line in _split_lines(carbon_copy)
    }
    carbon_copy_keys.discard("")

    if not carbon_copy_keys:
        return text

    output: list[str] = []
    for block in _split_blocks(text):
        clean_block = re.sub(r"^\s*\d+[.)]\s*", "", block).strip()
        if _normalize_comparison_text(clean_block) in carbon_copy_keys:
            continue
        output.append(block)

    return "\n".join(output).strip()


def _strip_external_source_artifacts(text: str) -> str:
    clean = re.sub(r"(?is)\[SOURCES\s*:.*?(?:\]\]|\n\s*\n|$)", "", text or "")
    clean = re.sub(r"\\n", " ", clean)

    output: list[str] = []
    for block in _split_blocks(clean):
        normalized = block.lower()
        has_source_marker = any(
            marker in normalized
            for marker in (
                '"type":"web"',
                '"type": "web"',
                '"url"',
                '"snippet"',
                "source:",
                "sources:",
            )
        )
        if has_source_marker and ("http://" in normalized or "https://" in normalized or "{" in block):
            continue
        if block.strip().startswith(("{", "[{")) and ("http://" in normalized or "https://" in normalized):
            continue
        output.append(block)

    return "\n".join(output).strip()


def _strip_markdown_artifacts(text: str) -> str:
    clean = re.sub(r"(?m)^\s{0,3}#{1,6}\s*", "", text)
    clean = re.sub(r"\*\*(.*?)\*\*", r"\1", clean)
    clean = re.sub(r"__(.*?)__", r"\1", clean)
    clean = clean.replace("```", "").replace("`", "")
    clean = re.sub(r"\[([^\]]+)\]\((?:https?://)?[^)]+\)", r"\1", clean)
    return clean.strip()


def _strip_evaluation_artifacts(text: str) -> str:
    output: list[str] = []

    for block in _split_blocks(text):
        if _is_evaluation_artifact_block(block):
            continue

        clean = re.sub(
            r",?\s*(?:serta|dan)?\s*sebagai baseline untuk uji[^,.;]*(?P<sep>[,.;])",
            lambda match: match.group("sep") if match.group("sep") == "," else ".",
            block,
            flags=re.IGNORECASE,
        )
        clean = re.sub(r"\bpada baseline\b", "pada", clean, flags=re.IGNORECASE)
        clean = re.sub(r"\bbaseline\b", "", clean, flags=re.IGNORECASE)
        clean = re.sub(r"\s{2,}", " ", clean).strip()
        clean = re.sub(r"\s+([,.;])", r"\1", clean)
        clean = re.sub(r",\s*,", ",", clean).strip(" ,")

        if clean:
            output.append(clean)

    return "\n".join(output).strip()


def _is_evaluation_artifact_block(block: str) -> bool:
    normalized = block.strip().lower()
    if not normalized:
        return False

    if "baseline paragraf ini digunakan" in normalized:
        return True

    evaluation_starters = (
        "poin nomor",
        "memo ini juga",
        "penyampaian data dan laporan ini digunakan",
    )
    if normalized.startswith(evaluation_starters) and "baseline" in normalized:
        return True

    if normalized.startswith("sebagai baseline"):
        return True

    if any(marker in normalized for marker in ("skenario evaluasi", "skenario revisi", "digenerate")):
        return True

    return "pengujian auto format" in normalized


def _strip_body_closing_sentences(text: str) -> str:
    output: list[str] = []

    for block in _split_blocks(text):
        if _looks_like_closing_block(block):
            output.append(block)
            continue

        trailing_closing = _split_trailing_closing_sentence(block)
        if trailing_closing:
            body_part, closing_part = trailing_closing
            if body_part:
                output.append(body_part)
            output.append(closing_part)
            continue

        clean = re.sub(r"(?is)\s+dengan\s+demikian\b.*$", "", block).strip()
        clean = re.sub(r"(?is)\s+demikian\b.*$", "", clean).strip()
        clean = re.sub(r"(?is)\s+atas\s+perhatian\b.*?terima\s+kasih\.?$", "", clean).strip()

        if clean and not _looks_like_closing_block(clean):
            output.append(clean)

    return "\n".join(output).strip()


def _split_trailing_closing_sentence(block: str) -> tuple[str, str] | None:
    match = re.search(
        r"(?is)(?P<body>.+?[.;])\s+"
        r"(?P<closing>(?:dengan\s+demikian|demikian|atas\s+perhatian|atas\s+kerja\s+sama|"
        r"dimohon|mohon\s+arahan\s+lebih\s+lanjut|mohon\s+tindak\s+lanjut|mohon\s+untuk\s+dapat|"
        r"kami\s+harapkan)\b.+)$",
        block.strip(),
    )
    if not match:
        return None

    body_part = match.group("body").strip()
    closing_part = match.group("closing").strip()
    if not body_part or not _looks_like_closing_block(closing_part):
        return None
    return body_part, closing_part


def _separate_generated_closing(body: str, config: dict[str, str]) -> tuple[str, str]:
    blocks = _split_blocks(body)
    if not blocks:
        return body, ""

    generated_closing = ""
    while blocks:
        if _looks_like_closing_block(blocks[-1]):
            closing_block = blocks.pop().strip()
            if not config.get("closing") and not generated_closing:
                generated_closing = closing_block
            continue

        trailing_closing = _split_trailing_closing_sentence(blocks[-1])
        if trailing_closing:
            body_part, closing_block = trailing_closing
            blocks[-1] = body_part
            if not config.get("closing") and not generated_closing:
                generated_closing = closing_block
            continue

        break

    return "\n".join(blocks).strip(), generated_closing


def _clean_official_language_fragments(text: str) -> str:
    output: list[str] = []

    replacements = (
        (r"\bkami ingin menyampaikan sebagai berikut\b", "dapat kami sampaikan sebagai berikut"),
        (r"\bkami ingin menyampaikan berikut\b", "dapat kami sampaikan sebagai berikut"),
        (r"\bberikut adalah beberapa\b", "disampaikan beberapa"),
        (r"\bTerimakasih\b", "Terima kasih"),
        (r"\bterimakasih\b", "terima kasih"),
    )

    for block in _split_blocks(text):
        clean = block
        clean = re.sub(r"(?i)\bmohon\s+untuk\s+mem\s+dan\b[^.;]*(?:[.;]|$)", "", clean)
        clean = re.sub(r"(?i)\buntuk\s+mem\s+dan\b[^.;]*(?:[.;]|$)", "", clean)
        for pattern, replacement in replacements:
            clean = re.sub(pattern, replacement, clean, flags=re.IGNORECASE)
        clean = _normalize_spacing_after_sentence_removal(clean)
        if clean:
            output.append(clean)

    return "\n".join(output).strip()


def _looks_like_closing_block(block: str) -> bool:
    normalized = _normalize_comparison_text(block)
    if not normalized:
        return False

    closing_starters = (
        "dengan demikian",
        "demikian",
        "atas perhatian",
        "atas kerja sama",
        "dimohon",
        "mohon arahan lebih lanjut",
        "mohon tindak lanjut",
        "mohon untuk dapat",
        "kami harapkan",
    )
    return normalized.startswith(closing_starters)


def _strip_trailing_fragments(text: str) -> str:
    blocks = _split_blocks(text)

    while blocks and _looks_like_trailing_fragment(blocks[-1]):
        blocks.pop()

    if blocks:
        blocks[-1] = re.sub(r"(?is)\s+(?:dengan|untuk\s+itu)$", "", blocks[-1]).strip(" ,;")
        if _looks_like_trailing_fragment(blocks[-1]):
            blocks.pop()

    return "\n".join(block for block in blocks if block).strip()


def _looks_like_trailing_fragment(block: str) -> bool:
    normalized = _normalize_comparison_text(block)
    return normalized in {
        "dengan",
        "dengan demikian",
        "untuk itu",
        "sehubungan hal tersebut",
    }


def _remove_configured_closing(text: str, closing: str) -> str:
    if not closing:
        return text

    closing_key = _normalize_comparison_text(closing)
    output = []
    for block in _split_blocks(text):
        clean = re.sub(re.escape(closing), "", block, flags=re.IGNORECASE).strip()
        if not clean:
            continue
        if _normalize_comparison_text(clean) == closing_key:
            continue
        output.append(clean)
    return "\n".join(output).strip()


def _strip_instruction_artifacts(text: str, config: dict[str, str]) -> str:
    clean = text
    for instruction in (config.get("revision_instruction", ""), config.get("additional_instruction", "")):
        clean = _remove_exact_instruction_text(clean, instruction)

    output: list[str] = []
    for block in _split_blocks(clean):
        if not _looks_like_instruction_artifact(block):
            output.append(block)
            continue

        sentences = _split_sentence_like_parts(block)
        kept = [sentence for sentence in sentences if not _looks_like_instruction_artifact(sentence)]
        cleaned_block = " ".join(sentence.strip() for sentence in kept if sentence.strip())
        cleaned_block = _normalize_spacing_after_sentence_removal(cleaned_block)
        if cleaned_block and not _looks_like_instruction_artifact(cleaned_block):
            output.append(cleaned_block)

    return "\n".join(output).strip()


def _strip_unconfigured_honorific_data(text: str, config: dict[str, str]) -> str:
    config_text = _normalize_comparison_text(
        "\n".join(
            str(config.get(key, ""))
            for key in ("basis", "content", "closing", "additional_instruction", "revision_instruction")
        )
    )
    if any(marker in config_text for marker in ("bapak ", "ibu ", "sdr.", "sdr ")):
        return text

    output: list[str] = []
    for block in _split_blocks(text):
        normalized = _normalize_comparison_text(block)
        has_unconfigured_honorific = re.search(r"\b(?:bapak|ibu|sdr\.?)\s+[a-z]", normalized) is not None
        if "pic" in normalized and "sebagai berikut" in normalized and not _config_contains_person_name(config_text):
            numbered = re.match(r"^(\d+[.)]\s+)", block)
            prefix = numbered.group(1) if numbered else ""
            output.append(
                f"{prefix}Penunjukan Person In Charge (PIC) pada tiap layanan agar ditetapkan oleh masing-masing unit."
            )
            continue
        if has_unconfigured_honorific and re.match(r"^[-*•]\s+", block):
            continue
        if has_unconfigured_honorific and "pic" in normalized:
            output.append("Penunjukan Person In Charge (PIC) pada tiap layanan agar ditetapkan oleh masing-masing unit.")
            continue
        output.append(block)

    return "\n".join(output).strip()


def _strip_unconfigured_operational_data(text: str, config: dict[str, str]) -> str:
    config_text = "\n".join(
        str(config.get(key, ""))
        for key in ("basis", "content", "closing", "additional_instruction", "revision_instruction")
    )
    normalized_config = _normalize_comparison_text(config_text)
    config_requests_incident_time = "waktu kejadian" in normalized_config
    config_requests_impacted_users = "pengguna terdampak" in normalized_config
    has_configured_time = _config_contains_explicit_time(config_text)

    if not (config_requests_incident_time or config_requests_impacted_users):
        return text

    output: list[str] = []
    for block in _split_blocks(text):
        normalized = _normalize_comparison_text(block)
        prefix_match = re.match(r"^(\d+[.)]\s+)", block)
        prefix = prefix_match.group(1) if prefix_match else ""

        if (
            config_requests_incident_time
            and not has_configured_time
            and re.search(r"\b(?:pukul|jam)\s+\d{1,2}[.:]\d{2}", normalized)
            and any(marker in normalized for marker in ("waktu kejadian", "kejadian", "kendala", "terjadi"))
        ):
            output.append(f"{prefix}Waktu kejadian ditetapkan berdasarkan laporan unit terkait.")
            continue

        if config_requests_impacted_users and _contains_unconfigured_impacted_users(normalized, normalized_config):
            output.append(f"{prefix}Pengguna terdampak diidentifikasi oleh unit terkait.")
            continue

        output.append(block)

    return "\n".join(output).strip()


def _config_contains_explicit_time(config_text: str) -> bool:
    return re.search(r"\b(?:pukul|jam)?\s*\d{1,2}[.:]\d{2}\s*(?:WIB)?\b", config_text, re.IGNORECASE) is not None


def _contains_unconfigured_impacted_users(normalized_block: str, normalized_config: str) -> bool:
    if "pengguna" not in normalized_block and "terdampak" not in normalized_block:
        return False
    if "seluruh staf persuratan" in normalized_block and "seluruh staf persuratan" not in normalized_config:
        return True
    if "seluruh pegawai" in normalized_block and "seluruh pegawai" not in normalized_config:
        return True
    if "seluruh staf" in normalized_block and "seluruh staf" not in normalized_config:
        return True
    return (
        re.search(r"\bpengguna(?:\s+yang)?\s+terdampak\b.*\bseluruh\b", normalized_block) is not None
        and "seluruh" not in normalized_config
    )


def _replace_empty_person_data_placeholders(text: str) -> str:
    blocks = _split_blocks(text)
    output: list[str] = []
    placeholder_labels: list[str] = []

    def flush_placeholders() -> None:
        nonlocal placeholder_labels
        if not placeholder_labels:
            return
        if len(placeholder_labels) >= 2:
            output.append(_generic_placeholder_replacement(placeholder_labels))
        placeholder_labels = []

    for block in blocks:
        label = _empty_person_data_placeholder_label(block)
        if label:
            placeholder_labels.append(label)
            continue

        flush_placeholders()
        output.append(block)

    flush_placeholders()
    return "\n".join(output).strip()


def _empty_person_data_placeholder_label(block: str) -> str | None:
    match = re.match(
        rf"^\s*(?:\d+[.)]\s*)?(?P<label>{PERSON_DATA_LABEL_PATTERN})\s*:\s*$",
        block.strip(),
        flags=re.IGNORECASE,
    )
    if not match:
        return None
    return _normalize_person_data_label(match.group("label"))


def _generic_placeholder_replacement(labels: list[str]) -> str:
    normalized = {_key_value_label_key(label) for label in labels}
    if normalized & {"nama", "nip", "jabatan", "unit kerja", "jadwal", "jadwal pendampingan"}:
        return "Staf pendamping dan kelengkapan data pegawai ditetapkan oleh unit terkait."
    if {"waktu kejadian", "pengguna terdampak"} <= normalized:
        return "Waktu kejadian dan pengguna terdampak ditetapkan berdasarkan laporan unit terkait."
    if "waktu kejadian" in normalized:
        return "Waktu kejadian ditetapkan berdasarkan laporan unit terkait."
    if "pengguna terdampak" in normalized:
        return "Pengguna terdampak diidentifikasi oleh unit terkait."
    if "peserta" in normalized:
        return "Daftar peserta ditetapkan oleh unit terkait sesuai kebutuhan kegiatan."
    if "nomor kontak" in normalized:
        return "Nomor kontak koordinasi ditetapkan oleh unit terkait."
    return "Data pendukung ditetapkan oleh unit terkait sesuai kebutuhan."


def _config_contains_person_name(config_text: str) -> bool:
    return bool(re.search(r"\bnama\s*:", config_text) or re.search(r"\b(?:bapak|ibu|sdr\.?)\s+[a-z]", config_text))


def _remove_exact_instruction_text(text: str, instruction: str) -> str:
    clean_instruction = (instruction or "").strip()
    if not clean_instruction:
        return text

    variants = {
        clean_instruction,
        clean_instruction.rstrip(". "),
        _normalize_comparison_text(clean_instruction),
    }
    clean = text
    for variant in variants:
        if not variant:
            continue
        clean = re.sub(re.escape(variant), "", clean, flags=re.IGNORECASE)

    return _normalize_spacing_after_sentence_removal(clean)


def _split_sentence_like_parts(block: str) -> list[str]:
    parts = re.findall(r"[^.!?]+(?:[.!?]+|$)", block)
    return parts if parts else [block]


def _looks_like_instruction_artifact(text: str) -> bool:
    normalized = _normalize_comparison_text(text)
    if not normalized:
        return False

    return any(re.search(pattern, normalized, flags=re.IGNORECASE) for pattern in INSTRUCTION_ARTIFACT_PATTERNS)


def _normalize_spacing_after_sentence_removal(text: str) -> str:
    clean = re.sub(r"\s{2,}", " ", text or "").strip()
    clean = re.sub(r"\s+([,.;:])", r"\1", clean)
    clean = re.sub(r"([,;:])\s*([.;])", r"\2", clean)
    clean = re.sub(r"\.{2,}", ".", clean)
    return clean.strip(" ,;")


def _dedupe_consecutive_blocks(text: str) -> str:
    output: list[str] = []
    previous_key = ""

    for block in _split_blocks(text):
        key = _normalize_comparison_text(block)
        if key == previous_key:
            continue
        output.append(block)
        previous_key = key

    return "\n".join(output).strip()


def _normalize_comparison_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower()).rstrip(".")


def _resolve_page_size(config: dict[str, str], body: str) -> str:
    requested = config.get("page_size", "folio")
    if requested in {"letter", "folio"}:
        return requested

    carbon_copy_lines = _split_lines(config.get("carbon_copy", ""))
    body_blocks = _split_blocks(body)
    measured_text = "\n".join(
        [
            config.get("basis", ""),
            body,
            config.get("closing", ""),
            "\n".join(carbon_copy_lines),
        ]
    ).strip()
    numbered_count = sum(1 for block in body_blocks if re.match(r"^\d+[.)]\s+.+$", block))
    effective_blocks = len(body_blocks) + len(carbon_copy_lines)
    estimated_lines = _estimate_rendered_lines(config, body)

    if len(measured_text) <= 880 and effective_blocks <= 10 and numbered_count <= 5:
        return "letter"

    if len(measured_text) <= 1120 and estimated_lines <= 10 and effective_blocks <= 12 and numbered_count <= 5:
        return "letter"

    return "folio"


def _should_use_compact_layout(config: dict[str, str], body: str) -> bool:
    if config.get("page_size") != "folio":
        return False

    body_blocks = _split_blocks(body)
    carbon_copy_lines = _split_lines(config.get("carbon_copy", ""))
    numbered_count = sum(1 for block in body_blocks if re.match(r"^\d+[.)]\s+.+$", block))
    return (
        len(body) >= 1500
        or len(body_blocks) + len(carbon_copy_lines) >= 8
        or numbered_count >= 5
        or len(carbon_copy_lines) >= 3
    )


def _signature_space_before(config: dict[str, str], body: str, *, compact: bool) -> float:
    if compact:
        body_length = len(body or "")
        carbon_copy_count = len(_split_lines(config.get("carbon_copy", "")))
        if config.get("page_size") == "folio" and carbon_copy_count < 3:
            if body_length <= 900:
                return 118
            if body_length <= 1400:
                return 100
        if body_length >= 2200 or carbon_copy_count >= 3:
            return SIGNATURE_SPACE_BEFORE_COMPACT_DENSE_PT
        return SIGNATURE_SPACE_BEFORE_COMPACT_PT

    body_length = len(body or "")
    carbon_copy_count = len(_split_lines(config.get("carbon_copy", "")))

    if body_length <= 700:
        space = SIGNATURE_SPACE_BEFORE_SHORT_PT
    elif body_length <= 1150:
        space = SIGNATURE_SPACE_BEFORE_MEDIUM_PT
    else:
        space = SIGNATURE_SPACE_BEFORE_DEFAULT_PT

    if carbon_copy_count >= 3:
        return max(SIGNATURE_SPACE_BEFORE_DEFAULT_PT, space - 32)
    if carbon_copy_count >= 2:
        return max(SIGNATURE_SPACE_BEFORE_DEFAULT_PT, space - 18)
    return space


def _estimate_rendered_lines(config: dict[str, str], body: str) -> int:
    blocks = _split_blocks(body)
    carbon_copy_lines = _split_lines(config.get("carbon_copy", ""))
    closing_lines = _split_blocks(config.get("closing", ""))
    basis_lines = _split_blocks(config.get("basis", ""))

    body_lines = sum(max(1, (len(block) + 91) // 92) for block in blocks)
    closing_unit = sum(max(1, (len(block) + 91) // 92) for block in closing_lines)

    return body_lines + len(carbon_copy_lines) + closing_unit + min(1, len(basis_lines))


def _body_alignment(block: str):
    return WD_ALIGN_PARAGRAPH.JUSTIFY if len(block) >= 180 else WD_ALIGN_PARAGRAPH.LEFT


def _collect_person_data_blocks(blocks: list[str], start_index: int) -> list[tuple[str, str]]:
    items: list[tuple[str, str]] = []

    for block in blocks[start_index:]:
        parsed = _parse_person_data_block(block)
        if not parsed:
            break
        items.append(parsed)

    return items


def _parse_person_data_block(block: str) -> tuple[str, str] | None:
    match = PERSON_DATA_PATTERN.match(block.strip())
    if not match:
        return None

    raw_label = match.group("label")
    label = _normalize_person_data_label(raw_label)
    value = _clean_key_value_value(label, match.group("value"))

    return label, value


def _normalize_person_data_label(raw_label: str) -> str:
    normalized = raw_label.lower().rstrip(".")
    normalized = re.sub(r"\s+", " ", normalized)

    if normalized.startswith("nama"):
        return "nama"
    if normalized == "nip":
        return "NIP"
    if normalized.startswith("pangkat/gol"):
        return "pangkat/gol."
    if normalized == "jabatan":
        return "jabatan"
    if normalized.startswith("unit kerja"):
        return "unit kerja"
    if normalized.startswith("jadwal"):
        return "jadwal pendampingan" if "pendampingan" in normalized else "jadwal"
    if normalized.startswith("hari/tanggal"):
        return "hari/tanggal"
    if normalized.startswith("pukul"):
        return "pukul"
    if normalized.startswith("tempat"):
        return "tempat"
    if normalized.startswith("agenda"):
        return "agenda"
    if normalized.startswith("estimasi durasi"):
        return "estimasi durasi"
    if normalized.startswith("durasi"):
        return "durasi"
    if normalized.startswith("batas waktu"):
        return "batas waktu"
    if normalized.startswith("lokasi"):
        return normalized
    if normalized.startswith("periode"):
        return "periode"
    if normalized.startswith("waktu pelaksanaan"):
        return "waktu pelaksanaan"
    if normalized.startswith("waktu kejadian"):
        return "waktu kejadian"
    if normalized.startswith("pengguna terdampak"):
        return "pengguna terdampak"
    if normalized == "peserta":
        return "peserta"
    if normalized.startswith("nomor kontak"):
        return "nomor kontak"
    return "keperluan"


def _normalize_person_data_sequences(text: str) -> str:
    label = PERSON_DATA_LABEL_PATTERN
    clean = re.sub(
        rf"(?i)(?<!^)(?<!\n)\s+(\d+[.)]\s*(?:{label})\s*(?::|adalah))",
        r"\n\1",
        text,
    ).strip()
    return _split_inline_person_data_sequences(clean)


def _split_inline_person_data_sequences(text: str) -> str:
    pattern = re.compile(
        rf"(?i)(?:^|[\s,;])(?:\d+[.)]\s*)?(?P<label>{PERSON_DATA_LABEL_PATTERN})\s*(?::|adalah)\s*"
    )
    output: list[str] = []

    for block in _split_blocks(text):
        matches = list(pattern.finditer(block))
        if len(matches) < 2:
            output.append(block)
            continue

        prefix = block[: matches[0].start()].strip(" ,;")
        if prefix:
            output.append(prefix)

        for index, match in enumerate(matches):
            next_start = matches[index + 1].start() if index + 1 < len(matches) else len(block)
            raw_value = block[match.end() : next_start]
            label_text = _normalize_person_data_label(match.group("label"))
            value = _clean_key_value_value(label_text, raw_value)
            if not value:
                continue
            output.append(f"{label_text}: {value}")

    return "\n".join(output).strip()


def _clean_inline_person_data_value(value: str) -> str:
    clean = re.sub(r"\s+", " ", value or "").strip(" ,;")
    clean = re.sub(r"^\d+[.)]\s+", "", clean).strip()
    return clean.rstrip(" ,;.")


def _clean_title(title: str) -> str:
    clean = re.sub(r"\s+", " ", (title or "").strip())
    if not clean:
        raise ValueError("Judul memo wajib diisi.")
    if len(clean) > 160:
        raise ValueError("Judul memo terlalu panjang.")
    return clean


def _normalize_generated_text(text: str) -> str:
    clean = (text or "").strip()
    if "[MODEL:" in clean and "]" in clean:
        clean = clean.split("]", 1)[1].strip()
    if _looks_like_ai_unavailable_message(clean):
        raise ValueError("Layanan AI sedang tidak tersedia. Silakan coba lagi nanti.")
    if not clean:
        raise ValueError("AI tidak menghasilkan isi memo.")
    return clean


def _looks_like_ai_unavailable_message(text: str) -> bool:
    normalized = _normalize_comparison_text(text).lstrip("❌ ").strip()
    return any(marker in normalized for marker in AI_UNAVAILABLE_MARKERS)


def _split_blocks(text: str) -> list[str]:
    blocks = [re.sub(r"\s+", " ", block.strip()) for block in re.split(r"\n{1,}", text)]
    return [block for block in blocks if block]


def _split_lines(text: str) -> list[str]:
    return [line.strip() for line in (text or "").splitlines() if line.strip()]


def _slugify(value: str) -> str:
    slug = re.sub(r"[^\w.\-]+", "-", value.lower(), flags=re.UNICODE)
    slug = re.sub(r"-{2,}", "-", slug).strip("._-")
    return slug or "memo"
