import os
import csv
import hashlib
import logging
import time
from pathlib import Path
from typing import List, Tuple

from langchain_chroma import Chroma
from langchain_core.documents import Document
import pdfplumber
from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.services.rag_config import (
    CHROMA_PATH,
    VECTOR_COLLECTION_NAME,
    PARENT_COLLECTION_NAME,
    TOKEN_CHUNK_SIZE,
    TOKEN_CHUNK_OVERLAP,
    AGGRESSIVE_BATCH_SIZE,
    BATCH_DELAY_SECONDS,
    MAX_TOKENS_PER_BATCH,
    MAX_EMBEDDING_DIM,
    EMBEDDING_MODELS,
)
from app.services.rag_embeddings import count_tokens, get_embeddings_with_fallback
from app.services.lightweight_text_splitter import LightweightRecursiveTextSplitter

logger = logging.getLogger(__name__)


def _clean_text(value) -> str:
    return "" if value is None else str(value).strip()


def _table_to_text(rows: list[list[object]]) -> str:
    lines: list[str] = []
    for row in rows:
        cleaned = [_clean_text(cell) for cell in row]
        if any(cleaned):
            lines.append(" | ".join(cleaned))
    return "\n".join(lines)


def _load_pdf_documents(file_path: str, filename: str) -> list[Document]:
    docs: list[Document] = []

    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            parts: list[str] = []
            text = (page.extract_text() or "").strip()
            if text:
                parts.append(text)

            for table in page.extract_tables() or []:
                table_text = _table_to_text(table)
                if table_text:
                    parts.append(table_text)

            content = "\n\n".join(parts).strip()
            if content:
                docs.append(
                    Document(
                        page_content=content,
                        metadata={"source": filename, "page": page_number},
                    )
                )

    return docs


def _load_docx_documents(file_path: str, filename: str) -> list[Document]:
    document = DocxDocument(file_path)
    parts: list[str] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        table_text = _table_to_text(
            [[cell.text for cell in row.cells] for row in table.rows]
        )
        if table_text:
            parts.append(table_text)

    content = "\n\n".join(parts).strip()
    return [Document(page_content=content, metadata={"source": filename})] if content else []


def _load_xlsx_documents(file_path: str, filename: str) -> list[Document]:
    workbook = load_workbook(file_path, data_only=True, read_only=True)
    docs: list[Document] = []

    try:
        for sheet in workbook.worksheets:
            rows = [
                [cell for cell in row]
                for row in sheet.iter_rows(values_only=True)
            ]
            table_text = _table_to_text(rows)
            if table_text:
                docs.append(
                    Document(
                        page_content=f"{sheet.title}\n\n{table_text}",
                        metadata={"source": filename, "sheet": sheet.title},
                    )
                )
    finally:
        workbook.close()

    return docs


def _load_csv_documents(file_path: str, filename: str) -> list[Document]:
    rows: list[list[str]] = []

    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            with open(file_path, newline="", encoding=encoding) as handle:
                sample = handle.read(4096)
                handle.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample)
                except csv.Error:
                    dialect = csv.excel
                rows = [row for row in csv.reader(handle, dialect)]
            break
        except UnicodeDecodeError:
            continue

    content = _table_to_text(rows)
    return [Document(page_content=content, metadata={"source": filename})] if content else []


def _load_text_documents(file_path: str, filename: str) -> list[Document]:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            content = Path(file_path).read_text(encoding=encoding).strip()
            if content:
                return [Document(page_content=content, metadata={"source": filename})]
            return []
        except UnicodeDecodeError:
            continue

    return []


def _load_documents_lightweight(file_path: str, filename: str) -> list[Document]:
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _load_pdf_documents(file_path, filename)
    if suffix == ".docx":
        return _load_docx_documents(file_path, filename)
    if suffix == ".xlsx":
        return _load_xlsx_documents(file_path, filename)
    if suffix == ".csv":
        return _load_csv_documents(file_path, filename)
    if suffix in {".txt", ".md", ".markdown"}:
        return _load_text_documents(file_path, filename)

    raise ValueError(f"Unsupported document type for lightweight ingest: {suffix or 'unknown'}")


def process_document(file_path: str, filename: str, user_id: str = "unknown"):
    try:
        logger.info(f"=== Processing document: {filename} ===")
        logger.info(f"File path: {file_path}")
        logger.info(f"File exists: {os.path.exists(file_path)}")
        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            logger.info(f"File size: {file_size:,} bytes ({file_size / 1024 / 1024:.2f} MB)")

        import time as _time
        _load_start = _time.time()
        logger.info("Step 1: Loading document dengan parser ringan lokal...")
        docs = _load_documents_lightweight(file_path, filename)
        elapsed = _time.time() - _load_start
        logger.info(f"   ✅ Lightweight loader selesai: {len(docs)} bagian dalam {elapsed:.1f}s")

        if not docs:
            raise ValueError(
                "Tidak ada teks yang dapat diekstrak dari dokumen ini. "
                "PDF scan/gambar memerlukan pipeline OCR terpisah."
            )

        logger.info(f"Loaded {len(docs)} document(s)")

        total_content = "".join([doc.page_content for doc in docs])
        total_chars = len(total_content)
        estimated_tokens = count_tokens(total_content)
        logger.info(f"Total content: {total_chars:,} chars, ~{estimated_tokens:,} tokens")

        try:
            from app.config_loader import get_pdr_config
            pdr_cfg = get_pdr_config()
        except Exception:
            pdr_cfg = {}
        pdr_enabled       = pdr_cfg.get('enabled', False)
        pdr_child_size    = int(pdr_cfg.get('child_chunk_size', 256))
        pdr_child_overlap = int(pdr_cfg.get('child_chunk_overlap', 32))
        pdr_parent_size   = int(pdr_cfg.get('parent_chunk_size', TOKEN_CHUNK_SIZE))
        pdr_parent_overlap= int(pdr_cfg.get('parent_chunk_overlap', TOKEN_CHUNK_OVERLAP))

        if pdr_enabled:
            logger.info(
                "Step 2: PDR Chunking (parent=%d token, child=%d token)...",
                pdr_parent_size, pdr_child_size,
            )
            parent_splitter = LightweightRecursiveTextSplitter(
                chunk_size=pdr_parent_size,
                chunk_overlap=pdr_parent_overlap,
                length_function=count_tokens,
                add_start_index=True,
                separators=["\n\n", "\n", ". ", " ", ""],
            )
            parent_chunks = parent_splitter.split_documents(docs)

            child_splitter = LightweightRecursiveTextSplitter(
                chunk_size=pdr_child_size,
                chunk_overlap=pdr_child_overlap,
                length_function=count_tokens,
                add_start_index=True,
                separators=["\n\n", "\n", ". ", " ", ""],
            )

            child_chunks = []
            pdr_parent_docs = []

            for p_idx, parent in enumerate(parent_chunks):
                raw_key = f"{filename}:{user_id}:{p_idx}:{parent.page_content[:50]}"
                parent_id = hashlib.md5(raw_key.encode()).hexdigest()

                parent_meta = {
                    "filename": filename,
                    "user_id": str(user_id),
                    "chunk_type": "parent",
                    "parent_id": parent_id,
                    "parent_index": p_idx,
                }
                pdr_parent_docs.append((parent_id, parent.page_content, parent_meta))

                children = child_splitter.split_documents([parent])
                for c_idx, child in enumerate(children):
                    child.metadata["chunk_type"] = "child"
                    child.metadata["parent_id"]  = parent_id
                    child.metadata["child_index"] = c_idx
                    child_chunks.append(child)

            chunks = child_chunks
            logger.info(
                "✅ PDR: %d parent chunks + %d child chunks (ratio %.1f:1)",
                len(parent_chunks), len(child_chunks),
                len(child_chunks) / max(len(parent_chunks), 1),
            )

        else:
            logger.info(
                "Step 2: Token-Aware Recursive Chunking (chunk_size=%d, overlap=%d)...",
                TOKEN_CHUNK_SIZE, TOKEN_CHUNK_OVERLAP,
            )
            text_splitter = LightweightRecursiveTextSplitter(
                chunk_size=TOKEN_CHUNK_SIZE,
                chunk_overlap=TOKEN_CHUNK_OVERLAP,
                length_function=count_tokens,
                add_start_index=True,
                separators=["\n\n", "\n", ". ", " ", ""],
            )
            chunks = text_splitter.split_documents(docs)
            pdr_parent_docs = []
            logger.info(f"Created {len(chunks)} token-aware chunks")

        if chunks:
            chunk_tokens = [count_tokens(chunk.page_content) for chunk in chunks]
            avg_tokens = sum(chunk_tokens) / len(chunk_tokens)
            max_tokens_val = max(chunk_tokens)
            min_tokens_val = min(chunk_tokens)
            logger.info(
                "Chunk stats: avg=%d tokens, min=%d, max=%d",
                avg_tokens, min_tokens_val, max_tokens_val,
            )

        logger.info("Step 3: Initializing embedding model dengan cascading fallback...")
        current_model_index = 0
        embeddings, provider_name, current_model_index = get_embeddings_with_fallback(current_model_index)

        if embeddings is None:
            raise Exception("Semua embedding provider gagal. Tidak dapat memproses dokumen.")

        current_embedding_dim = pdr_cfg.get('embedding_dim', MAX_EMBEDDING_DIM)

        for idx, chunk in enumerate(chunks):
            chunk.metadata["filename"] = filename
            chunk.metadata["user_id"]  = str(user_id)
            chunk.metadata["embedding_model"] = provider_name
            if idx == 0:
                logger.info("🔍 INGEST: Storing chunk metadata - filename='%s', user_id='%s'", filename, str(user_id))

        logger.info(f"Step 4: Smart Batching & Embedding Generation...")
        logger.info(
            "Max batch size: %d chunks OR %d tokens (whichever is smaller)",
            AGGRESSIVE_BATCH_SIZE, MAX_TOKENS_PER_BATCH,
        )
        logger.info("Total capacity: 2M TPM across 4 models (4 x 500K TPM)")

        vectorstore = Chroma(
            collection_name=VECTOR_COLLECTION_NAME,
            embedding_function=embeddings,
            persist_directory=CHROMA_PATH
        )

        if pdr_enabled and pdr_parent_docs:
            try:
                parent_store = Chroma(
                    collection_name=PARENT_COLLECTION_NAME,
                    persist_directory=CHROMA_PATH,
                )
                raw_col = parent_store._collection
                p_ids   = [pid for pid, _, _ in pdr_parent_docs]
                p_texts = [txt for _, txt, _ in pdr_parent_docs]
                p_metas = [meta for _, _, meta in pdr_parent_docs]
                raw_col.upsert(
                    ids=p_ids,
                    documents=p_texts,
                    metadatas=p_metas,
                    embeddings=[[0.0] * MAX_EMBEDDING_DIM] * len(p_ids),
                )
                logger.info("✅ PDR: %d parent chunks disimpan ke ChromaDB", len(pdr_parent_docs))
            except Exception as pe:
                logger.warning("⚠️  Gagal menyimpan parent PDR: %s", pe)

        successful_chunks = 0
        failed_chunks = 0

        smart_batches = []
        current_batch = []
        current_batch_tokens = 0

        for chunk in chunks:
            chunk_tokens = count_tokens(chunk.page_content)

            would_exceed_tokens = (current_batch_tokens + chunk_tokens) > MAX_TOKENS_PER_BATCH
            would_exceed_count = len(current_batch) >= AGGRESSIVE_BATCH_SIZE

            if (would_exceed_tokens or would_exceed_count) and current_batch:
                smart_batches.append((current_batch, current_batch_tokens))
                current_batch = [chunk]
                current_batch_tokens = chunk_tokens
            else:
                current_batch.append(chunk)
                current_batch_tokens += chunk_tokens

        if current_batch:
            smart_batches.append((current_batch, current_batch_tokens))

        total_batches = len(smart_batches)
        logger.info(f"Created {total_batches} smart batches (token-aware)")

        for batch_index, (batch, batch_tokens) in enumerate(smart_batches, 1):
            try:
                if batch_index > 1:
                    time.sleep(BATCH_DELAY_SECONDS)

                logger.info(f"Processing batch {batch_index}/{total_batches}: {len(batch)} chunks, {batch_tokens:,} tokens...")
                vectorstore.add_documents(batch)
                successful_chunks += len(batch)
                logger.info(f"✅ Batch {batch_index}/{total_batches} success | Progress: {successful_chunks}/{len(chunks)} chunks")

            except Exception as batch_error:
                error_msg = str(batch_error)
                logger.error(f"❌ Batch {batch_index} error: {error_msg}")

                is_rate_limit = any(indicator in error_msg.lower() for indicator in [
                    "429", "rate limit", "resource_exhausted", "quota", "503", "too many requests"
                ])
                is_token_limit = any(indicator in error_msg.lower() for indicator in [
                    "413", "tokens_limit_reached", "too large", "request body too large"
                ])

                if is_token_limit and len(batch) > 1:
                    mid = len(batch) // 2
                    sub_batches = [batch[:mid], batch[mid:]]
                    logger.warning(
                        f"⚠️  Batch {batch_index} terlalu besar ({batch_tokens:,} tokens tiktoken). "
                        f"Auto-split → {len(sub_batches[0])} + {len(sub_batches[1])} chunks."
                    )
                    for si, sub in enumerate(sub_batches, 1):
                        try:
                            time.sleep(BATCH_DELAY_SECONDS)
                            vectorstore.add_documents(sub)
                            successful_chunks += len(sub)
                            logger.info(f"   ✅ Sub-batch {si}/2 berhasil: {len(sub)} chunks")
                        except Exception as sub_err:
                            logger.error(f"   ❌ Sub-batch {si}/2 gagal: {sub_err}")
                            failed_chunks += len(sub)

                elif is_rate_limit and current_model_index < len(EMBEDDING_MODELS) - 1:
                    if current_model_index < len(EMBEDDING_MODELS) - 1:
                        logger.warning(f"🚫 Rate limit detected! Cascading to next model tier...")

                        current_model_index += 1
                        embeddings, provider_name, current_model_index = get_embeddings_with_fallback(current_model_index)

                        if embeddings is None:
                            failed_chunks += len(batch)
                            logger.error(f"❌ All 4 models exhausted! Failed after {successful_chunks} chunks")
                            raise Exception(f"Semua 4 embedding models gagal (2M TPM exhausted) setelah {successful_chunks} chunks berhasil.")

                        vectorstore = Chroma(
                            collection_name=VECTOR_COLLECTION_NAME,
                            embedding_function=embeddings,
                            persist_directory=CHROMA_PATH
                        )

                        remaining_start = sum(len(b[0]) for b in smart_batches[:batch_index-1])
                        for remaining_chunk in chunks[remaining_start:]:
                            remaining_chunk.metadata["embedding_model"] = provider_name

                        retry_delay = 2.0
                        max_retries = 3

                        for retry in range(max_retries):
                            try:
                                logger.info(f"🔄 Retry {retry + 1}/{max_retries} dengan {provider_name}...")
                                time.sleep(retry_delay)
                                vectorstore.add_documents(batch)
                                successful_chunks += len(batch)
                                logger.info(f"✅ Batch {batch_index} berhasil dengan {provider_name} (retry {retry + 1})")
                                break
                            except Exception as retry_error:
                                retry_error_msg = str(retry_error)
                                logger.warning(f"⚠️ Retry {retry + 1} failed: {retry_error_msg}")

                                is_token_limit_retry = any(indicator in retry_error_msg.lower() for indicator in ["413", "tokens_limit_reached", "too large"])

                                if is_token_limit_retry:
                                    logger.error(f"❌ Batch {batch_index} exceeds token limit even after cascade")
                                    failed_chunks += len(batch)
                                    break

                                is_rate_limit_retry = any(indicator in retry_error_msg.lower() for indicator in ["429", "rate limit", "resource_exhausted"])
                                if is_rate_limit_retry:
                                    if retry < max_retries - 1:
                                        retry_delay *= 2
                                        continue
                                    else:
                                        from app.services.rag_config import EMBEDDING_MODELS
                                        if current_model_index < len(EMBEDDING_MODELS) - 1:
                                            current_model_index += 1
                                            embeddings, provider_name, current_model_index = get_embeddings_with_fallback(current_model_index)
                                            if embeddings:
                                                vectorstore = Chroma(
                                                    collection_name=VECTOR_COLLECTION_NAME,
                                                    embedding_function=embeddings,
                                                    persist_directory=CHROMA_PATH
                                                )
                                                remaining_start = sum(len(b[0]) for b in smart_batches[:batch_index-1])
                                                for remaining_chunk in chunks[remaining_start:]:
                                                    remaining_chunk.metadata["embedding_model"] = provider_name
                                                continue

                                if retry == max_retries - 1:
                                    logger.error(f"❌ Batch {batch_index} gagal setelah {max_retries} retries")
                                    failed_chunks += len(batch)
                                break
                else:
                    is_token_limit = any(indicator in error_msg.lower() for indicator in ["413", "tokens_limit_reached", "too large", "body too large"])

                    if is_token_limit:
                        logger.error(f"❌ Batch {batch_index} exceeds token limit ({batch_tokens:,} tokens)")
                        logger.error(f"💡 Suggestion: Reduce TOKEN_CHUNK_SIZE or AGGRESSIVE_BATCH_SIZE in .env")
                        failed_chunks += len(batch)
                    else:
                        logger.error(f"❌ Batch {batch_index} gagal (non-rate-limit atau no fallback)")
                        failed_chunks += len(batch)

        success_rate = (successful_chunks / len(chunks) * 100) if len(chunks) > 0 else 0
        logger.info(f"{'='*60}")
        logger.info(f"✅ Document '{filename}' processing completed")
        logger.info(f"Success: {successful_chunks}/{len(chunks)} chunks ({success_rate:.1f}%)")
        logger.info(f"Failed: {failed_chunks} chunks")
        logger.info(f"Final embedding model: {provider_name}")
        logger.info(f"Total tokens processed: ~{estimated_tokens:,}")
        logger.info(f"{'='*60}")

        if failed_chunks > 0:
            return True, f"Document processed dengan {failed_chunks}/{len(chunks)} chunks gagal"

        return True, "Document processed successfully dengan Token-Aware Chunking & Aggressive Batching."

    except Exception as e:
        logger.error(f"❌ Error processing document '{filename}': {type(e).__name__}: {str(e)}")
        return False, str(e)


def delete_document_vectors(filename: str):
    try:
        from app.services.rag_config import (
            CHROMA_PATH,
            VECTOR_COLLECTION_NAME,
            PARENT_COLLECTION_NAME,
        )
        embeddings, provider_name, _ = get_embeddings_with_fallback()

        if embeddings is None:
            return False, "Tidak dapat menginisialisasi embedding model untuk delete operation."

        vectorstore = Chroma(
            collection_name=VECTOR_COLLECTION_NAME,
            embedding_function=embeddings,
            persist_directory=CHROMA_PATH
        )

        vectorstore.delete(where={"filename": filename})

        try:
            parent_store = Chroma(
                collection_name=PARENT_COLLECTION_NAME,
                persist_directory=CHROMA_PATH,
            )
            raw_col = parent_store._collection
            raw_col.delete(where={"$and": [{"filename": filename}, {"chunk_type": "parent"}]})
            logger.info("✅ PDR parent chunks for %s deleted", filename)
        except Exception as pe:
            logger.debug("PDR parent delete skipped (mungkin non-PDR dokumen): %s", pe)

        logger.info("✅ Vectors for %s deleted successfully using %s", filename, provider_name)
        return True, f"Vectors for {filename} deleted successfully."
    except Exception as e:
        logger.error(f"❌ Error deleting vectors for {filename}: {str(e)}")
        return False, str(e)
