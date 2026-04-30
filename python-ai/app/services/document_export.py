from __future__ import annotations

import csv
import html
import re
from dataclasses import dataclass
from io import BytesIO, StringIO
from typing import Any

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document as DocxDocument
from openpyxl import Workbook
from weasyprint import HTML as WeasyHTML


PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
CSV_MIME = "text/csv; charset=utf-8"


@dataclass(slots=True)
class ExportArtifact:
    filename: str
    mime_type: str
    content: bytes


def _sanitize_filename(filename: str | None, fallback: str = "ista-ai-export") -> str:
    base = (filename or "").strip()
    if not base:
        base = fallback

    base = re.sub(r"[^\w.\-]+", "-", base, flags=re.UNICODE)
    base = re.sub(r"-{2,}", "-", base).strip("._-")

    return base or fallback


def _normalize_target_format(target_format: str) -> str:
    target = (target_format or "").strip().lower()
    if target not in {"pdf", "docx", "xlsx", "csv"}:
        raise ValueError(f"Unsupported export format: {target_format}")
    return target


def _wrap_html(content_html: str, title: str) -> str:
    safe_title = html.escape(title)
    return f"""<!DOCTYPE html>
<html lang="id">
  <head>
    <meta charset="utf-8" />
    <meta name="viewport" content="width=device-width, initial-scale=1" />
    <title>{safe_title}</title>
    <style>
      :root {{
        color-scheme: light;
      }}
      body {{
        font-family: Inter, Arial, Helvetica, sans-serif;
        color: #0f172a;
        line-height: 1.6;
        margin: 32px;
      }}
      h1, h2, h3, h4, h5, h6 {{
        line-height: 1.25;
        margin: 1.1em 0 0.45em;
      }}
      p {{
        margin: 0 0 0.8em;
      }}
      table {{
        border-collapse: collapse;
        width: 100%;
        margin: 1em 0;
      }}
      th, td {{
        border: 1px solid #cbd5e1;
        padding: 8px 10px;
        vertical-align: top;
        text-align: left;
      }}
      th {{
        background: #f8fafc;
        font-weight: 700;
      }}
      blockquote {{
        border-left: 4px solid #cbd5e1;
        margin: 1em 0;
        padding-left: 14px;
        color: #334155;
      }}
      pre {{
        background: #f8fafc;
        border: 1px solid #e2e8f0;
        border-radius: 8px;
        padding: 12px 14px;
        overflow: auto;
        white-space: pre-wrap;
      }}
      code {{
        font-family: "SFMono-Regular", Consolas, Monaco, monospace;
        font-size: 0.95em;
      }}
      ul, ol {{
        margin: 0 0 0.8em 1.5em;
        padding: 0;
      }}
    </style>
  </head>
  <body>
    <h1>{safe_title}</h1>
    {content_html}
  </body>
</html>"""


def _extract_text_lines(content_html: str) -> list[str]:
    soup = BeautifulSoup(content_html or "", "html.parser")
    text = soup.get_text("\n", strip=True)
    lines = [line.strip() for line in text.splitlines()]
    return [line for line in lines if line]


def _extract_tables_from_html(content_html: str) -> list[dict[str, Any]]:
    soup = BeautifulSoup(content_html or "", "html.parser")
    tables: list[dict[str, Any]] = []

    for index, table in enumerate(soup.find_all("table"), start=1):
        rows: list[list[str]] = []

        for row in table.find_all("tr"):
            cells = [
                re.sub(r"\s+", " ", cell.get_text(" ", strip=True)).strip()
                for cell in row.find_all(["th", "td"])
            ]
            if any(cell for cell in cells):
                rows.append(cells)

        if not rows:
            continue

        header = rows[0] if table.find("th") else []
        body_rows = rows[1:] if header else rows
        tables.append({
            "name": f"Table {index}",
            "header": header,
            "rows": body_rows,
        })

    return tables


def _render_docx(content_html: str, title: str) -> bytes:
    document = DocxDocument()
    document.core_properties.title = title

    soup = BeautifulSoup(content_html or "", "html.parser")
    root = soup.body or soup

    document.add_heading(title, level=1)

    for node in root.children:
        if isinstance(node, NavigableString):
            text = _clean_inline_text(str(node))
            if text:
                document.add_paragraph(text)
            continue

        if not isinstance(node, Tag):
            continue

        tag_name = node.name.lower()
        text = _clean_inline_text(node.get_text(" ", strip=True))

        if tag_name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            level = min(max(int(tag_name[1]), 1), 9)
            if text:
                document.add_heading(text, level=level)
            continue

        if tag_name in {"div", "section", "article"}:
            _append_nested_docx_content(document, node)
            continue

        if tag_name in {"p", "blockquote"}:
            if text:
                style = "Intense Quote" if tag_name == "blockquote" else None
                document.add_paragraph(text, style=style)
            continue

        if tag_name in {"ul", "ol"}:
            style = "List Bullet" if tag_name == "ul" else "List Number"
            for li in node.find_all("li", recursive=False):
                li_text = _clean_inline_text(li.get_text(" ", strip=True))
                if li_text:
                    document.add_paragraph(li_text, style=style)
            continue

        if tag_name == "table":
            rows = [
                [
                    _clean_inline_text(cell.get_text(" ", strip=True))
                    for cell in row.find_all(["th", "td"])
                ]
                for row in node.find_all("tr")
            ]
            rows = [row for row in rows if any(cell for cell in row)]
            if not rows:
                continue

            table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for col_index, value in enumerate(row):
                    table.cell(row_index, col_index).text = value
            continue

        if tag_name == "pre":
            if text:
                document.add_paragraph(text)
            continue

        if text:
            document.add_paragraph(text)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _append_nested_docx_content(document: DocxDocument, node: Tag) -> None:
    for child in node.children:
        if isinstance(child, NavigableString):
            text = _clean_inline_text(str(child))
            if text:
                document.add_paragraph(text)
            continue

        if not isinstance(child, Tag):
            continue

        tag_name = child.name.lower()
        text = _clean_inline_text(child.get_text(" ", strip=True))
        if not text:
            continue

        if tag_name in {"div", "section", "article"}:
            _append_nested_docx_content(document, child)
            continue

        if tag_name in {"ul", "ol"}:
            style = "List Bullet" if tag_name == "ul" else "List Number"
            for li in child.find_all("li", recursive=False):
                li_text = _clean_inline_text(li.get_text(" ", strip=True))
                if li_text:
                    document.add_paragraph(li_text, style=style)
            continue

        if tag_name == "table":
            rows = [
                [
                    _clean_inline_text(cell.get_text(" ", strip=True))
                    for cell in row.find_all(["th", "td"])
                ]
                for row in child.find_all("tr")
            ]
            rows = [row for row in rows if any(cell for cell in row)]
            if not rows:
                continue

            table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for col_index, value in enumerate(row):
                    table.cell(row_index, col_index).text = value
            continue

        if tag_name in {"p", "blockquote"}:
            document.add_paragraph(text)
            continue

        document.add_paragraph(text)


def _clean_inline_text(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def _render_xlsx(tables: list[dict[str, Any]], content_html: str, title: str) -> bytes:
    workbook = Workbook()
    default_sheet = workbook.active
    workbook.remove(default_sheet)

    if not tables:
        sheet = workbook.create_sheet(title="Jawaban")
        lines = _extract_text_lines(content_html)
        if not lines:
            sheet.cell(row=1, column=1, value="")
        else:
            for row_index, line in enumerate(lines, start=1):
                sheet.cell(row=row_index, column=1, value=line)
    else:
        for index, table_data in enumerate(tables, start=1):
            sheet_name = _sanitize_sheet_name(table_data.get("name") or f"Table {index}")
            sheet = workbook.create_sheet(title=sheet_name)
            row_cursor = 1

            header = [str(cell) for cell in table_data.get("header", []) if cell is not None]
            if header:
                for col_index, value in enumerate(header, start=1):
                    sheet.cell(row=row_cursor, column=col_index, value=value)
                row_cursor += 1

            for row in table_data.get("rows", []):
                for col_index, value in enumerate(row, start=1):
                    sheet.cell(row=row_cursor, column=col_index, value=value)
                row_cursor += 1

            if row_cursor == 1:
                sheet.cell(row=1, column=1, value=title)

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _render_csv(tables: list[dict[str, Any]], content_html: str) -> bytes:
    buffer = StringIO()
    writer = csv.writer(buffer)

    if not tables:
        lines = _extract_text_lines(content_html)
        if not lines:
            writer.writerow([""])
        else:
            for line in lines:
                writer.writerow([line])
    else:
        for index, table_data in enumerate(tables):
            if index > 0:
                writer.writerow([])

            header = [str(cell) for cell in table_data.get("header", []) if cell is not None]
            if header:
                writer.writerow(header)

            for row in table_data.get("rows", []):
                writer.writerow([str(cell) for cell in row])

    return buffer.getvalue().encode("utf-8")


def _sanitize_sheet_name(name: str) -> str:
    cleaned = re.sub(r"[\[\]\:\*\?\/\\]+", " ", name).strip()
    cleaned = cleaned[:31] if cleaned else "Sheet"
    return cleaned or "Sheet"


def export_content(content_html: str, target_format: str, file_name: str | None = None) -> ExportArtifact:
    target = _normalize_target_format(target_format)
    base_name = _sanitize_filename(file_name)
    title = base_name.replace("-", " ").strip() or "ISTA AI Export"

    if target == "pdf":
        wrapped_html = _wrap_html(content_html, title)
        return ExportArtifact(
            filename=f"{base_name}.pdf",
            mime_type=PDF_MIME,
            content=WeasyHTML(string=wrapped_html).write_pdf(),
        )

    if target == "docx":
        return ExportArtifact(
            filename=f"{base_name}.docx",
            mime_type=DOCX_MIME,
            content=_render_docx(content_html, title),
        )

    tables = _extract_tables_from_html(content_html)

    if target == "xlsx":
        return ExportArtifact(
            filename=f"{base_name}.xlsx",
            mime_type=XLSX_MIME,
            content=_render_xlsx(tables, content_html, title),
        )

    return ExportArtifact(
        filename=f"{base_name}.csv",
        mime_type=CSV_MIME,
        content=_render_csv(tables, content_html),
    )
