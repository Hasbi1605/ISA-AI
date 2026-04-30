from __future__ import annotations

import csv
import os
import shutil
import subprocess
import tempfile
from io import BytesIO, StringIO
from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt
from openpyxl import Workbook, load_workbook

from app.services.document_content import extract_document_content_html
from app.services.document_export import (
    CSV_MIME,
    DOCX_MIME,
    PDF_MIME,
    XLSX_MIME,
    ExportArtifact,
    _normalize_target_format,
    _render_csv,
    _render_xlsx,
    _sanitize_filename,
    export_content,
)
from app.services.table_extraction import extract_tables_from_file


SUPPORTED_SOURCE_FORMATS = {"pdf", "docx", "xlsx", "csv"}


def convert_document_file(file_path: str, target_format: str, filename: str | None = None, file_name: str | None = None) -> ExportArtifact:
    target = _normalize_target_format(target_format)
    source = _source_format(file_path, filename)
    base_name = _sanitize_filename(file_name or Path(filename or file_path).stem, fallback="dokumen")

    if source not in SUPPORTED_SOURCE_FORMATS:
        raise ValueError(f"Unsupported source format: {source or 'unknown'}")

    if source == target:
        return ExportArtifact(
            filename=f"{base_name}.{target}",
            mime_type=_mime_type_for_format(target),
            content=Path(file_path).read_bytes(),
        )

    if target == "pdf":
        return _convert_to_pdf(file_path, source, filename, base_name)

    if target == "docx":
        return _convert_to_docx(file_path, source, filename, base_name)

    if target == "xlsx":
        return _convert_to_xlsx(file_path, source, filename, base_name)

    return _convert_to_csv(file_path, source, filename, base_name)


def _source_format(file_path: str, filename: str | None) -> str:
    suffix = Path(filename or file_path).suffix.lower().lstrip(".")
    return suffix


def _mime_type_for_format(target: str) -> str:
    return {
        "pdf": PDF_MIME,
        "docx": DOCX_MIME,
        "xlsx": XLSX_MIME,
        "csv": CSV_MIME,
    }[target]


def _convert_to_pdf(file_path: str, source: str, filename: str | None, base_name: str) -> ExportArtifact:
    if source in {"docx", "xlsx", "csv"}:
        try:
            return ExportArtifact(
                filename=f"{base_name}.pdf",
                mime_type=PDF_MIME,
                content=_libreoffice_convert(file_path, "pdf", filename),
            )
        except RuntimeError:
            pass

    content_html = extract_document_content_html(file_path, filename=filename)
    artifact = export_content(content_html, "pdf", base_name)
    return ExportArtifact(filename=f"{base_name}.pdf", mime_type=artifact.mime_type, content=artifact.content)


def _convert_to_docx(file_path: str, source: str, filename: str | None, base_name: str) -> ExportArtifact:
    if source == "pdf":
        try:
            return ExportArtifact(
                filename=f"{base_name}.docx",
                mime_type=DOCX_MIME,
                content=_pdf_to_docx(file_path),
            )
        except Exception:
            pass

        content_html = extract_document_content_html(file_path, filename=filename)
        artifact = export_content(content_html, "docx", base_name)
        return ExportArtifact(filename=f"{base_name}.docx", mime_type=artifact.mime_type, content=artifact.content)

    if source in {"xlsx", "csv"}:
        return ExportArtifact(
            filename=f"{base_name}.docx",
            mime_type=DOCX_MIME,
            content=_spreadsheet_to_docx(file_path, filename),
        )

    try:
        return ExportArtifact(
            filename=f"{base_name}.docx",
            mime_type=DOCX_MIME,
            content=_libreoffice_convert(file_path, "docx", filename),
        )
    except RuntimeError:
        content_html = extract_document_content_html(file_path, filename=filename)
        artifact = export_content(content_html, "docx", base_name)
        return ExportArtifact(filename=f"{base_name}.docx", mime_type=artifact.mime_type, content=artifact.content)


def _convert_to_xlsx(file_path: str, source: str, filename: str | None, base_name: str) -> ExportArtifact:
    if source == "csv":
        return ExportArtifact(
            filename=f"{base_name}.xlsx",
            mime_type=XLSX_MIME,
            content=_csv_to_xlsx(file_path, Path(filename or file_path).stem),
        )

    tables = _tables_from_document(file_path, source)
    content_html = extract_document_content_html(file_path, filename=filename) if not tables else ""

    return ExportArtifact(
        filename=f"{base_name}.xlsx",
        mime_type=XLSX_MIME,
        content=_render_xlsx(tables, content_html, base_name),
    )


def _convert_to_csv(file_path: str, source: str, filename: str | None, base_name: str) -> ExportArtifact:
    if source == "xlsx":
        return ExportArtifact(
            filename=f"{base_name}.csv",
            mime_type=CSV_MIME,
            content=_xlsx_to_csv(file_path),
        )

    tables = _tables_from_document(file_path, source)
    content_html = extract_document_content_html(file_path, filename=filename) if not tables else ""

    return ExportArtifact(
        filename=f"{base_name}.csv",
        mime_type=CSV_MIME,
        content=_render_csv(tables, content_html),
    )


def _tables_from_document(file_path: str, source: str) -> list[dict]:
    if source not in {"pdf", "docx"}:
        return []

    try:
        return extract_tables_from_file(file_path)
    except ValueError:
        return []


def _libreoffice_convert(file_path: str, target_extension: str, filename: str | None = None) -> bytes:
    executable = shutil.which("soffice") or shutil.which("libreoffice")
    if executable is None:
        raise RuntimeError("LibreOffice tidak tersedia.")

    suffix = Path(filename or file_path).suffix or Path(file_path).suffix
    input_name = f"source{suffix}"

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        input_path = temp_path / input_name
        output_dir = temp_path / "out"
        profile_dir = temp_path / "profile"
        output_dir.mkdir()
        shutil.copyfile(file_path, input_path)

        command = [
            executable,
            "--headless",
            "--nologo",
            "--nofirststartwizard",
            "--nodefault",
            "--nolockcheck",
            f"-env:UserInstallation=file://{profile_dir}",
            "--convert-to",
            target_extension,
            "--outdir",
            str(output_dir),
            str(input_path),
        ]

        try:
            completed = subprocess.run(
                command,
                env={**os.environ, "HOME": temp_dir},
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=120,
                check=False,
            )
        except subprocess.TimeoutExpired as exc:
            raise RuntimeError("Konversi LibreOffice melewati batas waktu.") from exc

        if completed.returncode != 0:
            detail = completed.stderr.decode("utf-8", "ignore") or completed.stdout.decode("utf-8", "ignore")
            raise RuntimeError(detail.strip() or "Konversi LibreOffice gagal.")

        candidates = list(output_dir.glob(f"*.{target_extension}"))
        if not candidates:
            detail = completed.stdout.decode("utf-8", "ignore") or completed.stderr.decode("utf-8", "ignore")
            raise RuntimeError(detail.strip() or "File hasil konversi tidak ditemukan.")

        return candidates[0].read_bytes()


def _pdf_to_docx(file_path: str) -> bytes:
    from pdf2docx import Converter

    with tempfile.TemporaryDirectory() as temp_dir:
        output_path = Path(temp_dir) / "converted.docx"
        converter = Converter(file_path)
        try:
            converter.convert(str(output_path), start=0, end=None)
        finally:
            converter.close()

        return output_path.read_bytes()


def _spreadsheet_to_docx(file_path: str, filename: str | None = None) -> bytes:
    sheets = _spreadsheet_rows(file_path, filename)
    document = DocxDocument()
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)

    for sheet_index, (sheet_name, rows) in enumerate(sheets):
        if sheet_index > 0:
            document.add_page_break()

        document.add_heading(sheet_name, level=1)
        normalized_rows = _trim_table(rows)
        if not normalized_rows:
            document.add_paragraph("Tidak ada data.")
            continue

        max_columns = max(len(row) for row in normalized_rows)
        table = document.add_table(rows=len(normalized_rows), cols=max_columns)
        table.style = "Table Grid"
        table.autofit = True

        for row_index, row in enumerate(normalized_rows):
            for col_index in range(max_columns):
                value = row[col_index] if col_index < len(row) else ""
                cell = table.cell(row_index, col_index)
                cell.text = str(value)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _csv_to_xlsx(file_path: str, sheet_name: str) -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = _safe_sheet_name(sheet_name)

    for row_index, row in enumerate(_read_csv_rows(file_path), start=1):
        for col_index, value in enumerate(row, start=1):
            sheet.cell(row=row_index, column=col_index, value=value)

    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def _xlsx_to_csv(file_path: str) -> bytes:
    rows_by_sheet = _spreadsheet_rows(file_path, file_path)
    buffer = StringIO()
    writer = csv.writer(buffer)

    for sheet_index, (sheet_name, rows) in enumerate(rows_by_sheet):
        if sheet_index > 0:
            writer.writerow([])

        if len(rows_by_sheet) > 1:
            writer.writerow([f"Sheet: {sheet_name}"])

        for row in _trim_table(rows):
            writer.writerow(row)

    return buffer.getvalue().encode("utf-8")


def _spreadsheet_rows(file_path: str, filename: str | None) -> list[tuple[str, list[list[str]]]]:
    suffix = Path(filename or file_path).suffix.lower()

    if suffix == ".csv":
        return [(Path(filename or file_path).stem or "CSV", _read_csv_rows(file_path))]

    workbook = load_workbook(file_path, data_only=True, read_only=True)
    try:
        sheets: list[tuple[str, list[list[str]]]] = []
        for sheet in workbook.worksheets:
            rows = [
                ["" if cell is None else str(cell) for cell in row]
                for row in sheet.iter_rows(values_only=True)
            ]
            sheets.append((sheet.title, rows))
        return sheets
    finally:
        workbook.close()


def _read_csv_rows(file_path: str) -> list[list[str]]:
    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            with open(file_path, newline="", encoding=encoding) as handle:
                sample = handle.read(4096)
                handle.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample)
                except csv.Error:
                    dialect = csv.excel

                return [row for row in csv.reader(handle, dialect)]
        except UnicodeDecodeError:
            continue

    return []


def _trim_table(rows: list[list[str]]) -> list[list[str]]:
    cleaned = [[str(cell).strip() for cell in row] for row in rows]

    while cleaned and not any(cleaned[-1]):
        cleaned.pop()

    max_columns = 0
    for row in cleaned:
        for index, value in enumerate(row, start=1):
            if value:
                max_columns = max(max_columns, index)

    return [row[:max_columns] for row in cleaned] if max_columns > 0 else []


def _safe_sheet_name(name: str) -> str:
    cleaned = "".join(" " if char in "[]:*?/\\" else char for char in name).strip()
    return (cleaned or "Sheet")[:31]
