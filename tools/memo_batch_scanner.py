#!/usr/bin/env python3
"""Scan generated memo batches for deterministic official-format regressions."""

from __future__ import annotations

import argparse
import json
import random
import re
import statistics
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

try:
    import pdfplumber
except ImportError:  # pragma: no cover - scanner still works without PDF geometry.
    pdfplumber = None


HEADER_LINES = {
    "KEMENTERIAN SEKRETARIAT NEGARA RI",
    "KEMENTERIAN SEKRETARIAT NEGARA REPUBLIK INDONESIA",
    "SEKRETARIAT PRESIDEN",
    "ISTANA KEPRESIDENAN YOGYAKARTA",
}

KEY_VALUE_LABELS = {
    "hari/tanggal",
    "tanggal",
    "waktu",
    "pukul",
    "jam",
    "tempat",
    "agenda",
    "lokasi",
    "lokasi asal",
    "lokasi tujuan",
    "periode",
    "jadwal",
    "waktu pelaksanaan",
    "waktu kejadian",
    "batas waktu",
    "nama",
    "nip",
    "jabatan",
    "unit kerja",
    "nomor kontak",
    "peserta",
}

METADATA_LABEL_KEYS = {"yth", "dari", "hal", "tanggal"}

LABEL_ALIASES = {
    "waktu": "pukul",
    "jam": "pukul",
    "tanggal": "hari/tanggal",
}

HIGH_RISK_EVAL_NUMBERS = {"EVAL-05", "EVAL-06", "EVAL-12", "EVAL-13", "EVAL-19", "EVAL-21", "EVAL-39"}

PROMPT_LEAK_PATTERNS = {
    "catatan": re.compile(r"\bcatatan(?:\s+internal)?\s*:", re.I),
    "revision_instruction": re.compile(r"\binstruksi\s+revisi\b", re.I),
    "format_instruction": re.compile(r"\bsesuai\s+format\s+yang\s+berlaku\b", re.I),
    "manual_closing_instruction": re.compile(r"\bpenutup\s+manual\b.*\bapa\s+adanya\b", re.I),
    "preserve_manual_closing": re.compile(r"\b(?:pertahankan|dipertahankan)\s+penutup\s+manual\b", re.I),
    "sources_artifact": re.compile(r"\[SOURCES:|```|^\s*[-*]\s+", re.I | re.M),
}

CLOSING_PATTERN = re.compile(
    r"^(?:demikian|atas\s+perhatian|atas\s+kerja\s+sama|dimohon|mohon\s+tindak\s+lanjut|"
    r"partisipasi\s+aktif)",
    re.I,
)

TIME_PATTERN = re.compile(
    r"\b\d{1,2}[.:]\d{2}(?:\s*(?:s\.?d\.?|-|hingga|sampai)\s*\d{1,2}[.:]\d{2})?\s*WIB\b",
    re.I,
)


@dataclass
class VersionEntry:
    memo_id: int
    title: str
    version_number: int
    export_name: str
    docx_name: str
    configuration: dict[str, Any]
    searchable_text: str
    revision_instruction: str | None


def normalize(text: str) -> str:
    text = (text or "").lower()
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^0-9a-zà-ÿ/.: -]+", "", text)
    return text.strip()


def compact(text: str) -> str:
    return re.sub(r"[^0-9a-zà-ÿ]+", "", normalize(text))


def label_key(label: str) -> str:
    key = normalize(label).rstrip(".")
    return LABEL_ALIASES.get(key, key)


def load_versions(config_path: Path) -> list[VersionEntry]:
    payload = json.loads(config_path.read_text())
    versions: list[VersionEntry] = []
    for memo in payload.get("memos", []):
        for version in memo.get("versions", []):
            versions.append(
                VersionEntry(
                    memo_id=int(memo["id"]),
                    title=str(memo["title"]),
                    version_number=int(version["version_number"]),
                    export_name=str(version["export_name"]),
                    docx_name=str(version["docx_name"]),
                    configuration=version.get("configuration") or {},
                    searchable_text=str(version.get("searchable_text") or ""),
                    revision_instruction=version.get("revision_instruction"),
                )
            )
    return versions


def paragraph_alignment_name(paragraph) -> str:
    alignment = paragraph.alignment
    if alignment is None:
        return "INHERITED"
    if alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
        return "JUSTIFY"
    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return "CENTER"
    if alignment == WD_ALIGN_PARAGRAPH.RIGHT:
        return "RIGHT"
    if alignment == WD_ALIGN_PARAGRAPH.LEFT:
        return "LEFT"
    return str(alignment)


def spacing_before_twips(paragraph) -> int:
    spacing = paragraph._p.get_or_add_pPr().find(qn("w:spacing"))
    if spacing is None:
        return 0
    return int(spacing.get(qn("w:before"), "0"))


def spacing_after_twips(paragraph) -> int:
    spacing = paragraph._p.get_or_add_pPr().find(qn("w:spacing"))
    if spacing is None:
        return 0
    return int(spacing.get(qn("w:after"), "0"))


def is_numbered(text: str) -> bool:
    return re.match(r"^\s*\d+[.)]\s+|^\s*\d+[.)]\t", text) is not None


def is_bullet(text: str) -> bool:
    return re.match(r"^\s*[-*•]\s+", text) is not None


def numbered_value(text: str) -> int | None:
    match = re.match(r"^\s*(\d+)[.)](?:\s+|\t)", text)
    return int(match.group(1)) if match else None


def clean_cell_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def docx_table_rows(document: Document) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        for row in table.rows:
            cells = [clean_cell_text(cell.text) for cell in row.cells]
            if len(cells) < 3:
                continue
            label = cells[0].rstrip(":").strip()
            value = cells[2].strip()
            rows.append(
                {
                    "table_index": table_index,
                    "label": label,
                    "label_key": label_key(label),
                    "value": value,
                }
            )
    return rows


def body_key_value_rows(table_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        row
        for row in table_rows
        if row["table_index"] > 0
        and row["label_key"] in KEY_VALUE_LABELS
        and row["label_key"] not in METADATA_LABEL_KEYS
    ]


def body_paragraphs(document: Document, signatory: str) -> list[dict[str, Any]]:
    paragraphs: list[dict[str, Any]] = []
    in_carbon_copy = False
    footer_markers = (
        "Dokumen ini telah ditandatangani",
        "yang diterbitkan oleh Balai Sertifikasi",
    )
    for index, paragraph in enumerate(document.paragraphs):
        text = clean_cell_text(paragraph.text)
        if not text:
            continue
        if text in HEADER_LINES:
            continue
        if text == "MEMORANDUM" or text.startswith("Nomor "):
            continue
        if text == "Tembusan:":
            in_carbon_copy = True
            continue
        if in_carbon_copy:
            continue
        if text in {"QR", "TTE"}:
            continue
        if signatory and text == signatory:
            continue
        if any(text.startswith(marker) for marker in footer_markers):
            continue
        paragraphs.append(
            {
                "index": index,
                "text": text,
                "alignment": paragraph_alignment_name(paragraph),
                "space_before_twips": spacing_before_twips(paragraph),
                "space_after_twips": spacing_after_twips(paragraph),
                "number": numbered_value(text),
                "is_numbered": is_numbered(text),
                "is_bullet": is_bullet(text),
            }
        )
    return paragraphs


def extract_config_text(configuration: dict[str, Any]) -> str:
    return "\n".join(str(value) for value in configuration.values() if value is not None)


def check_duplicate_key_values(paragraphs: list[dict[str, Any]], table_rows: list[dict[str, Any]]) -> list[dict[str, str]]:
    kv_rows = body_key_value_rows(table_rows)
    configured_values = [compact(row["value"]) for row in kv_rows if len(compact(row["value"])) >= 4]
    findings: list[dict[str, str]] = []
    for paragraph in paragraphs:
        text = paragraph["text"]
        normalized_text = normalize(text)
        compact_text = compact(text)
        if not (paragraph["is_numbered"] or paragraph["is_bullet"] or ":" in text or "pada:" in normalized_text):
            continue
        for row in kv_rows:
            label = row["label_key"]
            value = compact(row["value"])
            label_hits = {label}
            if label == "pukul":
                label_hits.update({"waktu", "jam"})
            if label == "hari/tanggal":
                label_hits.add("tanggal")
            has_label = any(hit in normalized_text for hit in label_hits)
            has_value = value and value in compact_text
            shared_values = sum(1 for configured_value in configured_values if configured_value and configured_value in compact_text)
            if (has_label and (has_value or shared_values >= 1)) or shared_values >= 2:
                findings.append(
                    {
                        "paragraph": text,
                        "label": row["label"],
                        "value": row["value"],
                    }
                )
                break
    return findings


def check_numbering_restart(paragraphs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    seen_numbers: list[int] = []
    for paragraph in paragraphs:
        number = paragraph["number"]
        if number is None:
            continue
        if seen_numbers and number == 1:
            findings.append(
                {
                    "sequence": [*seen_numbers, number],
                    "paragraph": paragraph["text"],
                }
            )
        seen_numbers.append(number)
    return findings


def check_prompt_leakage(body_text: str) -> list[dict[str, str]]:
    findings: list[dict[str, str]] = []
    for name, pattern in PROMPT_LEAK_PATTERNS.items():
        for match in pattern.finditer(body_text):
            snippet = body_text[max(0, match.start() - 60) : match.end() + 60]
            findings.append({"type": name, "snippet": re.sub(r"\s+", " ", snippet).strip()})
    return findings


def check_unconfigured_facts(body_text: str, configuration: dict[str, Any]) -> list[dict[str, str]]:
    config_text = normalize(extract_config_text(configuration))
    compact_config = compact(config_text)
    findings: list[dict[str, str]] = []

    for time in sorted(set(TIME_PATTERN.findall(body_text))):
        if compact(time) not in compact_config:
            findings.append({"type": "time_not_in_config", "value": time})

    if "pic" not in config_text and "penanggung jawab" not in config_text:
        if re.search(r"\b(?:PIC|penanggung jawab|Person in Charge)\b", body_text):
            findings.append({"type": "pic_not_in_config", "value": "PIC/penanggung jawab"})

    staff_patterns = (
        r"\bseluruh\s+staf\b",
        r"\bstaf\s+persuratan\b",
        r"\boperator\s+persuratan\b",
    )
    for pattern in staff_patterns:
        match = re.search(pattern, body_text, flags=re.I)
        if match and compact(match.group(0)) not in compact_config:
            findings.append({"type": "staff_or_unit_not_in_config", "value": match.group(0)})

    honorific_match = re.search(r"\b(?:Bapak|Ibu|Sdr\.?)\s+[A-Z][A-Za-zÀ-ÿ]+", body_text)
    if honorific_match and compact(honorific_match.group(0)) not in compact_config:
        findings.append({"type": "honorific_name_not_in_config", "value": honorific_match.group(0)})

    return findings


def check_closing(paragraphs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    closing_paragraphs = [p for p in paragraphs if CLOSING_PATTERN.search(p["text"])]
    findings: list[dict[str, Any]] = []
    if not closing_paragraphs:
        findings.append({"type": "missing_closing", "detail": "No formal closing paragraph detected."})
        return findings

    for paragraph in closing_paragraphs:
        if paragraph["is_numbered"] or paragraph["is_bullet"]:
            findings.append({"type": "closing_inside_list", "paragraph": paragraph["text"]})
        if paragraph["space_before_twips"] < 160:
            findings.append(
                {
                    "type": "closing_spacing_too_tight",
                    "paragraph": paragraph["text"],
                    "space_before_twips": paragraph["space_before_twips"],
                }
            )
    return findings


def check_alignment(paragraphs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    for paragraph in paragraphs:
        text = paragraph["text"]
        if paragraph["is_numbered"] or paragraph["is_bullet"]:
            continue
        if len(text) < 24:
            continue
        if paragraph["alignment"] != "JUSTIFY":
            findings.append({"paragraph": text, "alignment": paragraph["alignment"]})
    return findings


def check_docx_layout(document: Document, table_rows: list[dict[str, Any]]) -> list[dict[str, Any]]:
    findings: list[dict[str, Any]] = []
    key_value_table_indexes = {row["table_index"] for row in body_key_value_rows(table_rows)}
    for table_index, table in enumerate(document.tables):
        if table_index not in key_value_table_indexes:
            continue
        table_borders = table._tbl.tblPr.find(qn("w:tblBorders"))
        if table_borders is not None:
            findings.append({"type": "key_value_table_has_borders", "table_index": table_index})
        next_element = table._tbl.getnext()
        if next_element is not None and next_element.tag.endswith("}p"):
            spacing = next_element.find(qn("w:pPr"))
            spacing = spacing.find(qn("w:spacing")) if spacing is not None else None
            after = int(spacing.get(qn("w:after"), "0")) if spacing is not None else 0
            if after < 160:
                findings.append({"type": "key_value_spacing_after_too_tight", "table_index": table_index, "after": after})
    metadata_rows = [row for row in table_rows if row["table_index"] == 0]
    if metadata_rows:
        metadata_labels = {row["label"].lower().rstrip(".") for row in metadata_rows}
        if not {"yth", "dari", "hal", "tanggal"} & metadata_labels:
            findings.append({"type": "metadata_table_unexpected_labels", "labels": sorted(metadata_labels)})
    return findings


def pdf_metrics(pdf_path: Path) -> dict[str, Any] | None:
    if pdfplumber is None or not pdf_path.exists():
        return None
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            page = pdf.pages[0]
            words = page.extract_words() or []
            lines = page.lines or []
            long_lines = [line for line in lines if abs(line.get("x1", 0) - line.get("x0", 0)) >= 250]
            separator_y = min((line["top"] for line in long_lines), default=None)
            body_words = [word for word in words if separator_y is not None and word["top"] > separator_y + 10]
            signature_words = [word for word in words if word["text"] == "Deni" or word["text"] == "Mulyana"]
            return {
                "page_width": round(page.width, 1),
                "page_height": round(page.height, 1),
                "first_word_y": round(min((word["top"] for word in words), default=0), 1),
                "separator_y": round(separator_y, 1) if separator_y is not None else None,
                "body_start_y": round(min((word["top"] for word in body_words), default=0), 1) if body_words else None,
                "separator_to_body_gap": round(min((word["top"] for word in body_words), default=0) - separator_y, 1)
                if separator_y is not None and body_words
                else None,
                "signature_y": round(min((word["top"] for word in signature_words), default=0), 1)
                if signature_words
                else None,
            }
    except Exception as exc:  # pragma: no cover - report scanner errors without failing.
        return {"error": str(exc)}


def official_baseline(official_dir: Path | None) -> dict[str, Any]:
    if not official_dir or not official_dir.exists() or pdfplumber is None:
        return {}
    metrics = [pdf_metrics(path) for path in sorted(official_dir.glob("*.pdf"))]
    metrics = [metric for metric in metrics if metric and "error" not in metric]
    baseline: dict[str, Any] = {"count": len(metrics)}
    for key in ("first_word_y", "separator_y", "body_start_y", "separator_to_body_gap", "signature_y"):
        values = [metric[key] for metric in metrics if isinstance(metric.get(key), (int, float)) and metric.get(key)]
        if values:
            baseline[key] = {
                "median": round(statistics.median(values), 1),
                "min": round(min(values), 1),
                "max": round(max(values), 1),
            }
    return baseline


def compare_pdf_to_official(metric: dict[str, Any] | None, baseline: dict[str, Any]) -> list[dict[str, Any]]:
    if not metric or not baseline:
        return []
    findings: list[dict[str, Any]] = []
    tolerances = {
        "first_word_y": 45,
        "separator_y": 90,
        "separator_to_body_gap": 45,
        "signature_y": 130,
    }
    for key, tolerance in tolerances.items():
        value = metric.get(key)
        base = baseline.get(key, {})
        median = base.get("median") if isinstance(base, dict) else None
        if not isinstance(value, (int, float)) or not isinstance(median, (int, float)):
            continue
        delta = value - median
        if abs(delta) > tolerance:
            findings.append({"type": f"pdf_{key}_outside_tolerance", "value": value, "official_median": median, "delta": round(delta, 1)})
    return findings


def scan_version(batch_dir: Path, entry: VersionEntry, official: dict[str, Any]) -> dict[str, Any]:
    docx_path = batch_dir / "docx" / entry.docx_name
    pdf_path = batch_dir / entry.export_name
    result: dict[str, Any] = {
        "memo_id": entry.memo_id,
        "title": entry.title,
        "version": entry.version_number,
        "docx": entry.docx_name,
        "pdf": entry.export_name,
        "eval_number": str(entry.configuration.get("number", "")),
        "checks": {},
    }
    if not docx_path.exists():
        result["checks"]["missing_docx"] = [{"path": str(docx_path)}]
        return result

    document = Document(str(docx_path))
    signatory = str(entry.configuration.get("signatory", ""))
    paragraphs = body_paragraphs(document, signatory)
    table_rows = docx_table_rows(document)
    body_text = "\n".join(paragraph["text"] for paragraph in paragraphs)
    pdf_metric = pdf_metrics(pdf_path)

    checks = {
        "duplicate_key_values": check_duplicate_key_values(paragraphs, table_rows),
        "numbering_restart": check_numbering_restart(paragraphs),
        "prompt_leakage": check_prompt_leakage(body_text),
        "unconfigured_facts": check_unconfigured_facts(body_text, entry.configuration),
        "closing": check_closing(paragraphs),
        "alignment": check_alignment(paragraphs),
        "layout_docx": check_docx_layout(document, table_rows),
        "layout_pdf_tolerance": compare_pdf_to_official(pdf_metric, official),
    }
    result["checks"] = checks
    result["pdf_metrics"] = pdf_metric
    result["key_value_labels"] = sorted({row["label_key"] for row in body_key_value_rows(table_rows)})
    return result


def has_findings(scan: dict[str, Any]) -> bool:
    return any(scan["checks"].get(name) for name in scan.get("checks", {}))


def high_risk_selection(entries: list[VersionEntry], scans: list[dict[str, Any]]) -> list[dict[str, Any]]:
    by_docx = {scan["docx"]: scan for scan in scans}
    selected: dict[str, set[str]] = {}

    def add(entry: VersionEntry, reason: str) -> None:
        selected.setdefault(entry.docx_name, set()).add(reason)

    for entry in entries:
        number = str(entry.configuration.get("number", ""))
        scan = by_docx.get(entry.docx_name, {})
        labels = set(scan.get("key_value_labels", []))
        if any(marker in number for marker in HIGH_RISK_EVAL_NUMBERS) or labels:
            add(entry, "key-value")
        if entry.version_number > 1:
            add(entry, "revision-v2")
        if str(entry.configuration.get("page_size", "")).lower() == "folio" or "folio" in entry.title.lower() or "panjang" in entry.title.lower():
            add(entry, "long-or-folio")
        if str(entry.configuration.get("carbon_copy", "")).count("\n") >= 1 or "tembusan" in entry.title.lower():
            add(entry, "many-carbon-copy")

    selected_names = set(selected)
    rest = [entry for entry in entries if entry.docx_name not in selected_names]
    random.Random(10).shuffle(rest)
    for entry in rest[:5]:
        add(entry, "deterministic-sample")

    output = []
    for entry in entries:
        reasons = selected.get(entry.docx_name)
        if reasons:
            output.append(
                {
                    "memo_id": entry.memo_id,
                    "title": entry.title,
                    "version": entry.version_number,
                    "docx": entry.docx_name,
                    "reasons": sorted(reasons),
                }
            )
    return output


def regression_memory(regression_root: Path) -> dict[str, Any]:
    batches: dict[str, Any] = {}
    for batch_dir in sorted(regression_root.glob("batch *")):
        text_dir = batch_dir / "text"
        text_files = list(text_dir.glob("*.txt")) if text_dir.exists() else list(batch_dir.glob("*.txt"))
        counts = {name: 0 for name in PROMPT_LEAK_PATTERNS}
        counts.update({"pukul_zero": 0, "dash_bullet": 0, "ordinal_words": 0})
        for path in text_files:
            text = path.read_text(errors="ignore")
            for name, pattern in PROMPT_LEAK_PATTERNS.items():
                if pattern.search(text):
                    counts[name] += 1
            if re.search(r"\bpukul\s*[: ]\s*00\s+WIB\b", text, re.I):
                counts["pukul_zero"] += 1
            if re.search(r"(?m)^\s*-\s+\S+", text):
                counts["dash_bullet"] += 1
            if re.search(r"\b(?:Pertama|Kedua|Ketiga|Keempat)\b", text):
                counts["ordinal_words"] += 1
        batches[batch_dir.name] = {"text_files": len(text_files), "pattern_counts": counts}
    return batches


def summarize(scans: list[dict[str, Any]]) -> dict[str, Any]:
    check_counts: dict[str, int] = {}
    blocker_docs: list[str] = []
    for scan in scans:
        for name, findings in scan.get("checks", {}).items():
            count = len(findings)
            check_counts[name] = check_counts.get(name, 0) + count
            if count and name in {"duplicate_key_values", "numbering_restart", "prompt_leakage", "closing", "alignment"}:
                blocker_docs.append(scan["docx"])
    return {
        "documents": len(scans),
        "documents_with_findings": sum(1 for scan in scans if has_findings(scan)),
        "blocker_documents": sorted(set(blocker_docs)),
        "check_counts": check_counts,
    }


def write_markdown(report: dict[str, Any], output_path: Path) -> None:
    lines = [
        "# Memo Batch Auto Scan",
        "",
        f"Batch: `{report['batch_dir']}`",
        f"Documents scanned: {report['summary']['documents']}",
        f"Documents with findings: {report['summary']['documents_with_findings']}",
        "",
        "## Check Counts",
        "",
    ]
    for name, count in sorted(report["summary"]["check_counts"].items()):
        lines.append(f"- `{name}`: {count}")
    lines.extend(["", "## Blocker Documents", ""])
    if report["summary"]["blocker_documents"]:
        for docx in report["summary"]["blocker_documents"]:
            lines.append(f"- `{docx}`")
    else:
        lines.append("- None")
    lines.extend(["", "## High-Risk Manual Review Queue", ""])
    for item in report["high_risk_manual_review"]:
        lines.append(f"- `{item['docx']}`: {', '.join(item['reasons'])}")
    lines.extend(["", "## Top Findings", ""])
    for scan in report["scans"]:
        if not has_findings(scan):
            continue
        lines.append(f"### {scan['docx']}")
        for name, findings in scan["checks"].items():
            if findings:
                lines.append(f"- `{name}`: {len(findings)}")
                for finding in findings[:3]:
                    lines.append(f"  - `{json.dumps(finding, ensure_ascii=False)[:300]}`")
        lines.append("")
    output_path.write_text("\n".join(lines).rstrip() + "\n")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--batch-dir", required=True, type=Path)
    parser.add_argument("--official-dir", type=Path)
    parser.add_argument("--regression-root", type=Path, default=Path("evaluation/memo"))
    parser.add_argument("--output-json", type=Path)
    parser.add_argument("--output-md", type=Path)
    args = parser.parse_args()

    batch_dir = args.batch_dir
    config_path = batch_dir / "CONFIG_AND_REVISION_PROMPTS_ACTUAL.json"
    if not config_path.exists():
        config_path = batch_dir / "CONFIG_AND_REVISION_PROMPTS.json"
    entries = load_versions(config_path)
    official = official_baseline(args.official_dir)
    scans = [scan_version(batch_dir, entry, official) for entry in entries]
    report = {
        "batch_dir": str(batch_dir),
        "config_path": str(config_path),
        "official_baseline": official,
        "summary": summarize(scans),
        "high_risk_manual_review": high_risk_selection(entries, scans),
        "regression_memory": regression_memory(args.regression_root),
        "scans": scans,
    }

    output_json = args.output_json or (batch_dir / "AUTO_SCAN_REPORT.json")
    output_md = args.output_md or (batch_dir / "AUTO_SCAN_REPORT.md")
    output_json.write_text(json.dumps(report, ensure_ascii=False, indent=2))
    write_markdown(report, output_md)
    print(json.dumps(report["summary"], ensure_ascii=False, indent=2))
    print(f"Wrote {output_json}")
    print(f"Wrote {output_md}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
